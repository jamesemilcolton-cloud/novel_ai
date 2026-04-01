from __future__ import annotations

import os
import re
import inspect
import math
import shutil
import subprocess
import sys
import threading
import time
from collections import Counter, OrderedDict
from copy import deepcopy
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any, Callable


# ============================================================
# Configuration
# ============================================================

WRITING_DIR = Path.home() / "writing"
NOVEL_AI_SCRIPT_PATH = WRITING_DIR / "ai" / "novel_ai" / "novel_ai.py"
NOVEL_PROJECT_DIR = WRITING_DIR / "novel_project"
PROJECT_MEMORY_DIR = NOVEL_PROJECT_DIR / "memory"
PROJECT_ANALYSIS_DIR = NOVEL_PROJECT_DIR / "analysis"
CHAPTERS_DIR = NOVEL_PROJECT_DIR / "chapters"
MANUSCRIPT_DIR = NOVEL_PROJECT_DIR / "manuscript"
MANUSCRIPT_PATH = MANUSCRIPT_DIR / "novel.txt"
CONTINUITY_REPORTS_DIR = PROJECT_ANALYSIS_DIR / "continuity_reports"
TIMELINE_LOGS_DIR = PROJECT_ANALYSIS_DIR / "timeline_logs"
BOOK_INTEGRITY_REPORTS_DIR = PROJECT_ANALYSIS_DIR / "book_integrity_reports"
REBUILD_LOG_DIR = PROJECT_ANALYSIS_DIR / "rebuild_logs"
FULL_NOVEL_PROCESSOR_LOG_DIR = NOVEL_PROJECT_DIR / "logs"
FULL_NOVEL_PROCESSOR_LOG_PATH = FULL_NOVEL_PROCESSOR_LOG_DIR / "processor_log.txt"
DRAFTS_DIR = NOVEL_PROJECT_DIR / "drafts"
PROJECT_BACKUPS_DIR = NOVEL_PROJECT_DIR / "backups"
CANON_MEMORY_BACKUPS_DIR = PROJECT_MEMORY_DIR / "backups"
RESEARCH_DIR = NOVEL_PROJECT_DIR / "research"
INSPIRATIONS_DIR = NOVEL_PROJECT_DIR / "inspirations"
RESEARCH_INTEGRITY_REPORTS_DIR = RESEARCH_DIR / "integrity_reports"
WORLD_PLAUSIBILITY_REPORTS_DIR = PROJECT_ANALYSIS_DIR / "world_plausibility_reports"
CANON_MEMORY_PATH = PROJECT_MEMORY_DIR / "canon_memory.txt"
CONTINUITY_INDEX_PATH = PROJECT_MEMORY_DIR / "continuity_index.txt"
UNPARSED_MEMORY_SUGGESTIONS_LOG_PATH = PROJECT_ANALYSIS_DIR / "unparsed_memory_suggestions.log"
SCENE_SUMMARIES_PATH = PROJECT_MEMORY_DIR / "scene_summaries.txt"
IDEAS_PATH = PROJECT_MEMORY_DIR / "ideas.txt"
WORLD_RULES_PATH = PROJECT_MEMORY_DIR / "world.txt"
STORY_STATE_PATH = PROJECT_MEMORY_DIR / "story_state.txt"
TIMELINE_THREADS_PATH = PROJECT_MEMORY_DIR / "timeline_threads.txt"
CHAPTER_FILENAME_PATTERN = re.compile(r"chapter_(\d+)\.txt$")
SUGGESTION_PATTERN = re.compile(
    r"^\s*(\d+)\.\s*(.+?)\s*(?:→|->)\s*\[([^\]]+)\]\s*$",
    re.MULTILINE,
)
CHAPTER_HEADER_PATTERN = re.compile(r"^CHAPTER\s+(\d+)\s*$")
CATEGORY_HEADER_PATTERN = re.compile(r"^\[(.+)\]\s*$")
FACT_STATE_PATTERN = re.compile(r"^(.*?)(?:\s*\((ACTIVE|RESOLVED)\))?$")
ANSI_ESCAPE_PATTERN = re.compile(r"\033\[[0-9;]*[A-Za-z]")
BRACKETED_PASTE_PATTERN = re.compile(r"(?:\033\[|\^\[\[?)(?:200~|201~|E)|\[\[200~|\[\[201~")
DISALLOWED_CONTROL_CHAR_PATTERN = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")

MODEL_NAME = os.getenv("OPENAI_MODEL", "gpt-4.1-mini")
MAIN_TEMPERATURE = 0.8
SCENE_TEMPERATURE = 0.1
CONTINUITY_TEMPERATURE = 0.0
PROOFREAD_TEMPERATURE = 0.1
IDEA_RESURFACE_TEMPERATURE = 0.3
DRAFT_PASS_TEMPERATURE = 0.2
RESEARCH_TEMPERATURE = 0.0
RESEARCH_SCENE_TEMPERATURE = 0.0
RECAP_TEMPERATURE = 0.3

CONTINUITY_CHAPTER_WINDOW = 3
MAX_SCENE_SUMMARIES = 5
MAX_CONVERSATION_TURNS = 6
MAX_CANON_CHARACTERS = 12000
DEFAULT_ANALYSIS_CHUNK_WORDS = 6000

SPINNER_FRAMES = ["|", "/", "-", "\\"]
SPINNER_INTERVAL_SECONDS = 0.1

MAIN_SYSTEM_PROMPT = """You are a thoughtful AI novel-writing assistant.
Help the user think through story ideas, scenes, structure, tone, character, and prose.
Use the provided memory carefully and naturally.
Be creative, clear, and practical.
Do not invent persistent facts unless the user states them.
"""

RECAP_SYSTEM_PROMPT = """You are a calm narrative orientation assistant.

Your task is to help the writer remember where the story currently is.

You must:

- describe current narrative situation
- reflect emotional tone
- identify immediate story focus
- mention active unresolved pressures
- describe story momentum

You must NOT:

- give writing advice
- suggest future plot directions
- critique pacing
- analyse themes
- perform continuity checking
- mention screenplay alignment
- invent new story facts

Tone must feel:

clear
grounded
present-focused
mentally stabilising

Return output in this exact structure:

STORY RECAP

Current Situation
- ...

Immediate Focus
- ...

Emotional Atmosphere
- ...

Active Narrative Pressures
- ...

Momentum Direction
- ...
"""

INSPIRATION_ANALYSIS_SYSTEM_PROMPT = """You are a professional writing technique analyst.

You are given structured inspiration data grouped by writing technique.

Your job is to:

1) Extract clear technique patterns from EACH category
2) Compare those patterns to the user's writing
3) Explain how to apply those techniques
4) Score technical alignment by category
5) Identify the biggest technique gap

-----------------------------------------------------

RULES:

- Evaluate how closely the user's writing aligns with techniques found in inspiration data
- Do NOT judge writing quality subjectively
- Do NOT give vague praise
- DO NOT rewrite the user's text
- DO NOT suggest plot ideas
- DO NOT create new story content
- DO NOT copy inspiration text directly
- Keep scores consistent with analysis
- Do not inflate scores

-----------------------------------------------------

SCORING RULES:

- Scores represent ALIGNMENT with inspiration techniques only
- Scores are NOT general writing quality
- Use full 0-10 range when warranted; avoid clustering around 6-8
- Be honest and consistent
- You must determine whether each category is present in the user's text.
- If a category is not present:
  - mark it as N/A
  - do not analyse it
  - do not include it in scoring
- Do not guess or force analysis where none exists.

-----------------------------------------------------

FOCUS AREAS:

PROSE:
- sentence rhythm
- sentence length variation
- tone

DIALOGUE:
- subtext
- conflict through speech
- indirect communication

DESCRIPTION:
- sensory detail
- environmental influence
- mood creation

PACING:
- sentence speed
- paragraph density
- action vs reflection balance

TENSION:
- stakes clarity
- pressure building
- sentence tightening

DEVICES:
- structural techniques
- narrative framing
- stylistic tools

-----------------------------------------------------

OUTPUT FORMAT:

INSPIRATION ANALYSIS

PROSE:
- pattern
- pattern

DIALOGUE:
- pattern

DESCRIPTION:
- pattern

PACING:
- pattern

TENSION:
- pattern

DEVICES:
- pattern

---

YOUR TEXT ANALYSIS:
- issue
- issue

---

ALIGNMENT SCORES (0-10):

PROSE: X/10
- short explanation

DIALOGUE: X/10 or N/A
- short explanation (for N/A use reason like "no dialogue present")

DESCRIPTION: X/10 or N/A
- short explanation (for N/A include reason)

PACING: X/10
- short explanation

TENSION: X/10 or N/A
- short explanation (for N/A include reason)

DEVICES: X/10 or N/A
- short explanation (for N/A use reason like "no structural devices used")

---

OVERALL ALIGNMENT:

X.X / 10

This is an average of numeric category scores only (ignore N/A categories).

---

BIGGEST GAP:

- Identify the single weakest category
- Explain clearly what is missing compared to inspiration techniques

---

HOW TO APPLY:
- actionable technique
- actionable technique

---

EXAMPLES:
- short quote from user text
- explain technique (DO NOT rewrite)
"""

INSPIRATION_DIALOGUE_MARKERS = ('"', "“", "”")
INSPIRATION_DEVICES_MARKERS = (
    "epigraph",
    "journal entry",
    "log entry",
    "transcript",
    "interlude",
    "appendix",
    "footnote",
    "[",
    "]",
    "***",
    "---",
)
INSPIRATION_SENSORY_TERMS = (
    "saw",
    "seen",
    "looked",
    "glow",
    "dark",
    "light",
    "heard",
    "sound",
    "noise",
    "silent",
    "smell",
    "scent",
    "taste",
    "felt",
    "cold",
    "hot",
    "warm",
    "rough",
    "smooth",
    "wind",
    "rain",
    "dust",
    "air",
    "room",
    "street",
    "forest",
    "metal",
)
INSPIRATION_TENSION_TERMS = (
    "risk",
    "danger",
    "threat",
    "deadline",
    "urgent",
    "uncertain",
    "if",
    "might",
    "could",
    "afraid",
    "fear",
    "pressure",
    "stakes",
    "before",
    "or else",
)

RESEARCH_SCENE_SYSTEM_PROMPT = """You are a hard-science realism consultant.

You must analyse:

- physics accuracy
- engineering feasibility
- environmental realism
- energy requirements
- scale realism
- survivability factors
- technological plausibility

You must NOT:

- write story prose
- suggest narrative changes
- judge writing quality
- invent fictional science unless clearly labelled theoretical

Be factual, structured and concise."""

WORLD_CONSISTENCY_CHUNK_SYSTEM_PROMPT = """You are a strict science-fiction world logic auditor.

Analyse ONLY for factual world consistency problems.

Focus on:
technology limits,
environment realism,
space mechanics,
energy logic,
AI behavioural consistency,
authority structures,
tone drift,
scale inflation.

Do NOT:
give writing advice,
rewrite text,
comment on prose quality.

If no issue in this chunk output EXACTLY:

NO WORLD ISSUES IN THIS CHUNK

Otherwise output:

WORLD ISSUE
- description
"""

WORLD_CONSISTENCY_SYNTHESIS_SYSTEM_PROMPT = """You are a senior world consistency auditor.

Combine all chunk findings into one final report.

Remove duplicates.
Group related problems.
Be concise.

If no issues exist output EXACTLY:

WORLD CONSISTENCY REPORT

World logic remains consistent.

Otherwise output:

WORLD CONSISTENCY REPORT

- issue
- issue
"""

CHARACTER_CONSISTENCY_CHUNK_SYSTEM_PROMPT = """You are a strict psychological continuity auditor.

Analyse ONLY behavioural consistency.

Focus on:
personality stability,
motivation logic,
emotional continuity,
injury behaviour,
relationship logic,
competence realism,
authority behaviour,
knowledge continuity.

Do NOT:
give writing advice,
rewrite scenes,
comment on prose quality.

If no issue exists in this chunk output EXACTLY:

NO CHARACTER ISSUES IN THIS CHUNK

Otherwise output:

CHARACTER ISSUE
- description
"""

CHARACTER_CONSISTENCY_SYNTHESIS_SYSTEM_PROMPT = """You are a senior character continuity auditor.

Combine all chunk findings into one final report.

- Remove duplicates
- Merge related issues
- Keep concise
- Preserve factual tone

If no issues exist output EXACTLY:

CHARACTER CONSISTENCY REPORT

Character behaviour remains consistent.

Otherwise output:

CHARACTER CONSISTENCY REPORT

- issue
- issue
"""

RESEARCH_DEPTH_OPTIONS: dict[str, str] = {
    "1": "Surface realism",
    "2": "Hard sci-fi realism",
    "3": "Ultra deep technical realism",
}

RESEARCH_STYLE_OPTIONS: dict[str, str] = {
    "1": "Scientific report",
    "2": "Teaching explanation",
    "3": "Practical notes",
}

SCREENPLAY_SOURCE_PATH = NOVEL_PROJECT_DIR / "sources" / "screenplay.pdf"
ALLOWED_MEMORY_CATEGORIES = (
    "Character",
    "Relationship",
    "World",
    "Location",
    "Object",
    "Timeline",
    "Injury",
    "Mission State — Active",
    "Mission State — Resolved",
    "Psychological State — Active",
    "Psychological State — Resolved",
    "Relationship State — Active",
    "Relationship State — Resolved",
    "Technology State — Active",
    "Technology State — Resolved",
    "Foreshadowing Setup",
    "Foreshadowing Payoff",
)


ACTIVE_TO_RESOLVED_CATEGORY = {
    "Mission State — Active": "Mission State — Resolved",
    "Psychological State — Active": "Psychological State — Resolved",
    "Relationship State — Active": "Relationship State — Resolved",
    "Technology State — Active": "Technology State — Resolved",
}

STORY_STATE_SIGNAL_KEYWORDS = (
    "urgency",
    "urgent",
    "countdown",
    "deadline",
    "danger",
    "threat",
    "risk",
    "unstable",
    "instability",
    "worsening",
    "deteriorating",
    "hidden",
    "secret",
    "withheld",
    "pressure",
    "strain",
    "conflict",
    "uncertainty",
    "hazard",
    "environmental",
    "approaching",
    "mission",
    "failure",
    "collapse",
    "panic",
    "fear",
    "tension",
    "mystery",
    "exposed",
    "running out",
    "time pressure",
)

STORY_STATE_PRIORITY_CATEGORIES = {
    "Injury",
    "Mission State — Active",
    "Mission State — Resolved",
    "Psychological State — Active",
    "Psychological State — Resolved",
    "Relationship State — Active",
    "Relationship State — Resolved",
    "Technology State — Active",
    "Technology State — Resolved",
    "Timeline",
    "Character",
    "Relationship",
    "World",
    "Location",
    "Object",
}

STORY_STATE_HEADING = "[StoryState]"
STORY_STATE_STATUS_PATTERN = re.compile(r"^STATE:\s*(ACTIVE|RESOLVED)\s*$", re.IGNORECASE)
STORY_STATE_FIRST_SEEN_PATTERN = re.compile(r"^FIRST_SEEN:\s*Chapter\s+(\d+)\s*$", re.IGNORECASE)
STORY_STATE_RESOLVED_PATTERN = re.compile(r"^RESOLVED_IN:\s*Chapter\s+(\d+)\s*$", re.IGNORECASE)


SCENE_SYSTEM_PROMPT = """You are a Canon Memory Extraction Engine for a long-form novel system.

Your job is to read the provided chapter text and extract story-critical continuity facts,
even when they are implied through descriptive narrative.

Canon memory must preserve narrative lifecycle canon, not just flat facts.
Focus on what changes, escalates, resolves, or sets up future payoff.
It must NOT read like encyclopaedic worldbuilding notes.

Detect and preserve, when strongly supported by the text:

- Persistent world facts
- Character traits, behavioural signals, and injuries
- Relationship changes
- Mission condition changes
- Technology states
- Psychological pressure states
- Narrative tension activations
- Narrative tension resolutions
- Foreshadowing setups
- Foreshadowing payoffs

Prioritize, in roughly this order:

1. Mission risk, survival pressure, or major condition changes
2. Narrative tensions being activated, escalated, or resolved
3. Character psychological pressure, behaviour shifts, or revealing traits
4. Relationship tension, trust shifts, alliance changes, or responsibility shifts
5. Injuries or physical condition changes
6. Technology states affecting survival, access, timing, or the plot
7. Foreshadowing setups or payoffs
8. Persistent world rules or facts affecting later continuity
9. Important tracked objects or location changes

Deprioritize:

- Decorative environment description
- General setting layout
- Atmospheric flavour without continuity consequences
- Generic background lore
- Technical explanation unless it changes risk, capability, or future continuity

If the chapter is calm, store psychological, relational, or mission-progress movement instead.
Prefer fewer high-impact facts over many low-importance facts.
Each fact must be short, concrete, actionable, and non-duplicative.
Prefer narrative-active facts over descriptive facts.
Avoid extracting facts already implied by a stronger fact.
Do not invent or infer beyond what is strongly supported.

Encode narrative states inside the fact text itself while keeping the output format unchanged.
For example:
- Shield deployment mechanism jammed creating mission risk -> [World]
- Dr Manfrid showing hesitation under pressure indicating psychological strain -> [Character]
- Crew tension increases due to mission failure risk -> [Relationship]
- Impending neutron star event creates unresolved mission threat -> [Timeline]

Return output in this exact structure:

Memory suggestions:

1. Canon fact text -> [Category]
2. Canon fact text -> [Category]
3. Canon fact text -> [Category]

Allowed categories:

Character
Timeline
World
Object
Relationship
Injury
Location

Rules:

- Extract only strong continuity-relevant facts.
- Use exactly one allowed category per fact.
- Encode mission states, technology states, psychological pressure, tension activation or resolution, and foreshadowing inside the fact text instead of creating new categories.
- Detect cause-effect changes.
- Detect risk escalation.
- Detect responsibility shifts.
- Detect emotional or behavioural signals.
- Detect mission progress markers.
- Prefer facts that show a change, activation, escalation, resolution, setup, or payoff.
- Use short, concrete statements, not explanations.
- Keep only the minimum set of facts needed to preserve continuity.
- Do not repeat identical or meaningfully redundant facts.
- Do not invent.
- If no strong canon facts exist, return:

Memory suggestions:

None
"""

SCENE_SUMMARY_SYSTEM_PROMPT = """You are a strict isolated Narrative Analysis Engine for a novel-writing project.

You must analyze ONLY the material provided in this single request.
You must ignore chat history and never rely on prior conversation state.
Do not provide general writing advice, critique, or rewriting.
Use the canon memory, previous chapter summaries, and screenplay text only as reference material for comparison.
Do not invent unsupported facts.

Return output in this exact structure and order:

MEMORY SUGGESTIONS

1. Fact text → [Category]
2. Fact text → [Category]

STORY STATE UPDATES

ACTIVATE
1. Description of unresolved narrative pressure
2. Description of unresolved narrative pressure

RESOLVE
1. Short reference label
2. Short reference label

CHAPTER STRUCTURE NOTE

<one of the required chapter-ending decisions>

SCREENPLAY ALIGNMENT NOTE

<alignment note>

RESOLUTION SUGGESTIONS

1. Short reference label
2. Short reference label

Rules for MEMORY SUGGESTIONS:
- Use a numbered list.
- Use short, concrete canon facts from the scene only.
- Only extract facts that matter for long-term continuity.
- Detect new canon facts introduced in the scene.
- Detect when a new unresolved narrative pressure, problem, mystery, danger, mission condition, relationship condition, psychological condition, or technology condition begins, and mark it as an Active state.
- Detect when a previously introduced problem, tension, mystery, danger, mission condition, relationship condition, psychological condition, or technology condition is resolved, and mark it as a Resolved state.
- Detect emotional or psychological turning points.
- Detect technology or mission condition changes.
- Detect setup or payoff of foreshadowing elements.
- Prioritize active mission problems, psychological changes, relationship shifts, injuries, technology state, location changes, foreshadowing, survival rules, and tracked objects.
- Deprioritize decorative description, layout, atmosphere, and generic lore.
- If the chapter is calm, prefer psychological or relational movement.
- Prefer fewer high-impact facts over many weak ones.
- Avoid duplicating facts already implied by stronger facts.
- Do NOT invent facts not present in the scene.
- Every suggestion must use exactly one of these categories: Character, Relationship, World, Location, Object, Timeline, Injury, Mission State — Active, Mission State — Resolved, Psychological State — Active, Psychological State — Resolved, Relationship State — Active, Relationship State — Resolved, Technology State — Active, Technology State — Resolved, Foreshadowing Setup, Foreshadowing Payoff.
- If there are no strong canon facts, return exactly:
MEMORY SUGGESTIONS

None
- STORY STATE UPDATES must still be included even if MEMORY SUGGESTIONS is None.

Rules for STORY STATE UPDATES:
- Track cinematic unresolved pressures that should persist across chapters.
- Detect urgency, countdowns, approaching danger, worsening conditions, hidden problems, emotional strain, rising conflict, strategic uncertainty, environmental threat, time pressure, secrecy, withheld information, and mission instability.
- Track both plot-level threats and internal/relationship pressure when they materially affect future narrative movement.
- Under ACTIVATE, list only unresolved conditions that should become persistent Story States.
- Under RESOLVE, list short labels for persistent Story States clearly stabilised, closed, or resolved by this scene.
- If a new pressure is clearly a continuation of an existing state, use wording that closely matches the existing state instead of inventing a separate one.
- Descriptions must be short, concrete, and cinematic.
- If no items apply in a subsection, write exactly: None

Rules for CHAPTER STRUCTURE NOTE:
- Evaluate chapter break strength based on whether the scene creates a strong location shift, tension pivot, emotional resolution, reveal, reversal, decision point, cliffhanger, or meaningful pause in the narrative lifecycle.
- Return exactly one of these lines:
Strong natural chapter ending point.
Moderate possible chapter ending.
Likely NOT a chapter ending.

Rules for SCREENPLAY ALIGNMENT NOTE:
- If screenplay text is not provided, say exactly: No screenplay source available for comparison.
- If screenplay text is provided and the scene diverges, briefly explain the divergence in motivation or events.
- If screenplay text is provided and aligned, say exactly: Scene remains consistent with screenplay intent.

Rules for RESOLUTION SUGGESTIONS:
- Look at canon memory facts marked ACTIVE.
- If the scene clearly resolves or concludes a situation, suggest it.
- Use a numbered list.
- Use a short reference label, not the full fact text.
- If no resolutions apply, write exactly:
None
"""

CONTINUITY_SYSTEM_PROMPT = """You are a strict continuity editor for a novel project.

Your task is to compare canon memory, previous chapters, and the selected chapter.
Return ONLY factual continuity issues.

Allowed issue types:
- injuries disappearing or changing without explanation
- characters appearing in incorrect locations
- timeline contradictions
- broken world rules
- object continuity errors
- relationship inconsistencies
- knowledge inconsistencies

Rules:
- Do not give writing advice.
- Do not rewrite any text.
- Do not praise the writing.
- Do not mention style, pacing, tone, or quality.
- Do not speculate beyond the provided material.
- If there are no continuity issues, say exactly: "No factual continuity issues found."

Return output in this exact structure:

CONTINUITY REPORT

- Issue 1
- Issue 2
"""

BOOK_INTEGRITY_SYSTEM_PROMPT = """You are a professional novel structural editor.

Analyse the FULL novel draft.

Return a BOOK INTEGRITY REPORT.

You must evaluate:

STRUCTURE

- pacing imbalance
- weak openings
- rushed climaxes
- chapter length inconsistency

CONTINUITY

- injury continuity errors
- location contradictions
- timeline breaks
- knowledge inconsistencies

CHARACTER ARCS

- unresolved motivations
- inconsistent emotional progression
- sudden unexplained behavioural shifts

PLOT THREAD STATUS

- unresolved setups
- dropped conflicts
- premature resolutions

TENSION FLOW

- flat narrative zones
- spikes without build-up
- missing escalation

WORLD CONSISTENCY

- broken rules
- technology drift
- setting contradictions

Do NOT rewrite text.
Do NOT give story ideas.
Do NOT praise writing.

Return format:

BOOK INTEGRITY REPORT

<sections listed above>"""

PROOFREAD_SYSTEM_PROMPT = """You are a professional novel editor and formatter.

Your job is to REWRITE the provided text into a clean, publication-ready novel format.

You must perform THREE steps in order:

1) Rewrite and correct the text
2) Summarise what you changed
3) Provide writing improvement suggestions

All three steps are mandatory.

You MUST output a fully corrected version of the text.

Do NOT provide suggestions first.
Do NOT analyse before rewriting.

-----------------------------------------------------

APPLY THESE RULES:

LANGUAGE:
- Correct all grammar, spelling, and punctuation
- Use STRICT British English spelling (UK English)
- Improve sentence clarity where needed

PARAGRAPH STRUCTURE:
- Break large text blocks into proper paragraphs
- Start a new paragraph when:
  - A new character speaks
  - Focus or action changes
  - A new narrative beat occurs

DIALOGUE:
- Use double quotation marks
- Each new speaker must start a new paragraph
- Dialogue tags remain in the same paragraph

SCENE STRUCTURE:
- Insert scene breaks using:
  ***
- ONLY when there is a clear time jump, location change, or perspective shift
- Do NOT overuse scene breaks

INDENTATION:
- Format as clean novel paragraphs
- First paragraph after a break or scene = no indent

STYLE:
- Preserve original tone and intent
- Do NOT add new story content
- Only rewrite for correctness, clarity, and formatting

-----------------------------------------------------

OUTPUT FORMAT:

CORRECTED TEXT:

<fully rewritten novel-formatted text>

---

CHANGES MADE:

<clear summary of corrections>

---

WRITING IMPROVEMENTS:

- stronger verbs
- adjective upgrades
- improved vocabulary
- phrasing improvements

-----------------------------------------------------

STRICT RULES:

- DO NOT leave the text unchanged
- DO NOT output suggestions before corrected text
- MUST include CHANGES MADE section
- MUST include WRITING IMPROVEMENTS section
- DO NOT treat improvements as optional
- DO NOT skip explanation of what was changed
- DO NOT explain grammar
- DO NOT analyse writing quality
- DO NOT add new story elements
- Provide replacement upgrades where possible (example: "walked quickly" → "hurried")

If no improvements are needed, still output the corrected text.
"""



IDEA_RESURFACE_SYSTEM_PROMPT = """You are a strategic story editor.

Your job is to:
- Read the current chapter
- Read full canon memory
- Read the list of stored ideas

Then:

Select ONLY ideas that strongly fit the story at this point.

STRICT RULES:
- Do NOT force ideas
- Do NOT suggest weak connections
- Do NOT suggest ideas that don't naturally fit
- If confidence is low, return:

No ideas that could work here.

If ideas DO fit:

Return:

IDEA RESURFACING

1. <idea text> → <short reason why it fits now>
2. <idea text> → <short reason why it fits now>

Be concise.
No fluff.
No writing advice.
No rewriting."""


CHAPTER_SUMMARY_SYSTEM_PROMPT = """You are a narrative summarisation engine for a novel project.

You must analyse ONLY the provided chapter text and canon memory.

Return output in this exact structure:

CHAPTER SUMMARY

Key Events:

- ...
- ...

Character Movement:

- ...
- ...

New World Information:

- ...

Tension Movement:

- Rising
  OR
- Stable
  OR
- Climax
  OR
- Resolution

Unresolved Threads Introduced:

- ...

Resolved Threads:

- ...

STORY STATE UPDATES

ACTIVATE
- ...

RESOLVE
- ...

Rules:

- Be factual.
- Do not give writing advice.
- Do not rewrite scenes.
- Do not praise the writing.
- Do not speculate beyond the text.
- In STORY STATE UPDATES, capture unresolved narrative escalation, psychological pressure, relationship strain, mission risk, environmental threat, secrecy, and strategic instability that should persist across chapters.
- Use ACTIVATE for ongoing pressures and RESOLVE for tensions clearly stabilised or concluded.
- If nothing applies under a subsection, write `- None`."""


DRAFT_PASS_SYSTEM_PROMPT_TEMPLATE = """You are a professional novel developmental editor.

You must analyse the provided text.

You must NOT rewrite any part of the story.
You must NOT give grammar corrections.
You must NOT give publishing advice.
You must NOT invent new plot ideas.

You must ONLY evaluate the requested dimension:

{dimension_instructions}

Return output in this format:

<{dimension_name} PASS>

Strengths:
- bullet points

Weaknesses:
- bullet points

Suggested Improvements:
- bullet points
"""


# ============================================================
# Filesystem helpers
# ============================================================


def ensure_project_files() -> None:
    """Create the expected project folders and files if they do not already exist."""
    PROJECT_MEMORY_DIR.mkdir(parents=True, exist_ok=True)
    PROJECT_ANALYSIS_DIR.mkdir(parents=True, exist_ok=True)
    CHAPTERS_DIR.mkdir(parents=True, exist_ok=True)
    MANUSCRIPT_DIR.mkdir(parents=True, exist_ok=True)
    (NOVEL_PROJECT_DIR / "sources").mkdir(parents=True, exist_ok=True)
    CONTINUITY_REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    TIMELINE_LOGS_DIR.mkdir(parents=True, exist_ok=True)
    BOOK_INTEGRITY_REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    REBUILD_LOG_DIR.mkdir(parents=True, exist_ok=True)
    RESEARCH_DIR.mkdir(parents=True, exist_ok=True)
    RESEARCH_INTEGRITY_REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    WORLD_PLAUSIBILITY_REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    PROJECT_BACKUPS_DIR.mkdir(parents=True, exist_ok=True)
    CANON_MEMORY_PATH.touch(exist_ok=True)
    CONTINUITY_INDEX_PATH.touch(exist_ok=True)
    SCENE_SUMMARIES_PATH.touch(exist_ok=True)
    IDEAS_PATH.touch(exist_ok=True)
    WORLD_RULES_PATH.touch(exist_ok=True)



def atomic_write(path: Path, text: str) -> None:
    """Write text to a temporary file and atomically replace the destination."""
    tmp = path.with_suffix(path.suffix + ".tmp")
    tmp.write_text(text, encoding="utf-8")
    tmp.replace(path)


def create_operation_backup(backup_type: str, source_path: Path | None = None, content: str | None = None) -> Path:
    """Create project-level timestamp backup for canon/chapter/draft restore operations."""
    PROJECT_BACKUPS_DIR.mkdir(parents=True, exist_ok=True)
    backup_timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M")
    backup_path = PROJECT_BACKUPS_DIR / f"backup_{backup_type}_{backup_timestamp}.txt"
    if content is not None:
        backup_content = content
    elif source_path is not None and source_path.exists():
        backup_content = source_path.read_text(encoding="utf-8")
    else:
        backup_content = ""
    atomic_write(backup_path, backup_content)
    return backup_path


def create_canon_memory_guard_backup() -> Path:
    """Create a timestamped canon memory backup for guarded writes."""
    CANON_MEMORY_BACKUPS_DIR.mkdir(parents=True, exist_ok=True)
    backup_timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    backup_path = CANON_MEMORY_BACKUPS_DIR / f"canon_backup_{backup_timestamp}.txt"
    backup_content = (
        CANON_MEMORY_PATH.read_text(encoding="utf-8")
        if CANON_MEMORY_PATH.exists()
        else ""
    )
    atomic_write(backup_path, backup_content)
    return backup_path


def guarded_write_canon_memory(chapters: list[dict[str, Any]]) -> None:
    """Safely write canon memory with pre-write backup and post-write validation."""
    backup_path = create_canon_memory_guard_backup()
    create_operation_backup("canon", source_path=CANON_MEMORY_PATH)
    rendered_memory = render_canon_memory(chapters)

    try:
        parse_canon_memory(rendered_memory)
    except Exception as exc:
        raise OSError(f"Canon memory guard rejected invalid render: {exc}") from exc

    atomic_write(CANON_MEMORY_PATH, rendered_memory)

    try:
        written_text = CANON_MEMORY_PATH.read_text(encoding="utf-8")
        parse_canon_memory(written_text)
        refresh_continuity_index(chapters)
    except Exception as exc:
        rollback_text = backup_path.read_text(encoding="utf-8")
        atomic_write(CANON_MEMORY_PATH, rollback_text)
        raise OSError(
            f"Canon memory guard restored from backup after validation failed: {exc}"
        ) from exc



def read_text_file(path: Path) -> str:
    """Read a text file and return cleaned text."""
    text = path.read_text(encoding="utf-8").strip()
    return text if text else "(empty)"



def load_pdf_text(path: Path) -> str:
    """Extract text from a PDF if possible, otherwise return a fallback note."""
    if not path.exists() or not path.is_file():
        return ""

    try:
        from pypdf import PdfReader
    except ImportError:
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            return "(screenplay PDF exists but no PDF reader library is installed)"

    try:
        reader = PdfReader(str(path))
        pages = []
        for page in reader.pages:
            page_text = page.extract_text() or ""
            cleaned_page_text = page_text.strip()
            if cleaned_page_text:
                pages.append(cleaned_page_text)
        return "\n\n".join(pages).strip()
    except Exception as exc:
        return f"(could not read screenplay PDF: {exc})"


def load_memory_block(*, full: bool = False) -> str:
    """Load canonical story memory for the main assistant and continuity checker."""
    canon_text = read_text_file(CANON_MEMORY_PATH)
    if full:
        return canon_text
    if len(canon_text) > MAX_CANON_CHARACTERS:
        return canon_text[-MAX_CANON_CHARACTERS:]
    return canon_text


def load_recap_canon_context() -> str:
    """Load only recap-relevant canon sections with safe scaling."""
    if not CANON_MEMORY_PATH.exists():
        return "(empty)"

    canon_text = clean_terminal_text(CANON_MEMORY_PATH.read_text(encoding="utf-8")).strip()
    if not canon_text:
        return "(empty)"

    chapters = parse_canon_memory(canon_text)
    if not chapters:
        if len(canon_text) <= MAX_CANON_CHARACTERS:
            return canon_text
        chunked = chunk_text_blocks([canon_text], max_chars=MAX_CANON_CHARACTERS)
        return chunked[-1] if chunked else canon_text[-MAX_CANON_CHARACTERS:]

    latest_canon_chapter = max(chapters, key=lambda chapter: int(chapter.get("number", 0)))
    active_states: list[str] = []
    unresolved_threads: list[str] = []
    latest_mission_context: list[str] = []

    for chapter in chapters:
        chapter_number = int(chapter.get("number", 0))
        for story_state in chapter.get("story_states", []):
            if str(story_state.get("state", "ACTIVE")).upper() != "ACTIVE":
                continue
            description = str(story_state.get("description", "")).strip()
            if not description:
                continue
            state_line = f"Chapter {chapter_number}: {description}"
            active_states.append(state_line)
            if should_track_story_state(description, "Story State"):
                unresolved_threads.append(state_line)

        for category, facts in chapter.get("categories", {}).items():
            category_lower = category.lower()
            for fact in facts:
                fact_text = normalize_fact_text(fact)
                if not fact_text:
                    continue
                if "mission" in category_lower or "mission" in fact_text.lower():
                    latest_mission_context.append(f"Chapter {chapter_number}: {fact_text}")

    latest_chapter_block = render_canon_memory([latest_canon_chapter]).strip()
    if not latest_chapter_block:
        latest_chapter_block = "(none)"

    if len(active_states) > 12:
        active_states = active_states[-12:]
    if len(unresolved_threads) > 12:
        unresolved_threads = unresolved_threads[-12:]
    if len(latest_mission_context) > 10:
        latest_mission_context = latest_mission_context[-10:]

    sections = [
        "MOST RECENT CANON CHAPTER BLOCK\n"
        f"{latest_chapter_block}",
        "ACTIVE STORY STATES\n"
        + ("\n".join(f"- {state}" for state in active_states) if active_states else "- None"),
        "UNRESOLVED THREADS\n"
        + ("\n".join(f"- {thread}" for thread in unresolved_threads) if unresolved_threads else "- None"),
        "LATEST MISSION CONTEXT\n"
        + (
            "\n".join(f"- {item}" for item in latest_mission_context)
            if latest_mission_context
            else "- None"
        ),
    ]
    recap_context = "\n\n".join(sections).strip()

    if len(canon_text) <= MAX_CANON_CHARACTERS and len(recap_context) <= MAX_CANON_CHARACTERS:
        return recap_context

    recap_chunks = chunk_text_blocks([recap_context], max_chars=MAX_CANON_CHARACTERS)
    if recap_chunks:
        return recap_chunks[-1]
    return recap_context[-MAX_CANON_CHARACTERS:]


def load_optional_recap_context(path: Path, label: str) -> str:
    """Load optional recap context file if present."""
    if not path.exists():
        return ""
    content = clean_terminal_text(path.read_text(encoding="utf-8")).strip()
    if not content:
        return ""
    return f"{label}:\n{content}"


def load_world_rules_block() -> str:
    """Load structured world rules for continuity checking."""
    if not WORLD_RULES_PATH.exists():
        return "(empty)"
    return read_text_file(WORLD_RULES_PATH)


def load_ideas_block() -> str:
    """Load saved writing ideas for idea suggestion requests."""
    if not IDEAS_PATH.exists():
        return ""
    return IDEAS_PATH.read_text(encoding="utf-8").strip()


def load_previous_scene_summaries_block() -> str:
    """Load previous scene summaries for isolated narrative analysis context."""
    if not SCENE_SUMMARIES_PATH.exists():
        return ""

    divider = "=" * 40
    summaries_text = SCENE_SUMMARIES_PATH.read_text(encoding="utf-8").strip()
    if not summaries_text:
        return ""

    parts = summaries_text.split(divider)
    entries: list[str] = []
    index = 1
    while index + 1 < len(parts):
        header = parts[index].strip()
        body = parts[index + 1].strip()
        if header and body:
            entries.append(f"{divider}\n{header}\n{divider}\n{body}")
        index += 2

    if not entries:
        return summaries_text

    return "\n\n".join(entries[-MAX_SCENE_SUMMARIES:])


def load_screenplay_block() -> str:
    """Load screenplay source text from PDF when available."""
    return load_pdf_text(SCREENPLAY_SOURCE_PATH)



def append_world_rule(category: str, rule_text: str) -> None:
    """Append one structured world rule without overwriting existing rules."""
    PROJECT_MEMORY_DIR.mkdir(parents=True, exist_ok=True)
    cleaned_category = category.strip()
    cleaned_rule_text = rule_text.strip()

    if not cleaned_category or not cleaned_rule_text:
        raise ValueError("Category and rule text are required.")

    entry = f"[{cleaned_category}]\n{cleaned_rule_text}"
    existing_content = (
        WORLD_RULES_PATH.read_text(encoding="utf-8")
        if WORLD_RULES_PATH.exists()
        else ""
    )

    with WORLD_RULES_PATH.open("a", encoding="utf-8") as file:
        if existing_content.strip():
            separator = "\n\n"
            if existing_content.endswith("\n\n"):
                separator = ""
            elif existing_content.endswith("\n"):
                separator = "\n"
            file.write(separator)
        file.write(entry + "\n")



def parse_canon_fact(fact_line: str) -> dict[str, str]:
    """Parse one canon fact line into text and lifecycle state."""
    match = FACT_STATE_PATTERN.fullmatch(fact_line.strip())
    if match is None:
        return {"text": fact_line.strip(), "state": "ACTIVE"}

    fact_text = match.group(1).strip()
    fact_state = (match.group(2) or "ACTIVE").strip().upper()
    return {
        "text": fact_text,
        "state": "RESOLVED" if fact_state == "RESOLVED" else "ACTIVE",
    }



def render_canon_fact(fact: str | dict[str, Any]) -> str:
    """Render one canon fact entry with its lifecycle state marker."""
    if isinstance(fact, dict):
        fact_text = str(fact.get("text", "")).strip()
        fact_state = str(fact.get("state", "ACTIVE")).strip().upper() or "ACTIVE"
    else:
        fact_text = str(fact).strip()
        fact_state = "ACTIVE"

    normalized_state = "RESOLVED" if fact_state == "RESOLVED" else "ACTIVE"
    return f"{fact_text} ({normalized_state})"



def get_fact_text(fact: str | dict[str, Any]) -> str:
    """Return canon fact text without any lifecycle marker."""
    if isinstance(fact, dict):
        return str(fact.get("text", "")).strip()
    return parse_canon_fact(str(fact)).get("text", "").strip()



def parse_canon_memory(content: str) -> list[dict[str, Any]]:
    """Parse canon memory text into ordered chapter/category blocks."""
    chapters: list[dict[str, Any]] = []
    current_chapter: dict[str, Any] | None = None
    current_category: str | None = None
    current_story_state: dict[str, Any] | None = None

    for raw_line in content.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        if stripped == "========================":
            continue

        chapter_match = CHAPTER_HEADER_PATTERN.fullmatch(stripped)
        if chapter_match is not None:
            current_story_state = None
            current_chapter = {
                "number": int(chapter_match.group(1)),
                "categories": OrderedDict(),
                "story_states": [],
            }
            chapters.append(current_chapter)
            current_category = None
            continue

        if current_story_state is not None:
            if not stripped:
                current_story_state = None
                continue

            if stripped == STORY_STATE_HEADING:
                current_chapter["story_states"].append(
                    {
                        "description": "",
                        "state": "ACTIVE",
                        "first_seen": current_chapter["number"],
                        "resolved_in": None,
                    }
                )
                current_story_state = current_chapter["story_states"][-1]
                continue

            state_match = STORY_STATE_STATUS_PATTERN.fullmatch(stripped)
            if state_match is not None:
                current_story_state["state"] = state_match.group(1).upper()
                continue

            first_seen_match = STORY_STATE_FIRST_SEEN_PATTERN.fullmatch(stripped)
            if first_seen_match is not None:
                current_story_state["first_seen"] = int(first_seen_match.group(1))
                continue

            resolved_match = STORY_STATE_RESOLVED_PATTERN.fullmatch(stripped)
            if resolved_match is not None:
                current_story_state["resolved_in"] = int(resolved_match.group(1))
                continue

            description = current_story_state.get("description", "")
            current_story_state["description"] = (
                f"{description} {stripped}".strip() if description else stripped
            )
            continue

        if not stripped:
            continue

        if stripped == STORY_STATE_HEADING and current_chapter is not None:
            current_story_state = {
                "description": "",
                "state": "ACTIVE",
                "first_seen": current_chapter["number"],
                "resolved_in": None,
            }
            current_chapter["story_states"].append(current_story_state)
            current_category = None
            continue

        category_match = CATEGORY_HEADER_PATTERN.fullmatch(stripped)
        if category_match is not None:
            if current_chapter is None:
                continue
            current_category = category_match.group(1).strip()
            current_chapter["categories"].setdefault(current_category, [])
            continue

        if stripped.startswith("-") and current_chapter is not None and current_category is not None:
            fact = stripped[1:].strip()
            if fact:
                current_chapter["categories"][current_category].append(parse_canon_fact(fact))

    return chapters



def render_canon_memory(chapters: list[dict[str, Any]]) -> str:
    """Render ordered chapter/category blocks back to canon memory text."""
    blocks: list[str] = []
    for chapter in chapters:
        lines = ["========================", f"CHAPTER {chapter['number']}"]
        for category, facts in chapter["categories"].items():
            lines.append("")
            lines.append(f"[{category}]")
            for fact in facts:
                lines.append(f"- {render_canon_fact(fact)}")
        for story_state in chapter.get("story_states", []):
            description = str(story_state.get("description", "")).strip()
            if not description:
                continue
            lines.extend(
                [
                    "",
                    STORY_STATE_HEADING,
                    description,
                    f"STATE: {str(story_state.get('state', 'ACTIVE')).upper()}",
                    f"FIRST_SEEN: Chapter {int(story_state.get('first_seen', chapter['number']))}",
                ]
            )
            resolved_in = story_state.get("resolved_in")
            if resolved_in is not None:
                lines.append(f"RESOLVED_IN: Chapter {int(resolved_in)}")
        blocks.append("\n".join(lines).rstrip())
    return "\n\n".join(blocks).rstrip() + ("\n" if blocks else "")


def order_memory_categories(categories: OrderedDict[str, list[dict[str, str]]]) -> OrderedDict[str, list[dict[str, str]]]:
    """Return categories in canonical display order, keeping unknown categories last."""
    ordered_categories: OrderedDict[str, list[dict[str, str]]] = OrderedDict()

    for category in ALLOWED_MEMORY_CATEGORIES:
        if category in categories and categories[category]:
            ordered_categories[category] = categories[category]

    for category, facts in categories.items():
        if category not in ordered_categories and facts:
            ordered_categories[category] = facts

    return ordered_categories



def normalize_fact_text(text: str | dict[str, Any]) -> str:
    """Normalize canon fact text for duplicate and near-duplicate comparison."""
    cleaned_text = get_fact_text(text).lower().strip()
    cleaned_text = re.sub(r"\s+", " ", cleaned_text)
    cleaned_text = cleaned_text.strip("\"'")
    cleaned_text = cleaned_text.replace("--", "")
    cleaned_text = re.sub(r"\s+", " ", cleaned_text)
    cleaned_text = cleaned_text.rstrip(".,;: ")
    return cleaned_text


def facts_are_similar(fact_a: str | dict[str, Any], fact_b: str | dict[str, Any]) -> bool:
    """Return True when two canon facts are duplicates or near-duplicates."""
    normalized_fact_a = normalize_fact_text(fact_a)
    normalized_fact_b = normalize_fact_text(fact_b)

    if normalized_fact_a == normalized_fact_b:
        return True

    words_a = normalized_fact_a.split()
    words_b = normalized_fact_b.split()
    if not words_a or not words_b:
        return False

    common_words = len(set(words_a) & set(words_b))
    overlap_ratio = common_words / max(len(words_a), len(words_b))
    return overlap_ratio >= 0.6


def normalize_story_state_description(description: str) -> str:
    """Normalize a story-state description for matching and deduplication."""
    return normalize_fact_text(description)


def story_states_are_similar(state_a: str | dict[str, Any], state_b: str | dict[str, Any]) -> bool:
    """Return True when two story-state descriptions represent the same persistent pressure."""
    text_a = (
        str(state_a.get("description", "")).strip()
        if isinstance(state_a, dict)
        else str(state_a).strip()
    )
    text_b = (
        str(state_b.get("description", "")).strip()
        if isinstance(state_b, dict)
        else str(state_b).strip()
    )
    return facts_are_similar(text_a, text_b)


def should_track_story_state(fact_text: str, category: str) -> bool:
    """Return True when a canon fact should also become a persistent cinematic Story State."""
    cleaned_fact = normalize_story_state_description(fact_text)
    if not cleaned_fact:
        return False

    if category in {
        "Mission State — Active",
        "Mission State — Resolved",
        "Psychological State — Active",
        "Psychological State — Resolved",
        "Relationship State — Active",
        "Relationship State — Resolved",
        "Technology State — Active",
        "Technology State — Resolved",
        "Injury",
    }:
        return True

    if category not in STORY_STATE_PRIORITY_CATEGORIES:
        return False

    return any(keyword in cleaned_fact for keyword in STORY_STATE_SIGNAL_KEYWORDS)


def iter_story_states(chapters: list[dict[str, Any]]) -> list[tuple[dict[str, Any], dict[str, Any]]]:
    """Return chapter/state pairs for every parsed Story State."""
    pairs: list[tuple[dict[str, Any], dict[str, Any]]] = []
    for chapter in chapters:
        for story_state in chapter.get("story_states", []):
            pairs.append((chapter, story_state))
    return pairs


def add_or_update_story_state(
    chapters: list[dict[str, Any]],
    description: str,
    first_seen: int,
    state: str = "ACTIVE",
    resolved_in: int | None = None,
) -> bool:
    """Create or update one persistent Story State without duplicating continuations."""
    cleaned_description = description.strip()
    normalized_state = "RESOLVED" if state.strip().upper() == "RESOLVED" else "ACTIVE"
    if not cleaned_description:
        return False

    matched_chapter: dict[str, Any] | None = None
    matched_story_state: dict[str, Any] | None = None

    for chapter, story_state in iter_story_states(chapters):
        if not story_states_are_similar(story_state, cleaned_description):
            continue
        matched_chapter = chapter
        matched_story_state = story_state
        break

    if matched_story_state is not None and matched_chapter is not None:
        existing_description = str(matched_story_state.get("description", "")).strip()
        if len(cleaned_description) > len(existing_description):
            matched_story_state["description"] = cleaned_description
        matched_story_state["first_seen"] = min(
            int(matched_story_state.get("first_seen", first_seen)),
            first_seen,
        )
        if normalized_state == "RESOLVED":
            matched_story_state["state"] = "RESOLVED"
            matched_story_state["resolved_in"] = resolved_in or first_seen
        else:
            matched_story_state["state"] = "ACTIVE"
            matched_story_state["resolved_in"] = None
        return True

    target_chapter: dict[str, Any] | None = None
    for chapter in chapters:
        if chapter["number"] == first_seen:
            target_chapter = chapter
            break

    if target_chapter is None:
        target_chapter = {
            "number": first_seen,
            "categories": OrderedDict(),
            "story_states": [],
        }
        chapters.append(target_chapter)
        chapters.sort(key=lambda chapter: chapter["number"])

    target_chapter.setdefault("story_states", []).append(
        {
            "description": cleaned_description,
            "state": normalized_state,
            "first_seen": first_seen,
            "resolved_in": resolved_in if normalized_state == "RESOLVED" else None,
        }
    )
    return True


def sync_story_states_from_suggestions(
    chapters: list[dict[str, Any]],
    chapter_number: int,
    suggestions: list[tuple[int, str, str]],
) -> int:
    """Derive Story States from canon suggestions and update the canon-memory model."""
    updates = 0
    for _, fact, category in suggestions:
        if not should_track_story_state(fact, category):
            continue

        state = "ACTIVE"
        resolved_in: int | None = None
        if category in ACTIVE_TO_RESOLVED_CATEGORY.values():
            state = "RESOLVED"
            resolved_in = chapter_number

        if add_or_update_story_state(
            chapters,
            description=fact,
            first_seen=chapter_number,
            state=state,
            resolved_in=resolved_in,
        ):
            updates += 1

    return updates


def consolidate_story_states(chapters: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Deduplicate Story States across the whole novel while preserving earliest origin."""
    consolidated_chapters = deepcopy(chapters)
    original_states: list[dict[str, Any]] = []

    for chapter in consolidated_chapters:
        original_states.extend(
            [
                {
                    "description": str(state.get("description", "")).strip(),
                    "state": str(state.get("state", "ACTIVE")).upper(),
                    "first_seen": int(state.get("first_seen", chapter["number"])),
                    "resolved_in": state.get("resolved_in"),
                }
                for state in chapter.get("story_states", [])
                if str(state.get("description", "")).strip()
            ]
        )
        chapter["story_states"] = []

    for story_state in sorted(original_states, key=lambda state: state["first_seen"]):
        add_or_update_story_state(
            consolidated_chapters,
            description=story_state["description"],
            first_seen=story_state["first_seen"],
            state=story_state["state"],
            resolved_in=story_state["resolved_in"],
        )

    consolidated_chapters.sort(key=lambda chapter: chapter["number"])
    return consolidated_chapters


def append_to_canon_memory(
    chapter_number: int,
    selected_facts: list[tuple[str, str]],
) -> None:
    """Append selected canon facts into the appropriate chapter block."""
    ensure_project_files()
    cleaned_facts = [
        (fact.strip(), category.strip())
        for fact, category in selected_facts
        if fact.strip() and category.strip()
    ]
    if not cleaned_facts:
        print("No canon facts selected to save.")
        return

    existing_content = CANON_MEMORY_PATH.read_text(encoding="utf-8")
    chapters = parse_canon_memory(existing_content)

    target_chapter: dict[str, Any] | None = None
    for chapter in chapters:
        if chapter["number"] == chapter_number:
            target_chapter = chapter
            break

    if target_chapter is None:
        target_chapter = {
            "number": chapter_number,
            "categories": OrderedDict(),
            "story_states": [],
        }
        chapters.append(target_chapter)

    saved_count = 0
    skipped_duplicates = 0
    for fact, category in cleaned_facts:
        category_facts = target_chapter["categories"].setdefault(category, [])
        if any(facts_are_similar(existing_fact, fact) for existing_fact in category_facts):
            skipped_duplicates += 1
            print(f"Skipped duplicate canon fact: {fact}")
            continue
        category_facts.append({"text": fact, "state": "ACTIVE"})
        saved_count += 1

    if saved_count > 0:
        target_chapter["categories"] = order_memory_categories(target_chapter["categories"])
    story_state_updates = sync_story_states_from_suggestions(chapters, chapter_number, [
        (index, fact, category) for index, (fact, category) in enumerate(cleaned_facts, start=1)
    ])
    if saved_count > 0 or story_state_updates > 0:
        chapters = consolidate_story_states(chapters)
        guarded_write_canon_memory(chapters)

    print(f"Saved {saved_count} canon fact(s).")
    print(f"Skipped {skipped_duplicates} duplicate fact(s).")
    if story_state_updates > 0:
        print(f"Updated {story_state_updates} story state(s).")



def mark_fact_resolved(chapter_number: int, category: str, fact_text: str) -> bool:
    """Mark the first matching canon fact in a chapter/category as resolved."""
    ensure_project_files()
    if not CANON_MEMORY_PATH.exists():
        return False

    cleaned_category = category.strip()
    cleaned_fact_text = fact_text.strip()
    if not cleaned_category or not cleaned_fact_text:
        return False

    chapters = parse_canon_memory(CANON_MEMORY_PATH.read_text(encoding="utf-8"))
    updated = False

    for chapter in chapters:
        if chapter["number"] != chapter_number:
            continue

        category_facts = chapter["categories"].get(cleaned_category, [])
        for existing_fact in category_facts:
            if not facts_are_similar(existing_fact, cleaned_fact_text):
                continue
            if existing_fact.get("state") == "RESOLVED":
                return False
            existing_fact["state"] = "RESOLVED"
            updated = True
            break
        break

    if not updated:
        return False

    guarded_write_canon_memory(chapters)
    return True



def append_scene_summary(chapter_number: int, summary_text: str) -> None:
    """Append full narrative analysis output to the scene summaries log."""
    ensure_project_files()
    cleaned_summary = summary_text.strip()
    if not cleaned_summary:
        return

    divider = "=" * 40
    entry = f"{divider}\nCHAPTER {chapter_number}\n{divider}\n{cleaned_summary}\n"

    with SCENE_SUMMARIES_PATH.open("a", encoding="utf-8") as file:
        if SCENE_SUMMARIES_PATH.stat().st_size > 0:
            file.write("\n")
        file.write(entry)


def remove_scene_summary_block(chapter_number: int) -> bool:
    """Remove one chapter block from the scene summaries log if it exists."""
    if not SCENE_SUMMARIES_PATH.exists():
        return False

    content = SCENE_SUMMARIES_PATH.read_text(encoding="utf-8")
    if not content.strip():
        return False

    divider = "=" * 40
    escaped_divider = re.escape(divider)
    block_pattern = re.compile(
        rf"(?:^|\n){escaped_divider}\nCHAPTER {chapter_number}\n{escaped_divider}\n.*?"
        rf"(?=\n{escaped_divider}\nCHAPTER \d+\n{escaped_divider}\n|\Z)",
        re.DOTALL,
    )
    updated_content, removals = block_pattern.subn("", content, count=1)
    if removals == 0:
        return False

    normalized_content = updated_content.strip()
    atomic_write(
        SCENE_SUMMARIES_PATH,
        (normalized_content + "\n") if normalized_content else "",
    )
    return True



def append_idea(text: str) -> None:
    """Append a timestamped writing idea to the ideas log."""
    PROJECT_MEMORY_DIR.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now()
    with IDEAS_PATH.open("a", encoding="utf-8") as file:
        file.write(
            "\n".join(
                [
                    "========================",
                    timestamp.strftime("%Y-%m-%d %H:%M"),
                    "========================",
                    "Idea:",
                    text.strip(),
                    "",
                    "",
                ]
            )
        )



def build_chapter_memory_block(
    chapter_number: int,
    suggestions: list[tuple[int, str, str]],
) -> dict[str, Any]:
    """Convert extracted suggestions into one ordered canon-memory chapter block."""
    categories: OrderedDict[str, list[dict[str, str]]] = OrderedDict()
    for _, fact, category in suggestions:
        cleaned_fact = fact.strip()
        cleaned_category = category.strip()
        if not cleaned_fact or not cleaned_category:
            continue
        category_facts = categories.setdefault(cleaned_category, [])
        if any(facts_are_similar(existing_fact, cleaned_fact) for existing_fact in category_facts):
            continue
        category_facts.append({"text": cleaned_fact, "state": "ACTIVE"})

    chapter_block = {
        "number": chapter_number,
        "categories": order_memory_categories(categories),
        "story_states": [],
    }
    sync_story_states_from_suggestions([chapter_block], chapter_number, suggestions)
    return chapter_block



def insert_or_replace_chapter_block(
    chapters: list[dict[str, Any]],
    rebuilt_chapter: dict[str, Any],
) -> list[dict[str, Any]]:
    """Replace one chapter block and keep chapter ordering numeric."""
    filtered_chapters = [
        chapter
        for chapter in chapters
        if chapter["number"] != rebuilt_chapter["number"]
    ]
    filtered_chapters.append(rebuilt_chapter)
    filtered_chapters.sort(key=lambda chapter: chapter["number"])
    return filtered_chapters



def write_rebuild_log(
    mode: str,
    lines: list[str],
) -> Path:
    """Write a rebuild log file and return its path."""
    timestamp = datetime.utcnow()
    log_path = REBUILD_LOG_DIR / f"rebuild_{timestamp.strftime('%Y%m%d_%H%M%S')}.txt"
    content = "\n".join(
        [
            f"mode: {mode}",
            *lines,
            f"timestamp: {timestamp.strftime('%Y-%m-%d %H:%M:%S UTC')}",
        ]
    )
    atomic_write(log_path, content + "\n")
    return log_path


def validate_rebuilt_canon_text(text: str) -> bool:
    """Return True if rebuilt canon text is non-empty and contains categories."""
    if not text.strip():
        return False
    parsed = parse_canon_memory(text)
    if not parsed:
        return False
    return any(chapter.get("categories") for chapter in parsed)



def extract_chapter_number(path: Path) -> int | None:
    """Return the chapter number from a chapter filename."""
    match = CHAPTER_FILENAME_PATTERN.fullmatch(path.name)
    if match is None:
        return None
    return int(match.group(1))



def load_sorted_chapter_paths() -> list[Path]:
    """Return project chapter files sorted numerically."""
    chapter_paths = []
    for path in CHAPTERS_DIR.iterdir():
        if not path.is_file():
            continue
        chapter_number = extract_chapter_number(path)
        if chapter_number is None:
            continue
        chapter_paths.append((chapter_number, path))
    chapter_paths.sort(key=lambda item: item[0])
    return [path for _, path in chapter_paths]



def format_chapter_block(chapter_paths: list[Path]) -> str:
    """Read chapter files into one labeled text block."""
    sections = []
    for path in chapter_paths:
        sections.append(f"{path.name}:\n{path.read_text(encoding='utf-8').strip()}")
    return "\n\n".join(sections) if sections else "(none)"



def build_manuscript_text(chapter_paths: list[Path]) -> str:
    """Compile numbered chapter files into one manuscript string."""
    sections: list[str] = []

    for path in chapter_paths:
        chapter_number = extract_chapter_number(path)
        if chapter_number is None:
            continue

        chapter_text = path.read_text(encoding="utf-8").strip()
        sections.append(
            "\n".join(
                [
                    "========================",
                    f"CHAPTER {chapter_number}",
                    "========================",
                    "",
                    chapter_text,
                ]
            ).rstrip()
        )

    return "\n\n".join(sections) + ("\n\n" if sections else "")


def load_all_chapters() -> list[str]:
    """Load all chapter_<number>.txt files in numeric order as labeled text blocks."""
    if not CHAPTERS_DIR.exists() or not CHAPTERS_DIR.is_dir():
        return []

    chapter_blocks: list[str] = []
    for chapter_path in load_sorted_chapter_paths():
        chapter_number = extract_chapter_number(chapter_path)
        if chapter_number is None:
            continue
        try:
            chapter_text = clean_terminal_text(chapter_path.read_text(encoding="utf-8")).strip()
        except OSError as exc:
            print(f"Warning: Could not read {chapter_path.name}: {exc}")
            continue
        if not chapter_text:
            continue
        chapter_blocks.append(f"CHAPTER {chapter_number}\n\n{chapter_text}")

    return chapter_blocks


def build_safe_chunks(text: str, max_chars: int = 12000) -> list[str]:
    """Split text into bounded chunks on paragraph boundaries without cutting words."""
    cleaned_text = clean_terminal_text(text).strip()
    if not cleaned_text:
        return []

    paragraphs = [paragraph.strip() for paragraph in re.split(r"\n\s*\n+", cleaned_text) if paragraph.strip()]
    if not paragraphs:
        return []

    chunks: list[str] = []
    current_paragraphs: list[str] = []
    current_length = 0

    def append_current_chunk() -> None:
        nonlocal current_paragraphs, current_length
        if current_paragraphs:
            chunks.append("\n\n".join(current_paragraphs).strip())
            current_paragraphs = []
            current_length = 0

    def split_large_paragraph(paragraph_text: str) -> list[str]:
        words = paragraph_text.split()
        if not words:
            return []

        pieces: list[str] = []
        current_words: list[str] = []
        current_word_length = 0

        for word in words:
            word_length = len(word)
            separator = 1 if current_words else 0
            if current_words and current_word_length + separator + word_length > max_chars:
                pieces.append(" ".join(current_words))
                current_words = [word]
                current_word_length = word_length
                continue

            current_words.append(word)
            current_word_length += separator + word_length

        if current_words:
            pieces.append(" ".join(current_words))

        return pieces

    for paragraph in paragraphs:
        paragraph_parts = split_large_paragraph(paragraph)
        for paragraph_part in paragraph_parts:
            part_length = len(paragraph_part)
            separator_length = 2 if current_paragraphs else 0

            if current_paragraphs and current_length + separator_length + part_length > max_chars:
                append_current_chunk()

            current_paragraphs.append(paragraph_part)
            current_length += (2 if len(current_paragraphs) > 1 else 0) + part_length

    append_current_chunk()
    return chunks


def split_text_by_word_count(text: str, max_words: int = DEFAULT_ANALYSIS_CHUNK_WORDS) -> list[str]:
    """Split text into chunks of roughly max_words while preserving order."""
    cleaned_text = clean_terminal_text(text).strip()
    if not cleaned_text:
        return []

    words = cleaned_text.split()
    chunks: list[str] = []
    for index in range(0, len(words), max_words):
        chunks.append(" ".join(words[index:index + max_words]))
    return chunks


def print_large_manuscript_warning_if_needed(text: str) -> None:
    """Warn when manuscript enters large-novel scale processing."""
    if len(clean_terminal_text(text).split()) > 90000:
        print("Large manuscript mode active — processing may be slower.")


def write_full_novel_processor_log(
    command_name: str,
    chunk_count: int,
    success: bool,
) -> None:
    """Append full-novel processor execution status to the processor log."""
    timestamp = datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S UTC")
    try:
        FULL_NOVEL_PROCESSOR_LOG_DIR.mkdir(parents=True, exist_ok=True)
        with FULL_NOVEL_PROCESSOR_LOG_PATH.open("a", encoding="utf-8") as handle:
            handle.write(
                f"{timestamp} | command={command_name} | chunks={chunk_count} | "
                f"status={'success' if success else 'failure'}\n"
            )
    except OSError as exc:
        print(f"Warning: Could not write processor log: {exc}")


def process_chunks(
    client: Any,
    system_prompt: str,
    chunks: list[str],
    *,
    temperature: float = CONTINUITY_TEMPERATURE,
) -> list[str]:
    """Process chunk requests with one retry and continue on repeated failures."""
    summaries: list[str] = []
    ordered_chunks = [chunk.strip() for chunk in chunks if chunk.strip()]

    for chunk_index, chunk_text in enumerate(ordered_chunks, start=1):
        print(f"Processing chunk {chunk_index} / {len(ordered_chunks)}")
        response_text = ""
        for attempt in range(2):
            try:
                response_text = request_chat_completion(
                    client=client,
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": chunk_text},
                    ],
                    temperature=temperature,
                ).strip()
                break
            except Exception as exc:  # Keep terminal app stable for the user.
                if attempt == 0:
                    print(f"Warning: Chunk {chunk_index} failed, retrying once: {exc}")
                    continue
                print(f"Warning: Chunk {chunk_index} skipped after retry failure: {exc}")

        if response_text:
            summaries.append(response_text)

    return summaries


def synthesise_chunk_summaries(
    client: Any,
    system_prompt: str,
    summaries: list[str],
    *,
    temperature: float = CONTINUITY_TEMPERATURE,
) -> str:
    """Synthesize chunk summaries into one final global report."""
    cleaned_summaries = [summary.strip() for summary in summaries if summary.strip()]
    if not cleaned_summaries:
        return ""

    if len(cleaned_summaries) == 1:
        return cleaned_summaries[0]

    synthesis_payload = "\n\n".join(
        f"Chunk {index} summary:\n{summary}"
        for index, summary in enumerate(cleaned_summaries, start=1)
    )
    return request_chat_completion(
        client=client,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": synthesis_payload},
        ],
        temperature=temperature,
    )


def run_full_novel_processor(
    client: Any,
    command_name: str,
    chunk_system_prompt: str,
    synthesis_system_prompt: str,
    chunks: list[str],
    *,
    temperature: float = CONTINUITY_TEMPERATURE,
) -> str:
    """Run FULL_NOVEL_PROCESSOR chunk pass + synthesis + logging."""
    summaries = process_chunks(
        client=client,
        system_prompt=chunk_system_prompt,
        chunks=chunks,
        temperature=temperature,
    )

    if not summaries:
        write_full_novel_processor_log(command_name, len(chunks), False)
        return ""

    final_output = synthesise_chunk_summaries(
        client=client,
        system_prompt=synthesis_system_prompt,
        summaries=summaries,
        temperature=temperature,
    )
    success = bool(final_output.strip())
    write_full_novel_processor_log(command_name, len(chunks), success)
    return final_output


def chunk_text_blocks(blocks: list[str], max_chars: int = 12000) -> list[str]:
    """Group text blocks into bounded chunks for safer long-form analysis calls."""
    return build_safe_chunks("\n\n".join(blocks), max_chars=max_chars)


def split_manuscript_into_chunks(full_text: str) -> list[str]:
    """Split manuscript text into safe bounded chunks for full-book analysis."""
    return split_text_by_word_count(full_text, max_words=DEFAULT_ANALYSIS_CHUNK_WORDS)


def run_chunked_analysis(
    client: Any,
    system_prompt: str,
    chunks: list[str],
    *,
    temperature: float = CONTINUITY_TEMPERATURE,
) -> str:
    """Backward-compatible wrapper around FULL_NOVEL_PROCESSOR helpers."""
    summaries = process_chunks(
        client=client,
        system_prompt=system_prompt,
        chunks=chunks,
        temperature=temperature,
    )
    if not summaries:
        return ""

    synthesis_system_prompt = (
        "You are combining multiple analysis reports into one final coherent report.\n\n"
        "Rules:\n"
        "- Remove duplicate issues\n"
        "- Merge similar findings\n"
        "- Preserve factual accuracy\n"
        "- Maintain bullet structure\n"
        "- Do NOT invent new issues\n"
        "- Do NOT give writing advice unless original command allows it."
    )
    return synthesise_chunk_summaries(
        client=client,
        system_prompt=synthesis_system_prompt,
        summaries=summaries,
        temperature=temperature,
    )



def warn_for_missing_chapter_files(chapter_paths: list[Path]) -> None:
    """Print warnings for gaps in numbered chapter files without stopping execution."""
    chapter_numbers = [
        extract_chapter_number(path)
        for path in chapter_paths
    ]
    ordered_numbers = sorted(number for number in chapter_numbers if number is not None)
    if not ordered_numbers:
        return

    expected = set(range(ordered_numbers[0], ordered_numbers[-1] + 1))
    missing_numbers = sorted(expected - set(ordered_numbers))
    for missing_number in missing_numbers:
        print(f"Warning: Missing chapter_{missing_number}.txt")


def prompt_for_destructive_confirmation() -> bool:
    """Require an explicit yes confirmation before destructive operations."""
    print("This will overwrite current project state. Continue? (yes/no)")
    try:
        response = input("> ").strip()
    except EOFError:
        print()
        return False
    return response == "yes"


def create_canon_memory_backup() -> None:
    """Create a timestamped canon memory backup before rebuild operations."""
    backup_path = create_operation_backup("canon", source_path=CANON_MEMORY_PATH)
    print("Canon memory backup created.")
    print(f"Backup path: {backup_path}")


def build_continuity_index(chapters: list[dict[str, Any]]) -> str:
    """Build a lightweight continuity index from canon memory chapters."""
    lines: list[str] = ["CONTINUITY INDEX", ""]
    category_groups = (
        ("key injuries", {"Injury"}),
        ("mission milestones", {"Timeline", "Mission State — Active", "Mission State — Resolved"}),
        ("location changes", {"Location"}),
        ("major character events", {"Character", "Relationship", "Psychological State — Active", "Psychological State — Resolved", "Relationship State — Active", "Relationship State — Resolved"}),
    )

    for heading, categories in category_groups:
        lines.append(f"{heading}:")
        added = 0
        for chapter in sorted(chapters, key=lambda item: item["number"]):
            chapter_number = chapter["number"]
            for category_name in categories:
                for fact in chapter["categories"].get(category_name, []):
                    text = get_fact_text(fact)
                    if not text:
                        continue
                    lines.append(f"- Ch {chapter_number}: {text}")
                    added += 1
                    if added >= 40:
                        break
                if added >= 40:
                    break
            if added >= 40:
                break
        if added == 0:
            lines.append("- none recorded")
        lines.append("")

    return "\n".join(lines).rstrip() + "\n"


def refresh_continuity_index(chapters: list[dict[str, Any]]) -> None:
    """Persist continuity index alongside canon memory updates."""
    atomic_write(CONTINUITY_INDEX_PATH, build_continuity_index(chapters))



# ============================================================
# Input helpers
# ============================================================


def clean_terminal_text(text: str) -> str:
    """Remove terminal control artifacts from pasted multi-line text."""
    cleaned_text = ANSI_ESCAPE_PATTERN.sub("", text)
    cleaned_text = BRACKETED_PASTE_PATTERN.sub("", cleaned_text)
    cleaned_text = DISALLOWED_CONTROL_CHAR_PATTERN.sub("", cleaned_text)
    return cleaned_text.strip()


def collect_multiline_input(end_marker: str = "END") -> str:
    """Collect multi-line input until the user types the end marker."""
    lines: list[str] = []
    while True:
        try:
            line = input()
        except EOFError:
            print()
            break

        if line.strip() == end_marker:
            break
        lines.append(line)

    return clean_terminal_text("\n".join(lines))



def prompt_for_chapter_number(prompt_text: str = "Chapter number?") -> int | None:
    """Ask the user for a chapter number."""
    print(prompt_text)
    try:
        raw_value = input("> ").strip()
    except EOFError:
        print()
        return None

    if not raw_value:
        print("No chapter number entered.")
        return None

    if not raw_value.isdigit():
        print("Chapter number must be a positive integer.")
        return None

    chapter_number = int(raw_value)
    if chapter_number <= 0:
        print("Chapter number must be a positive integer.")
        return None

    return chapter_number



def prompt_for_confirmation(prompt_text: str) -> bool:
    """Ask the user a yes/no question and return True for yes."""
    print(prompt_text)
    try:
        response = input("> ").strip().lower()
    except EOFError:
        print()
        return False

    return response == "y"


def prompt_for_research_choice(
    prompt_text: str,
    options: dict[str, str],
) -> str | None:
    """Prompt for a numbered research option and return the selected key."""
    print(prompt_text)
    for key, label in options.items():
        print(f"{key} = {label}")
    try:
        choice = input("> ").strip()
    except EOFError:
        print()
        return None

    if choice not in options:
        print("Invalid selection.")
        return None

    return choice



def parse_memory_suggestions(result: str) -> list[tuple[int, str, str]]:
    """Parse numbered memory suggestions from the scene extractor output."""
    suggestions: list[tuple[int, str, str]] = []
    unparsed_lines: list[str] = []
    seen = set()
    flexible_pattern = re.compile(
        r"^\s*(?:[-*•]\s*)?(?:(\d+)\s*[.)-]?\s*)?(.+?)\s*(?:→|->|=>)\s*\[\s*([^\]]+?)\s*\]\s*$"
    )

    for line in result.splitlines():
        stripped = line.strip()
        if not stripped or stripped.lower().startswith("memory suggestions"):
            continue
        match = flexible_pattern.match(stripped)
        if not match:
            unparsed_lines.append(stripped)
            continue
        number = int(match.group(1)) if match.group(1) else len(suggestions) + 1
        fact = re.sub(r"\s+", " ", match.group(2)).strip(" -\t")
        category = re.sub(r"\s+", " ", match.group(3)).strip()
        if category not in ALLOWED_MEMORY_CATEGORIES or not fact:
            unparsed_lines.append(stripped)
            continue
        dedupe_key = (fact.lower(), category)
        if dedupe_key in seen:
            continue
        seen.add(dedupe_key)
        suggestions.append((number, fact, category))

    if unparsed_lines:
        try:
            timestamp = datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S UTC")
            with UNPARSED_MEMORY_SUGGESTIONS_LOG_PATH.open("a", encoding="utf-8") as handle:
                handle.write(f"{timestamp}\n")
                for line in unparsed_lines:
                    handle.write(f"- {line}\n")
                handle.write("\n")
        except OSError:
            pass
    return suggestions



def parse_resolution_suggestions(result_text: str) -> list[tuple[int, str]]:
    """Parse numbered resolution suggestions from the scene summary output."""
    match = re.search(
        r"RESOLUTION SUGGESTIONS\s*(.*?)(?:\n[A-Z][A-Z ]+\n|\Z)",
        result_text,
        re.DOTALL,
    )
    if match is None:
        return []

    section_text = match.group(1).strip()
    if not section_text or section_text == "None":
        return []

    suggestions: list[tuple[int, str]] = []
    for line in section_text.splitlines():
        numbered_match = re.match(r"^\s*(\d+)\.\s*(.+?)\s*$", line)
        if numbered_match is None:
            continue
        suggestions.append((int(numbered_match.group(1)), numbered_match.group(2).strip()))
    return suggestions


def parse_story_state_updates(result_text: str) -> tuple[list[str], list[str]]:
    """Parse STORY STATE UPDATES sections into activation and resolution labels."""
    match = re.search(
        r"STORY STATE UPDATES\s*(.*?)(?:\n[A-Z][A-Z ]+\n|\Z)",
        result_text,
        re.DOTALL,
    )
    if match is None:
        return [], []

    section_text = match.group(1).strip()
    activate_match = re.search(
        r"ACTIVATE\s*(.*?)(?:\nRESOLVE\s|\Z)",
        section_text,
        re.DOTALL,
    )
    resolve_match = re.search(r"RESOLVE\s*(.*)\Z", section_text, re.DOTALL)

    def parse_items(raw_text: str) -> list[str]:
        cleaned = raw_text.strip()
        if not cleaned or cleaned.lower() == "none":
            return []

        items: list[str] = []
        for line in cleaned.splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            item_match = re.match(r"^(?:\d+\.|-)\s*(.+?)\s*$", stripped)
            if item_match is None:
                continue
            item_text = item_match.group(1).strip()
            if item_text.lower() == "none":
                return []
            items.append(item_text)
        return items

    activated = parse_items(activate_match.group(1) if activate_match else "")
    resolved = parse_items(resolve_match.group(1) if resolve_match else "")
    return activated, resolved



def prompt_for_selection(max_number: int, prompt_text: str = "\nSelect numbers to save:") -> list[int] | None:
    """Ask the user which numbered suggestions should be applied."""
    print(prompt_text)
    try:
        raw_selection = input("> ").strip()
    except EOFError:
        print()
        return None

    if not raw_selection:
        print("No selections entered. Nothing saved.")
        return []

    tokens = raw_selection.replace(",", " ").split()
    selections: list[int] = []
    seen: set[int] = set()

    for token in tokens:
        if not token.isdigit():
            print(f"Invalid selection: {token}")
            return None
        number = int(token)
        if number < 1 or number > max_number:
            print(f"Selection out of range: {number}")
            return None
        if number not in seen:
            selections.append(number)
            seen.add(number)

    return selections


def resolution_label_matches_fact(label: str, fact: str) -> bool:
    """Return True when a short resolution label clearly points to an ACTIVE fact."""
    normalized_label = normalize_fact_text(label)
    normalized_fact = normalize_fact_text(fact)
    if not normalized_label or not normalized_fact:
        return False

    if normalized_label in normalized_fact or normalized_fact in normalized_label:
        return True

    label_words = normalized_label.split()
    if label_words and all(word in normalized_fact.split() for word in label_words):
        return True

    return facts_are_similar(label, fact)



def apply_resolutions(selected_labels: list[str], chapter_number: int | None = None) -> int:
    """Move matching ACTIVE canon facts and Story States into their resolved lifecycle state."""
    ensure_project_files()
    if not selected_labels or not CANON_MEMORY_PATH.exists():
        return 0

    canon_content = CANON_MEMORY_PATH.read_text(encoding="utf-8")
    chapters = parse_canon_memory(canon_content)
    resolved_count = 0

    for label in selected_labels:
        normalized_label = normalize_fact_text(label)
        if not normalized_label:
            continue

        matched_entry: tuple[dict[str, Any], str, str] | None = None
        story_state_resolved = False

        for chapter in chapters:
            for category in ACTIVE_TO_RESOLVED_CATEGORY:
                facts = chapter["categories"].get(category, [])
                for fact in facts:
                    if resolution_label_matches_fact(label, fact):
                        matched_entry = (chapter, category, fact)
                        break
                if matched_entry is not None:
                    break
            if matched_entry is not None:
                break

        if matched_entry is None:
            for _, story_state in iter_story_states(chapters):
                if (
                    str(story_state.get("state", "ACTIVE")).upper() == "ACTIVE"
                    and resolution_label_matches_fact(label, str(story_state.get("description", "")))
                ):
                    story_state["state"] = "RESOLVED"
                    story_state["resolved_in"] = chapter_number
                    resolved_count += 1
                    story_state_resolved = True
                    break
            if matched_entry is None:
                if story_state_resolved:
                    continue
                continue

        chapter, active_category, fact = matched_entry
        resolved_category = ACTIVE_TO_RESOLVED_CATEGORY[active_category]
        active_facts = chapter["categories"].get(active_category, [])
        if fact not in active_facts:
            continue

        active_facts.remove(fact)
        resolved_facts = chapter["categories"].setdefault(resolved_category, [])
        if fact not in resolved_facts:
            resolved_facts.append(fact)
        if not active_facts:
            chapter["categories"].pop(active_category, None)

        chapter["categories"] = order_memory_categories(chapter["categories"])
        resolved_count += 1

    if resolved_count > 0:
        chapters = consolidate_story_states(chapters)
        guarded_write_canon_memory(chapters)

    return resolved_count


def apply_story_state_updates(
    chapter_number: int,
    activated_states: list[str],
    resolved_labels: list[str],
) -> tuple[int, int]:
    """Persist Story State activations and resolutions from an analysis response."""
    ensure_project_files()
    chapters = parse_canon_memory(CANON_MEMORY_PATH.read_text(encoding="utf-8"))
    activations_applied = 0

    for description in activated_states:
        if add_or_update_story_state(
            chapters,
            description=description,
            first_seen=chapter_number,
            state="ACTIVE",
        ):
            activations_applied += 1

    resolutions_applied = 0
    for label in resolved_labels:
        for _, story_state in iter_story_states(chapters):
            if (
                str(story_state.get("state", "ACTIVE")).upper() == "ACTIVE"
                and resolution_label_matches_fact(label, str(story_state.get("description", "")))
            ):
                story_state["state"] = "RESOLVED"
                story_state["resolved_in"] = chapter_number
                resolutions_applied += 1
                break

    if activations_applied > 0 or resolutions_applied > 0:
        chapters = consolidate_story_states(chapters)
        guarded_write_canon_memory(chapters)

    return activations_applied, resolutions_applied


# ============================================================
# Global loading spinner
# ============================================================

_loading_lock = threading.Lock()
_loading_thread: threading.Thread | None = None
_loading_active = False
_loading_message = "Processing..."
_loading_depth = 0


def _spinner_loop() -> None:
    """Render a one-line spinner while loading is active."""
    frame_index = 0
    while True:
        with _loading_lock:
            if not _loading_active:
                break
            message = _loading_message
        frame = SPINNER_FRAMES[frame_index % len(SPINNER_FRAMES)]
        sys.stdout.write(f"\r{message} {frame}")
        sys.stdout.flush()
        frame_index += 1
        time.sleep(SPINNER_INTERVAL_SECONDS)


def start_loading(message: str = "Processing...") -> None:
    """Start the global loading spinner (supports nested calls)."""
    global _loading_active, _loading_thread, _loading_message, _loading_depth
    with _loading_lock:
        _loading_depth += 1
        _loading_message = message
        if _loading_active:
            return
        _loading_active = True
        _loading_thread = threading.Thread(target=_spinner_loop, daemon=True)
        _loading_thread.start()


def stop_loading() -> None:
    """Stop the global loading spinner safely and clear its line."""
    global _loading_active, _loading_thread, _loading_depth
    spinner_thread: threading.Thread | None = None
    with _loading_lock:
        if _loading_depth > 0:
            _loading_depth -= 1
        if _loading_depth > 0:
            return
        _loading_active = False
        spinner_thread = _loading_thread
        _loading_thread = None

    if spinner_thread is not None:
        spinner_thread.join(timeout=1.0)

    clear_width = shutil.get_terminal_size(fallback=(80, 20)).columns
    sys.stdout.write("\r" + (" " * max(0, clear_width - 1)) + "\r")
    sys.stdout.write("\n")
    sys.stdout.flush()


def run_with_loading(message: str, function: Callable[[], Any]) -> Any:
    """Run any callable with spinner lifecycle handled automatically."""
    start_loading(message)
    try:
        return function()
    finally:
        stop_loading()


# ============================================================
# OpenAI helpers
# ============================================================


def create_client() -> Any:
    """Create the OpenAI client using environment configuration."""
    try:
        from openai import OpenAI
    except ImportError as exc:
        raise RuntimeError(
            "The OpenAI package is not installed. Run: pip install openai"
        ) from exc

    return OpenAI()



def request_chat_completion(
    client: Any,
    messages: list[dict[str, str]],
    temperature: float,
    loading_message: str = "Processing...",
) -> str:
    """Send a chat request and return plain text output."""
    response = run_with_loading(
        loading_message,
        lambda: client.responses.create(
            model=MODEL_NAME,
            input=messages,
            temperature=temperature,
        ),
    )
    return response.output_text.strip()


# ============================================================
# Prompt builders
# ============================================================


def build_main_messages(
    memory_block: str,
    conversation_history: list[dict[str, str]],
    user_message: str,
) -> list[dict[str, str]]:
    """Build the message list for normal assistant chat."""
    messages: list[dict[str, str]] = [
        {
            "role": "system",
            "content": MAIN_SYSTEM_PROMPT,
        },
        {
            "role": "system",
            "content": f"Persistent novel memory:\n\n{memory_block}",
        },
    ]
    messages.extend(conversation_history)
    messages.append({"role": "user", "content": user_message})
    return messages



def build_scene_messages(
    scene_text: str,
) -> list[dict[str, str]]:
    """Build the isolated message list for scene extraction."""
    return [
        {
            "role": "system",
            "content": SCENE_SYSTEM_PROMPT,
        },
        {
            "role": "user",
            "content": scene_text,
        },
    ]


def build_scene_summary_messages(
    scene_text: str,
    canon_memory_block: str,
    previous_summaries_block: str,
    screenplay_block: str,
) -> list[dict[str, str]]:
    """Build the isolated message list for narrative scene analysis."""
    screenplay_text = screenplay_block.strip() if screenplay_block.strip() else "(none)"
    previous_summaries_text = (
        previous_summaries_block.strip() if previous_summaries_block.strip() else "(none)"
    )

    return [
        {
            "role": "system",
            "content": SCENE_SUMMARY_SYSTEM_PROMPT,
        },
        {
            "role": "user",
            "content": (
                f"Scene text:\n\n{scene_text}\n\n"
                f"Canon memory:\n\n{canon_memory_block}\n\n"
                f"Previous chapter summaries:\n\n{previous_summaries_text}\n\n"
                f"Screenplay text:\n\n{screenplay_text}"
            ),
        },
    ]



def build_continuity_messages(
    memory_block: str,
    world_rules_block: str,
    previous_chapter_block: str,
    latest_chapter_block: str,
    continuity_index_block: str,
    selected_chapter_name: str,
    selected_chapter_text: str,
) -> list[dict[str, str]]:
    """Build the isolated message list for continuity checking."""
    return [
        {
            "role": "system",
            "content": CONTINUITY_SYSTEM_PROMPT,
        },
        {
            "role": "user",
            "content": (
                f"Canon memory:\n\n{memory_block}\n\n"
                f"World rules:\n\n{world_rules_block}\n\n"
                f"Continuity index:\n\n{continuity_index_block}\n\n"
                f"Previous chapter:\n\n{previous_chapter_block}\n\n"
                f"Latest chapter:\n\n{latest_chapter_block}\n\n"
                f"Selected chapter ({selected_chapter_name}):\n\n{selected_chapter_text}"
            ),
        },
    ]


def build_book_integrity_messages(
    full_novel_block: str,
    canon_memory_block: str,
    world_rules_block: str,
) -> list[dict[str, str]]:
    """Build the isolated message list for full-book integrity analysis."""
    return [
        {
            "role": "system",
            "content": BOOK_INTEGRITY_SYSTEM_PROMPT,
        },
        {
            "role": "user",
            "content": (
                f"Full novel draft:\n\n{full_novel_block}\n\n"
                f"Canon memory:\n\n{canon_memory_block}\n\n"
                f"World rules:\n\n{world_rules_block}"
            ),
        },
    ]


def build_proofread_messages(text_to_proofread: str) -> list[dict[str, str]]:
    """Build the isolated message list for proofreading."""
    return [
        {
            "role": "system",
            "content": PROOFREAD_SYSTEM_PROMPT,
        },
        {
            "role": "user",
            "content": text_to_proofread,
        },
    ]


def copy_to_clipboard(text: str) -> bool:
    """Copy text to clipboard via wl-copy and keep terminal flow stable on errors."""
    try:
        subprocess.run(
            ["wl-copy"],
            input=text.encode(),
            check=True,
        )
    except Exception as exc:  # Keep terminal app stable for the user.
        print("Clipboard copy failed:", exc)
        return False
    return True


def extract_corrected_text(response: str) -> str:
    """Extract corrected text block from proofreading response."""
    if "CORRECTED TEXT:" in response:
        return response.split("CORRECTED TEXT:")[1].split("---")[0].strip()
    return response.strip()


def extract_proofread_section(response: str, header: str) -> str:
    """Extract one labelled proofread section from the model response."""
    escaped_header = re.escape(header)
    pattern = rf"{escaped_header}:\s*(.*?)(?:\n\s*---\s*\n|$)"
    match = re.search(pattern, response, flags=re.DOTALL)
    if not match:
        return ""
    return match.group(1).strip()




def build_idea_resurface_messages(
    chapter_text: str,
    canon_memory_block: str,
    ideas_block: str,
) -> list[dict[str, str]]:
    """Build the message list for resurfacing highly relevant saved ideas."""
    return [
        {
            "role": "system",
            "content": IDEA_RESURFACE_SYSTEM_PROMPT,
        },
        {
            "role": "user",
            "content": (
                f"Current chapter text:\n\n{chapter_text}\n\n"
                f"Canon memory:\n\n{canon_memory_block}\n\n"
                f"Stored ideas:\n\n{ideas_block}"
            ),
        },
    ]


def build_chapter_summary_messages(
    chapter_text: str,
    canon_memory_block: str,
) -> list[dict[str, str]]:
    """Build the message list for chapter-level narrative summarisation."""
    return [
        {
            "role": "system",
            "content": CHAPTER_SUMMARY_SYSTEM_PROMPT,
        },
        {
            "role": "user",
            "content": (
                f"Chapter text:\n\n{chapter_text}\n\n"
                f"Canon memory:\n\n{canon_memory_block}"
            ),
        },
    ]


def build_draft_pass_messages(
    text_to_analyse: str,
    dimension_name: str,
    dimension_instructions: str,
) -> list[dict[str, str]]:
    """Build the isolated message list for draft analysis passes."""
    system_prompt = DRAFT_PASS_SYSTEM_PROMPT_TEMPLATE.format(
        dimension_name=dimension_name,
        dimension_instructions=dimension_instructions,
    )
    return [
        {
            "role": "system",
            "content": system_prompt,
        },
        {
            "role": "user",
            "content": text_to_analyse,
        },
    ]


def build_research_messages(
    question: str,
    depth_choice: str,
    style_choice: str,
) -> list[dict[str, str]]:
    """Build a dedicated isolated scientific research prompt."""
    depth_label = RESEARCH_DEPTH_OPTIONS[depth_choice]
    style_label = RESEARCH_STYLE_OPTIONS[style_choice]

    depth_instructions = {
        "1": (
            "Depth target: Surface realism. Focus on accurate high-confidence fundamentals,"
            " plain-language scientific framing, and practical real-world constraints."
        ),
        "2": (
            "Depth target: Hard sci-fi realism. Include quantitative ranges where possible,"
            " engineering constraints, operational failure modes, and physics-based feasibility."
        ),
        "3": (
            "Depth target: Ultra deep technical realism. Provide highly technical treatment,"
            " equations or scaling-law style reasoning when relevant, edge-case constraints,"
            " and explicit uncertainty bounds."
        ),
    }[depth_choice]

    style_instructions = {
        "1": (
            "Output style: Scientific report. Use precise technical prose with clear headings,"
            " evidence-grounded claims, and concise analytical structure."
        ),
        "2": (
            "Output style: Teaching explanation. Use pedagogical, step-by-step explanation while"
            " remaining scientifically rigorous and fact-only."
        ),
        "3": (
            "Output style: Practical notes. Use concise bullet-point guidance focused on implementation,"
            " constraints, and decision-relevant facts."
        ),
    }[style_choice]

    system_prompt = (
        "You are a neutral scientific research consultant.\n"
        "Provide only real-world scientific facts and clearly marked uncertainty.\n"
        "Absolutely do not provide story suggestions, plot ideas, character advice, writing guidance,\n"
        "creativity prompts, or any connection to a novel, screenplay, chapters, canon memory,\n"
        "or narrative analysis.\n"
        "Do not include speculative fantasy. If mentioning frontier hypotheses, label them clearly as\n"
        "theoretical and not experimentally confirmed.\n"
        "Always cover, where relevant: scale, timeframes, physical limits, environmental effects,\n"
        "survivability, constraints, and engineering implications.\n"
        f"{depth_instructions}\n"
        f"{style_instructions}\n"
        f"Selected realism depth: {depth_label}.\n"
        f"Selected output style: {style_label}."
    )

    return [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": question},
    ]


def build_research_scene_messages(
    scene_chunk: str,
    canon_memory_block: str,
    world_rules_block: str,
    prior_findings: str = "",
) -> list[dict[str, str]]:
    """Build isolated hard-sci-fi realism analysis messages for a scene chunk."""
    prior_findings_block = prior_findings if prior_findings.strip() else "(none)"
    return [
        {"role": "system", "content": RESEARCH_SCENE_SYSTEM_PROMPT},
        {
            "role": "user",
            "content": (
                "Analyse this scene chunk for scientific realism only.\n"
                "Do not rewrite, do not suggest plot/dialogue/prose, and do not provide narrative critique.\n"
                "Use canon memory and world rules only as factual context constraints.\n\n"
                "Output format (strict):\n"
                "SCIENCE REALISM REPORT\n\n"
                "Relevant Real Science\n"
                "- bullet points\n\n"
                "Realistic Constraints\n"
                "- bullet points\n\n"
                "Already Realistic Elements\n"
                "- bullet points\n\n"
                "Unrealistic or Risky Elements\n"
                "- bullet points\n\n"
                "Scientific Detail Opportunities\n"
                "- micro realism ideas\n"
                "- instrumentation behaviour\n"
                "- environmental reactions\n\n"
                f"Prior findings from earlier chunks:\n{prior_findings_block}\n\n"
                f"Canon memory:\n{canon_memory_block}\n\n"
                f"World rules:\n{world_rules_block}\n\n"
                f"Scene chunk:\n{scene_chunk}"
            ),
        },
    ]


def build_research_apply_messages(
    research_notes: str,
    scene_text: str,
) -> list[dict[str, str]]:
    """Build isolated realism-audit messages from saved research notes and scene text."""
    system_prompt = (
        "You are a scientific realism auditor.\n"
        "Compare only the provided Research notes and Scene text.\n"
        "Do not use any external context, canon memory, screenplay source, or creative interpretation.\n"
        "Do not rewrite scenes.\n"
        "Do not suggest plot changes.\n"
        "Do not connect facts to story ideas.\n"
        "Do not give writing advice.\n"
        "Return ONLY a REALISM REPORT.\n\n"
        "Required sections:\n"
        "PHYSICS CONFLICTS\n"
        "ENVIRONMENTAL CONFLICTS\n"
        "SURVIVABILITY ISSUES\n"
        "SCALE / TIMEFRAME ERRORS\n"
        "ENGINEERING IMPOSSIBILITIES\n"
        "ENERGY OR FORCE INACCURACIES\n\n"
        "Rules:\n"
        "- Be factual\n"
        "- Be concise\n"
        "- Do NOT suggest solutions\n"
        "- Do NOT rewrite\n"
        "- Do NOT be creative\n"
        "- If no issues, return exactly: No realism conflicts detected."
    )

    return [
        {"role": "system", "content": system_prompt},
        {
            "role": "user",
            "content": f"Research notes:\n{research_notes}\n\nScene text:\n{scene_text}",
        },
    ]


def build_research_integrity_messages(research_corpus: str) -> list[dict[str, str]]:
    """Build an isolated prompt for auditing scientific consistency across saved research topics."""
    system_prompt = (
        "You are a scientific realism auditor.\n\n"
        "Your task:\n"
        "- Analyse all provided research notes.\n"
        "- Detect contradictions in:\n"
        "  - physics\n"
        "  - astronomy\n"
        "  - engineering feasibility\n"
        "  - biology / survivability\n"
        "  - environmental realism\n"
        "  - time scales\n"
        "  - energy requirements\n"
        "  - technological assumptions\n"
        "  - cause and effect realism\n\n"
        "Rules:\n"
        "- Do NOT provide writing advice.\n"
        "- Do NOT suggest plot ideas.\n"
        "- Do NOT connect analysis to any story.\n"
        "- Do NOT speculate unless clearly labelling \"theoretical physics\".\n"
        "- Focus only on real scientific plausibility.\n\n"
        "Output must follow exactly one of these formats:\n\n"
        "RESEARCH INTEGRITY REPORT\n\n"
        "Scientific contradictions:\n\n"
        "- item\n"
        "- item\n\n"
        "Scale or timeframe risks:\n\n"
        "- item\n"
        "- item\n\n"
        "Engineering feasibility risks:\n\n"
        "- item\n"
        "- item\n\n"
        "Unclear assumptions:\n\n"
        "- item\n"
        "- item\n\n"
        "OR, if no issues exist:\n\n"
        "RESEARCH INTEGRITY REPORT\n\n"
        "No scientific contradictions detected."
    )

    return [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": f"Research notes corpus:\n{research_corpus}"},
    ]


# ============================================================
# Extraction helpers
# ============================================================


def extract_memory_suggestions_for_text(
    client: Any,
    scene_text: str,
) -> list[tuple[int, str, str]]:
    """Run the shared isolated extractor and return parsed suggestions."""
    result = request_chat_completion(
        client=client,
        messages=build_scene_messages(scene_text),
        temperature=SCENE_TEMPERATURE,
    )
    return parse_memory_suggestions(result)


def extract_memory_suggestions_for_large_text(
    client: Any,
    scene_text: str,
) -> list[tuple[int, str, str]]:
    """Extract canon suggestions from very large chapter text using safe chunking."""
    chunks = build_safe_chunks(scene_text, max_chars=12000)
    if not chunks:
        return []

    if len(chunks) == 1:
        return extract_memory_suggestions_for_text(client=client, scene_text=chunks[0])

    chunk_reports = process_chunks(
        client=client,
        system_prompt=SCENE_SYSTEM_PROMPT,
        chunks=chunks,
        temperature=SCENE_TEMPERATURE,
    )

    merged: list[tuple[int, str, str]] = []
    for report in chunk_reports:
        parsed = parse_memory_suggestions(report)
        for _, fact_text, category in parsed:
            if any(
                existing_category == category and facts_are_similar(existing_fact, fact_text)
                for _, existing_fact, existing_category in merged
            ):
                continue
            merged.append((len(merged) + 1, fact_text, category))

    return merged


def generate_scene_summary_for_chapter(
    client: Any,
    chapter_number: int,
) -> str:
    """Run the scene summary prompt for one chapter file."""
    chapter_path = CHAPTERS_DIR / f"chapter_{chapter_number}.txt"
    if not chapter_path.exists() or not chapter_path.is_file():
        raise FileNotFoundError(f"Chapter file not found: {chapter_path}")

    scene_text = clean_terminal_text(chapter_path.read_text(encoding="utf-8"))
    if not scene_text:
        raise ValueError(f"Chapter {chapter_number} is empty.")

    return request_chat_completion(
        client=client,
        messages=build_scene_summary_messages(
            scene_text=scene_text,
            canon_memory_block=load_memory_block(),
            previous_summaries_block=load_previous_scene_summaries_block(),
            screenplay_block=load_screenplay_block(),
        ),
        temperature=SCENE_TEMPERATURE,
    )



# ============================================================
# Command handlers
# ============================================================


def handle_scene_summary(client: Any) -> None:
    """Analyse one pasted scene in a fully isolated request."""
    ensure_project_files()
    chapter_number = prompt_for_chapter_number()
    if chapter_number is None:
        return

    print("Paste scene. Type END on new line when finished.")
    scene_text = collect_multiline_input(end_marker="END")
    scene_text = clean_terminal_text(scene_text)

    if not scene_text:
        print("No scene entered.")
        return

    canon_memory_block = load_memory_block()
    previous_summaries_block = load_previous_scene_summaries_block()
    screenplay_block = load_screenplay_block()

    try:
        result = request_chat_completion(
            client=client,
            messages=build_scene_summary_messages(
                scene_text=scene_text,
                canon_memory_block=canon_memory_block,
                previous_summaries_block=previous_summaries_block,
                screenplay_block=screenplay_block,
            ),
            temperature=SCENE_TEMPERATURE,
        )
    except Exception as exc:  # Keep terminal app stable for the user.
        print(f"Scene summary failed: {exc}")
        return

    print()
    print(result)

    activated_states, resolved_story_states = parse_story_state_updates(result)
    try:
        activated_count, resolved_story_count = apply_story_state_updates(
            chapter_number=chapter_number,
            activated_states=activated_states,
            resolved_labels=resolved_story_states,
        )
    except OSError as exc:
        print(f"Could not update story states: {exc}")
        return

    if activated_count > 0 or resolved_story_count > 0:
        print(
            f"\nStory States updated: {activated_count} activated, "
            f"{resolved_story_count} resolved."
        )

    suggestions = parse_memory_suggestions(result)
    if not suggestions:
        print("\nNo structured canon suggestions found. Nothing saved.")
    else:
        selection_numbers = prompt_for_selection(
            max(number for number, _, _ in suggestions)
        )
        if selection_numbers is None:
            print("Nothing saved.")
        else:
            selected_lookup = {number for number in selection_numbers}
            selected_facts = [
                (fact, category)
                for number, fact, category in suggestions
                if number in selected_lookup
            ]

            if selected_facts:
                try:
                    append_to_canon_memory(chapter_number, selected_facts)
                except OSError as exc:
                    print(f"Could not save canon memory: {exc}")
                    return
            else:
                print("Nothing saved.")

    resolution_suggestions = parse_resolution_suggestions(result)
    if resolution_suggestions:
        print("\nSelect resolutions to apply:")
        for number, label_text in resolution_suggestions:
            print(f"{number}. {label_text}")

        resolution_numbers = prompt_for_selection(
            max(number for number, _ in resolution_suggestions),
            prompt_text="",
        )
        if resolution_numbers is None:
            print("Resolved 0 canon facts.")
        else:
            selected_resolution_numbers = {number for number in resolution_numbers}
            selected_labels = [
                label_text
                for number, label_text in resolution_suggestions
                if number in selected_resolution_numbers
            ]
            try:
                resolved_count = apply_resolutions(selected_labels, chapter_number=chapter_number)
            except OSError as exc:
                print(f"Could not update canon memory resolutions: {exc}")
                return
            print(f"Resolved {resolved_count} canon facts.")

    try:
        append_scene_summary(chapter_number, result)
    except OSError as exc:
        print(f"Narrative analysis generated, but could not save log: {exc}")
        return

    print(f"Narrative analysis log saved to {SCENE_SUMMARIES_PATH}.")



def handle_continuity_check(client: Any) -> None:
    """Run a continuity report in a fully isolated request."""
    ensure_project_files()
    print("Enter chapter filename:")
    try:
        selected_name = input("> ").strip()
    except EOFError:
        print()
        return

    if not selected_name:
        print("No chapter filename entered.")
        return

    if not CHAPTERS_DIR.exists():
        print(f"Missing chapters directory: {CHAPTERS_DIR}")
        return

    chapter_paths = load_sorted_chapter_paths()
    selected_path = CHAPTERS_DIR / selected_name

    if selected_path not in chapter_paths:
        print("Chapter not found.")
        return

    selected_index = chapter_paths.index(selected_path)
    previous_path = chapter_paths[selected_index - 1] if selected_index > 0 else None
    latest_path = chapter_paths[-1] if chapter_paths else None
    print("Using previous chapter + latest chapter + continuity index for context.")

    memory_block = load_memory_block(full=True)
    world_rules_block = load_world_rules_block()
    previous_chapter_block = format_chapter_block([previous_path]) if previous_path else "(none)"
    latest_chapter_block = format_chapter_block([latest_path]) if latest_path else "(none)"
    continuity_index_block = (
        read_text_file(CONTINUITY_INDEX_PATH)
        if CONTINUITY_INDEX_PATH.exists()
        else "(none)"
    )
    selected_chapter_text = selected_path.read_text(encoding="utf-8").strip()
    messages = build_continuity_messages(
        memory_block=memory_block,
        world_rules_block=world_rules_block,
        previous_chapter_block=previous_chapter_block,
        latest_chapter_block=latest_chapter_block,
        continuity_index_block=continuity_index_block,
        selected_chapter_name=selected_path.name,
        selected_chapter_text=selected_chapter_text,
    )

    try:
        report = request_chat_completion(
            client=client,
            messages=messages,
            temperature=CONTINUITY_TEMPERATURE,
        )
    except Exception as exc:  # Keep terminal app stable for the user.
        print(f"Continuity check failed: {exc}")
        return

    report_path = CONTINUITY_REPORTS_DIR / f"{selected_path.stem}_report.txt"
    try:
        atomic_write(report_path, report + "\n")
    except OSError as exc:
        print(f"Continuity report could not be saved: {exc}")
        return

    print()
    print(report)
    print()
    print(f"Continuity report saved to {report_path}.")



def handle_book_integrity(client: Any) -> None:
    """Run a full-novel structural and continuity integrity audit."""
    ensure_project_files()

    if not CHAPTERS_DIR.exists():
        print(f"Missing chapters directory: {CHAPTERS_DIR}")
        return

    chapter_paths = load_sorted_chapter_paths()
    if not chapter_paths:
        print("No chapter files found in ~/writing/novel_project/chapters/")
        return

    chapter_blocks = load_all_chapters()
    manuscript_text = "\n\n".join(chapter_blocks).strip()
    print_large_manuscript_warning_if_needed(manuscript_text)
    chunked_blocks = split_manuscript_into_chunks(manuscript_text)
    if not chunked_blocks:
        print("No chapter text found.")
        return

    canon_memory_block = load_memory_block(full=True)
    world_rules_block = (
        read_text_file(WORLD_RULES_PATH)
        if WORLD_RULES_PATH.exists()
        else "(not provided)"
    )

    try:
        analysis_chunks = [
            (
                f"Full novel draft:\n\n{chunk_text}\n\n"
                f"Canon memory:\n\n{canon_memory_block}\n\n"
                f"World rules:\n\n{world_rules_block}"
            )
            for chunk_text in chunked_blocks
        ]
        report = run_full_novel_processor(
            client=client,
            command_name="/book-integrity",
            chunk_system_prompt=BOOK_INTEGRITY_SYSTEM_PROMPT,
            synthesis_system_prompt=(
                "You are combining multiple analysis reports into one final coherent report.\n\n"
                "Rules:\n"
                "- Remove duplicate issues\n"
                "- Merge similar findings\n"
                "- Preserve factual accuracy\n"
                "- Maintain bullet structure\n"
                "- Do NOT invent new issues\n"
                "- Do NOT give writing advice unless original command allows it."
            ),
            chunks=analysis_chunks,
            temperature=CONTINUITY_TEMPERATURE,
        )
    except Exception as exc:  # Keep terminal app stable for the user.
        print(f"Book integrity analysis failed: {exc}")
        return

    timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M")
    report_path = BOOK_INTEGRITY_REPORTS_DIR / f"book_integrity_{timestamp}.txt"
    try:
        atomic_write(report_path, report + "\n")
    except OSError as exc:
        print(f"Book integrity report could not be saved: {exc}")
        return

    print()
    print(report)
    print()
    print(f"Book integrity report saved to {report_path}.")


def handle_world_consistency(client: Any) -> None:
    """Run a chunk-safe full-novel sci-fi world consistency audit."""
    ensure_project_files()

    if not PROJECT_MEMORY_DIR.exists() or not PROJECT_MEMORY_DIR.is_dir():
        print(f"Missing memory directory: {PROJECT_MEMORY_DIR}")
        return

    if not CANON_MEMORY_PATH.exists() or not CANON_MEMORY_PATH.is_file():
        print(f"Missing canon memory file: {CANON_MEMORY_PATH}")
        return

    if not CHAPTERS_DIR.exists() or not CHAPTERS_DIR.is_dir():
        print(f"Missing chapters directory: {CHAPTERS_DIR}")
        return

    try:
        canon_memory = clean_terminal_text(CANON_MEMORY_PATH.read_text(encoding="utf-8")).strip()
    except OSError as exc:
        print(f"Could not read canon memory: {exc}")
        return

    if not canon_memory:
        print(f"Canon memory is empty: {CANON_MEMORY_PATH}")
        return

    world_rules = "(none)"
    if WORLD_RULES_PATH.exists() and WORLD_RULES_PATH.is_file():
        try:
            loaded_world_rules = clean_terminal_text(
                WORLD_RULES_PATH.read_text(encoding="utf-8")
            ).strip()
        except OSError as exc:
            print(f"Could not read world rules: {exc}")
            return
        if loaded_world_rules:
            world_rules = loaded_world_rules

    chapter_paths = load_sorted_chapter_paths()
    if not chapter_paths:
        print("No chapter files found in ~/writing/novel_project/chapters/")
        return

    chapter_blocks = load_all_chapters()

    if not chapter_blocks:
        print("No chapter text found.")
        return

    full_manuscript_text = "\n\n".join(chapter_blocks)
    print_large_manuscript_warning_if_needed(full_manuscript_text)
    chunk_blocks = [
        (
            f"Canon memory:\n\n{canon_memory}\n\n"
            f"World rules:\n\n{world_rules}\n\n"
            f"Novel chunk:\n\n{chunk_text}"
        )
        for chunk_text in split_manuscript_into_chunks(full_manuscript_text)
    ]
    if not chunk_blocks:
        print("No chapter text found.")
        return

    try:
        final_report = run_full_novel_processor(
            client=client,
            command_name="/world-consistency",
            chunk_system_prompt=WORLD_CONSISTENCY_CHUNK_SYSTEM_PROMPT,
            synthesis_system_prompt=WORLD_CONSISTENCY_SYNTHESIS_SYSTEM_PROMPT,
            chunks=chunk_blocks,
            temperature=CONTINUITY_TEMPERATURE,
        ).strip()
    except Exception as exc:  # Keep terminal app stable for the user.
        print(f"World consistency audit failed: {exc}")
        return

    if not final_report:
        print("World consistency audit failed: no chunk summaries produced.")
        return

    print()
    print(final_report)


def handle_character_consistency(client: Any) -> None:
    """Run a chunk-safe full-novel character behaviour consistency audit."""
    ensure_project_files()

    if not PROJECT_MEMORY_DIR.exists() or not PROJECT_MEMORY_DIR.is_dir():
        print(f"Missing memory directory: {PROJECT_MEMORY_DIR}")
        return

    if not CANON_MEMORY_PATH.exists() or not CANON_MEMORY_PATH.is_file():
        print(f"Missing canon memory file: {CANON_MEMORY_PATH}")
        return

    if not CHAPTERS_DIR.exists() or not CHAPTERS_DIR.is_dir():
        print(f"Missing chapters directory: {CHAPTERS_DIR}")
        return

    try:
        canon_memory = clean_terminal_text(CANON_MEMORY_PATH.read_text(encoding="utf-8")).strip()
    except OSError as exc:
        print(f"Could not read canon memory: {exc}")
        return

    if not canon_memory:
        print(f"Canon memory is empty: {CANON_MEMORY_PATH}")
        return

    chapter_paths = load_sorted_chapter_paths()
    if not chapter_paths:
        print("No chapter files found in ~/writing/novel_project/chapters/")
        return

    chapter_blocks = load_all_chapters()

    if not chapter_blocks:
        print("No chapter text found.")
        return

    full_manuscript_text = "\n\n".join(chapter_blocks)
    print_large_manuscript_warning_if_needed(full_manuscript_text)
    chunk_blocks = [
        (
            f"Canon memory:\n\n{canon_memory}\n\n"
            f"Novel chunk:\n\n{chunk_text}"
        )
        for chunk_text in split_manuscript_into_chunks(full_manuscript_text)
    ]
    if not chunk_blocks:
        print("No chapter text found.")
        return

    try:
        final_report = run_full_novel_processor(
            client=client,
            command_name="/character-consistency",
            chunk_system_prompt=CHARACTER_CONSISTENCY_CHUNK_SYSTEM_PROMPT,
            synthesis_system_prompt=CHARACTER_CONSISTENCY_SYNTHESIS_SYSTEM_PROMPT,
            chunks=chunk_blocks,
            temperature=CONTINUITY_TEMPERATURE,
        ).strip()
    except Exception as exc:  # Keep terminal app stable for the user.
        print(f"Character consistency audit failed: {exc}")
        return

    if not final_report:
        print("Character consistency audit failed: no chunk summaries produced.")
        return

    print()
    print(final_report)


def handle_rebuild_memory(client: Any, command_text: str = "") -> None:
    """Rebuild canon memory for the whole novel or one chapter."""
    ensure_project_files()
    normalized_command = command_text.strip().lower()
    forced_full = normalized_command.startswith("/rebuild-memory full")
    forced_single = normalized_command.startswith("/rebuild-memory single")

    selection = ""
    if forced_full:
        selection = "1"
    elif forced_single:
        selection = "2"
    else:
        print("Rebuild options:")
        print()
        print("1) Rebuild entire novel")
        print("2) Rebuild single chapter")

        try:
            selection = input("> ").strip()
        except EOFError:
            print()
            return

    if selection == "1":
        if not prompt_for_destructive_confirmation():
            print("Rebuild aborted.")
            return
        try:
            pre_rebuild_backup_path = create_operation_backup("canon", source_path=CANON_MEMORY_PATH)
            print("Canon memory backup created.")
            print(f"Backup path: {pre_rebuild_backup_path}")
        except OSError as exc:
            print(f"Could not create canon memory backup: {exc}")
            return

        if not CHAPTERS_DIR.exists():
            print(f"Missing chapters directory: {CHAPTERS_DIR}")
            return

        chapter_paths = load_sorted_chapter_paths()
        if not chapter_paths:
            print("No chapter files found.")
            return

        rebuilt_chapters: list[dict[str, Any]] = []
        total_facts = 0

        for chapter_path in chapter_paths:
            chapter_number = extract_chapter_number(chapter_path)
            if chapter_number is None:
                continue

            print(f"Processing chapter {chapter_number}...")
            try:
                suggestions = extract_memory_suggestions_for_large_text(
                    client=client,
                    scene_text=chapter_path.read_text(encoding="utf-8").strip(),
                )
            except Exception as exc:  # Keep terminal app stable for the user.
                print(f"Rebuild failed while processing chapter {chapter_number}: {exc}")
                return

            rebuilt_chapters.append(build_chapter_memory_block(chapter_number, suggestions))
            total_facts += len(suggestions)
            print(f"Extracted {len(suggestions)} fact(s) from chapter {chapter_number}.")

        rebuilt_chapters.sort(key=lambda chapter: chapter["number"])
        rebuilt_chapters = consolidate_story_states(rebuilt_chapters)

        try:
            rendered_memory = render_canon_memory(rebuilt_chapters)
            temp_rebuild_path = CANON_MEMORY_PATH.with_suffix(".txt.rebuild_tmp")
            atomic_write(temp_rebuild_path, rendered_memory)
            rebuilt_text = temp_rebuild_path.read_text(encoding="utf-8")
            if not validate_rebuilt_canon_text(rebuilt_text):
                raise OSError("Rebuild validation failed: empty or missing categories.")
            atomic_write(CANON_MEMORY_PATH, rebuilt_text)
            refresh_continuity_index(rebuilt_chapters)
            temp_rebuild_path.unlink(missing_ok=True)
            log_path = write_rebuild_log(
                mode="FULL",
                lines=[
                    "chapters processed: "
                    + ", ".join(str(chapter["number"]) for chapter in rebuilt_chapters),
                    f"total facts extracted: {total_facts}",
                ],
            )
        except OSError as exc:
            try:
                rollback_text = pre_rebuild_backup_path.read_text(encoding="utf-8")
                atomic_write(CANON_MEMORY_PATH, rollback_text)
                print("Rebuild failed; canon memory restored from backup.")
            except OSError as restore_exc:
                print(f"Rebuild failed and restore failed: {restore_exc}")
                return
            print(f"Rebuild completed, but could not save files: {exc}")
            return

        print()
        print(
            f"Full rebuild complete. Wrote {len(rebuilt_chapters)} chapter block(s) "
            f"and {total_facts} fact(s) to {CANON_MEMORY_PATH}."
        )
        print(f"Rebuild log saved to {log_path}.")
        return

    if selection == "2":
        try:
            create_canon_memory_backup()
        except OSError as exc:
            print(f"Could not create canon memory backup: {exc}")
            return

        chapter_number = prompt_for_chapter_number()
        if chapter_number is None:
            return

        chapter_path = CHAPTERS_DIR / f"chapter_{chapter_number}.txt"
        if not chapter_path.exists() or not chapter_path.is_file():
            print(f"Chapter file not found: {chapter_path}")
            return

        print(f"Processing chapter {chapter_number}...")
        try:
            suggestions = extract_memory_suggestions_for_text(
                client=client,
                scene_text=chapter_path.read_text(encoding="utf-8").strip(),
            )
        except Exception as exc:  # Keep terminal app stable for the user.
            print(f"Single-chapter rebuild failed: {exc}")
            return

        rebuilt_chapter = build_chapter_memory_block(chapter_number, suggestions)
        existing_content = (
            CANON_MEMORY_PATH.read_text(encoding="utf-8")
            if CANON_MEMORY_PATH.exists()
            else ""
        )
        chapters = parse_canon_memory(existing_content)
        chapters = insert_or_replace_chapter_block(chapters, rebuilt_chapter)
        chapters = consolidate_story_states(chapters)

        try:
            guarded_write_canon_memory(chapters)
            log_path = write_rebuild_log(
                mode="SINGLE",
                lines=[
                    f"chapter number: {chapter_number}",
                    f"facts extracted: {len(suggestions)}",
                ],
            )
        except OSError as exc:
            print(f"Single-chapter rebuild completed, but could not save files: {exc}")
            return

        print()
        print(
            f"Single-chapter rebuild complete for chapter {chapter_number}. "
            f"Saved {len(suggestions)} fact(s) to {CANON_MEMORY_PATH}."
        )
        print(f"Rebuild log saved to {log_path}.")
        return

    print("Invalid selection. Enter 1 or 2.")


def handle_rebuild_summaries(client: Any) -> None:
    """Rebuild stored scene summaries from the current chapter files."""
    ensure_project_files()
    print("Rebuild summary options:")
    print()
    print("1) Full novel")
    print("2) Single chapter")

    try:
        selection = input("> ").strip()
    except EOFError:
        print()
        return

    if selection == "1":
        if not CHAPTERS_DIR.exists():
            print(f"Missing chapters directory: {CHAPTERS_DIR}")
            return

        chapter_paths = load_sorted_chapter_paths()
        if not chapter_paths:
            print("No chapter files found.")
            return

        try:
            atomic_write(SCENE_SUMMARIES_PATH, "")
        except OSError as exc:
            print(f"Could not clear scene summaries: {exc}")
            return

        for chapter_path in chapter_paths:
            chapter_number = extract_chapter_number(chapter_path)
            if chapter_number is None:
                continue

            print(f"Processing chapter {chapter_number}...")
            try:
                summary_text = generate_scene_summary_for_chapter(client, chapter_number)
                append_scene_summary(chapter_number, summary_text)
            except FileNotFoundError as exc:
                print(exc)
                continue
            except ValueError as exc:
                print(f"Skipped chapter {chapter_number}: {exc}")
                continue
            except Exception as exc:  # Keep terminal app stable for the user.
                print(f"Rebuild summaries failed while processing chapter {chapter_number}: {exc}")
                return

        print(f"Scene summaries rebuilt at {SCENE_SUMMARIES_PATH}.")
        return

    if selection == "2":
        chapter_number = prompt_for_chapter_number()
        if chapter_number is None:
            return

        chapter_path = CHAPTERS_DIR / f"chapter_{chapter_number}.txt"
        if not chapter_path.exists() or not chapter_path.is_file():
            print(f"Chapter file not found: {chapter_path}")
            return

        try:
            removed_existing = remove_scene_summary_block(chapter_number)
            print(f"Processing chapter {chapter_number}...")
            summary_text = generate_scene_summary_for_chapter(client, chapter_number)
            append_scene_summary(chapter_number, summary_text)
        except ValueError as exc:
            print(f"Could not rebuild chapter {chapter_number}: {exc}")
            return
        except OSError as exc:
            print(f"Could not update scene summaries: {exc}")
            return
        except Exception as exc:  # Keep terminal app stable for the user.
            print(f"Single-chapter summary rebuild failed: {exc}")
            return

        if removed_existing:
            print(f"Rebuilt scene summary for chapter {chapter_number}.")
        else:
            print(
                f"No existing scene summary block found for chapter {chapter_number}. "
                "Added a new one."
            )
        print(f"Scene summaries updated at {SCENE_SUMMARIES_PATH}.")
        return

    print("Invalid selection. Enter 1 or 2.")


def handle_proofread(client: Any, command_text: str = "") -> None:
    """Rewrite and format pasted text, then optionally copy clean output to clipboard."""
    print("Paste text to proofread. Type END on a new line when finished.")
    text_to_proofread = collect_multiline_input(end_marker="END")
    text_to_proofread = clean_terminal_text(text_to_proofread)

    if not text_to_proofread:
        print("No text provided.")
        return

    try:
        result = request_chat_completion(
            client=client,
            messages=build_proofread_messages(text_to_proofread),
            temperature=PROOFREAD_TEMPERATURE,
        )
    except Exception:  # Keep terminal app stable for the user.
        print("Proofread failed.")
        return

    corrected_text = extract_corrected_text(result)
    changes_made = extract_proofread_section(result, "CHANGES MADE")
    writing_improvements = extract_proofread_section(result, "WRITING IMPROVEMENTS")
    skip_clipboard_copy = " nocopy" in f" {command_text.strip().lower()} "

    print()
    print(corrected_text)

    if skip_clipboard_copy:
        print()
        print("📋 Clipboard copy skipped (/proofread nocopy).")
    else:
        copy_to_clipboard(corrected_text)
        print()
        print("📋 Clean text copied to clipboard. Paste into your editor.")

    print()
    print("CHANGES MADE:")
    print(changes_made or "No major correction categories were identified.")

    print()
    print("WRITING IMPROVEMENTS:")
    print(
        writing_improvements
        or "- No additional upgrades were found beyond the corrected rewrite."
    )




def handle_idea_resurface(client: Any) -> None:
    """Resurface only highly relevant saved ideas for the selected chapter."""
    ensure_project_files()
    chapter_number = prompt_for_chapter_number(prompt_text="Enter chapter number:")
    if chapter_number is None:
        return

    chapter_path = CHAPTERS_DIR / f"chapter_{chapter_number}.txt"
    if not chapter_path.exists() or not chapter_path.is_file():
        print(f"Chapter file not found: {chapter_path}")
        return

    ideas_block = load_ideas_block()
    if not ideas_block:
        print("No ideas available.")
        return

    chapter_text = chapter_path.read_text(encoding="utf-8").strip()
    canon_memory_block = load_memory_block()
    messages = build_idea_resurface_messages(
        chapter_text=chapter_text,
        canon_memory_block=canon_memory_block,
        ideas_block=ideas_block,
    )

    try:
        result = request_chat_completion(
            client=client,
            messages=messages,
            temperature=IDEA_RESURFACE_TEMPERATURE,
        )
    except Exception as exc:  # Keep terminal app stable for the user.
        print(f"Idea resurfacing failed: {exc}")
        return

    print()
    print(result)



def handle_chapter_summary(client: Any) -> None:
    """Summarise one full chapter using canon memory for narrative context."""
    ensure_project_files()
    chapter_number = prompt_for_chapter_number(prompt_text="Enter chapter number:")
    if chapter_number is None:
        return

    chapter_path = CHAPTERS_DIR / f"chapter_{chapter_number}.txt"
    if not chapter_path.exists() or not chapter_path.is_file():
        print(f"Chapter file not found: {chapter_path}")
        return

    chapter_text = clean_terminal_text(chapter_path.read_text(encoding="utf-8"))
    canon_memory_block = read_text_file(CANON_MEMORY_PATH)

    try:
        result = request_chat_completion(
            client=client,
            messages=build_chapter_summary_messages(
                chapter_text=chapter_text,
                canon_memory_block=canon_memory_block,
            ),
            temperature=SCENE_TEMPERATURE,
        )
    except Exception as exc:  # Keep terminal app stable for the user.
        print(f"Chapter summary failed: {exc}")
        return

    print()
    print(result)

    activated_states, resolved_story_states = parse_story_state_updates(result)
    try:
        activated_count, resolved_story_count = apply_story_state_updates(
            chapter_number=chapter_number,
            activated_states=activated_states,
            resolved_labels=resolved_story_states,
        )
    except OSError as exc:
        print(f"Could not update story states: {exc}")
        return

    if activated_count > 0 or resolved_story_count > 0:
        print(
            f"\nStory States updated: {activated_count} activated, "
            f"{resolved_story_count} resolved."
        )


def handle_recap(client: Any) -> None:
    """Generate a present-moment narrative recap from the latest chapter context."""
    ensure_project_files()

    chapter_paths = load_sorted_chapter_paths()
    if not chapter_paths:
        print("No chapters found.")
        return

    latest_chapter_path = chapter_paths[-1]
    chapter_number = extract_chapter_number(latest_chapter_path)
    latest_chapter_text = clean_terminal_text(latest_chapter_path.read_text(encoding="utf-8")).strip()
    if not latest_chapter_text:
        latest_chapter_text = "(empty)"

    canon_context = load_recap_canon_context()
    optional_context_blocks: list[str] = []
    story_state_context = load_optional_recap_context(STORY_STATE_PATH, "Optional story_state.txt")
    if story_state_context:
        optional_context_blocks.append(story_state_context)
    timeline_context = load_optional_recap_context(
        TIMELINE_THREADS_PATH,
        "Optional timeline_threads.txt",
    )
    if timeline_context:
        optional_context_blocks.append(timeline_context)
    optional_context = "\n\n".join(optional_context_blocks) if optional_context_blocks else "(none)"

    recap_messages = [
        {"role": "system", "content": RECAP_SYSTEM_PROMPT},
        {
            "role": "user",
            "content": (
                f"Latest chapter number: {chapter_number if chapter_number is not None else '(unknown)'}\n\n"
                f"Latest chapter text:\n\n{latest_chapter_text}\n\n"
                f"Scaled canon context:\n\n{canon_context}\n\n"
                f"Optional context:\n\n{optional_context}"
            ),
        },
    ]

    try:
        recap = request_chat_completion(
            client=client,
            messages=recap_messages,
            temperature=RECAP_TEMPERATURE,
        )
    except Exception as exc:  # Keep terminal app stable for the user.
        print(f"Recap failed: {exc}")
        return

    print()
    print(recap)


def handle_draft_pass(client: Any, command_text: str = "") -> None:
    """Run a structured editorial analysis pass without modifying project files."""
    normalized_command = command_text.strip().lower()
    flag_to_selection = {
        "--structure": "1",
        "--tension": "2",
        "--character": "3",
        "--clarity": "4",
    }
    preselected_mode = next(
        (selection for flag, selection in flag_to_selection.items() if flag in normalized_command),
        "",
    )

    if not preselected_mode:
        print("Select analysis mode:")
        print()
        print("1 — Structure")
        print("2 — Tension")
        print("3 — Character")
        print("4 — Clarity")

        try:
            mode_selection = input("> ").strip()
        except EOFError:
            print()
            return
    else:
        mode_selection = preselected_mode

    dimension_map = {
        "1": (
            "STRUCTURE",
            "STRUCTURE:\n"
            "- pacing\n"
            "- scene purpose\n"
            "- narrative momentum\n"
            "- chapter openings\n"
            "- chapter endings\n"
            "- exposition balance",
        ),
        "2": (
            "TENSION",
            "TENSION:\n"
            "- stakes clarity\n"
            "- urgency\n"
            "- mystery drive\n"
            "- threat escalation\n"
            "- emotional pressure",
        ),
        "3": (
            "CHARACTER",
            "CHARACTER:\n"
            "- motivation clarity\n"
            "- emotional realism\n"
            "- dialogue authenticity\n"
            "- relationship dynamics\n"
            "- behavioural consistency",
        ),
        "4": (
            "CLARITY",
            "CLARITY:\n"
            "- readability\n"
            "- sentence density\n"
            "- description overload\n"
            "- technical confusion\n"
            "- flow interruptions",
        ),
    }

    selected_dimension = dimension_map.get(mode_selection)
    if selected_dimension is None:
        print("Invalid selection. Enter 1, 2, 3, or 4.")
        return

    dimension_name, dimension_instructions = selected_dimension

    print()
    print("Run on:")
    print()
    print("1 — Single chapter")
    print("2 — Full novel")

    try:
        scope_selection = input("> ").strip()
    except EOFError:
        print()
        return

    text_to_analyse = ""
    chapter_chunks: list[str] = []
    if scope_selection == "1":
        print("Enter chapter number:")
        try:
            chapter_value = input("> ").strip()
        except EOFError:
            print()
            return

        if not chapter_value or not chapter_value.isdigit() or int(chapter_value) <= 0:
            print("Chapter number must be a positive integer.")
            return

        chapter_number = int(chapter_value)
        chapter_path = CHAPTERS_DIR / f"chapter_{chapter_number}.txt"
        if not chapter_path.exists() or not chapter_path.is_file():
            print(f"Chapter file not found: {chapter_path}")
            return

        text_to_analyse = clean_terminal_text(chapter_path.read_text(encoding="utf-8"))
    elif scope_selection == "2":
        if not CHAPTERS_DIR.exists() or not CHAPTERS_DIR.is_dir():
            print(f"Missing chapters directory: {CHAPTERS_DIR}")
            return

        chapter_paths = load_sorted_chapter_paths()
        if not chapter_paths:
            print("No chapter files found.")
            return

        chapter_blocks = load_all_chapters()
        full_manuscript_text = "\n\n".join(chapter_blocks)
        print_large_manuscript_warning_if_needed(full_manuscript_text)
        chapter_chunks = split_manuscript_into_chunks(full_manuscript_text)
        if not chapter_chunks:
            print("No chapter text found.")
            return
        text_to_analyse = "\n\n".join(chapter_chunks).strip()
    else:
        print("Invalid selection. Enter 1 or 2.")
        return

    if not text_to_analyse:
        print("No text found to analyse.")
        return

    try:
        if scope_selection == "2":
            draft_system_prompt = DRAFT_PASS_SYSTEM_PROMPT_TEMPLATE.format(
                dimension_name=dimension_name,
                dimension_instructions=dimension_instructions,
            )
            result = run_full_novel_processor(
                client=client,
                command_name="/draft-pass full",
                chunk_system_prompt=draft_system_prompt,
                synthesis_system_prompt=(
                    "You are combining multiple analysis reports into one final coherent report.\n\n"
                    "Rules:\n"
                    "- Remove duplicate issues\n"
                    "- Merge similar findings\n"
                    "- Preserve factual accuracy\n"
                    "- Maintain bullet structure\n"
                    "- Do NOT invent new issues\n"
                    "- Do NOT give writing advice unless original command allows it."
                ),
                chunks=chapter_chunks,
                temperature=DRAFT_PASS_TEMPERATURE,
            )
        else:
            result = request_chat_completion(
                client=client,
                messages=build_draft_pass_messages(
                    text_to_analyse=text_to_analyse,
                    dimension_name=dimension_name,
                    dimension_instructions=dimension_instructions,
                ),
                temperature=DRAFT_PASS_TEMPERATURE,
            )
    except Exception as exc:  # Keep terminal app stable for the user.
        print(f"Draft pass failed: {exc}")
        return

    print()
    print(result)


def handle_build_book() -> None:
    """Compile all numbered chapter files into a manuscript file."""
    if not CHAPTERS_DIR.exists():
        print("No chapter files found.")
        return

    chapter_paths = load_sorted_chapter_paths()
    if not chapter_paths:
        print("No chapter files found.")
        return
    warn_for_missing_chapter_files(chapter_paths)

    manuscript_text = build_manuscript_text(chapter_paths)

    try:
        MANUSCRIPT_DIR.mkdir(parents=True, exist_ok=True)
        atomic_write(MANUSCRIPT_PATH, manuscript_text)
    except OSError as exc:
        print(f"Manuscript build failed: {exc}")
        return

    print("Manuscript built successfully.")



def handle_draft_save() -> None:
    """Snapshot all chapters into a timestamped draft manuscript."""
    if not CHAPTERS_DIR.exists():
        print("No chapter files found.")
        return

    chapter_paths = load_sorted_chapter_paths()
    if not chapter_paths:
        print("No chapter files found.")
        return

    draft_text = build_manuscript_text(chapter_paths)
    timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M")
    draft_path = DRAFTS_DIR / f"draft_{timestamp}.txt"

    try:
        DRAFTS_DIR.mkdir(parents=True, exist_ok=True)
        if draft_path.exists():
            print("Draft snapshot already exists for this minute. Try again in a moment.")
            return
        atomic_write(draft_path, draft_text)
    except OSError as exc:
        print(f"Draft snapshot failed: {exc}")
        return

    print("Draft snapshot saved.")


def load_sorted_draft_paths() -> list[Path]:
    """Return saved draft files sorted by filename."""
    if not DRAFTS_DIR.exists():
        return []

    draft_paths: list[Path] = []
    for path in DRAFTS_DIR.iterdir():
        if path.is_file() and path.name.startswith("draft_") and path.suffix == ".txt":
            draft_paths.append(path)
    return sorted(draft_paths, key=lambda path: path.name)


def handle_draft_list() -> list[Path]:
    """Print a numbered list of draft files and return them."""
    draft_paths = load_sorted_draft_paths()
    if not draft_paths:
        print("No drafts found.")
        return []

    for index, path in enumerate(draft_paths, start=1):
        print(f"{index}. {path.name}")
    return draft_paths


def split_manuscript_into_chapters(text: str) -> list[tuple[int, str]]:
    """Split manuscript text into numbered chapter bodies using CHAPTER headers."""
    lines = text.splitlines()
    chapters: list[tuple[int, str]] = []
    current_number: int | None = None
    current_lines: list[str] = []

    def commit_current() -> None:
        nonlocal current_number, current_lines
        if current_number is None:
            return
        chapter_text = "\n".join(current_lines).strip()
        chapters.append((current_number, chapter_text))
        current_number = None
        current_lines = []

    line_index = 0
    while line_index < len(lines):
        stripped = lines[line_index].strip()

        if stripped == "========================" and line_index + 1 < len(lines):
            chapter_match = CHAPTER_HEADER_PATTERN.fullmatch(lines[line_index + 1].strip())
            if chapter_match is not None:
                commit_current()
                current_number = int(chapter_match.group(1))
                line_index += 2
                if (
                    line_index < len(lines)
                    and lines[line_index].strip() == "========================"
                ):
                    line_index += 1
                if line_index < len(lines) and not lines[line_index].strip():
                    line_index += 1
                continue

        chapter_match = CHAPTER_HEADER_PATTERN.fullmatch(stripped)
        if chapter_match is not None:
            commit_current()
            current_number = int(chapter_match.group(1))
            line_index += 1
            if line_index < len(lines) and not lines[line_index].strip():
                line_index += 1
            continue

        if current_number is not None:
            current_lines.append(lines[line_index])
        line_index += 1

    commit_current()
    return chapters


def handle_draft_load() -> None:
    """Restore chapter files from a selected draft snapshot."""
    draft_paths = handle_draft_list()
    if not draft_paths:
        return

    if not prompt_for_confirmation("Restore will overwrite current manuscript. Continue? (y/n)"):
        print("Draft load aborted.")
        return

    print("Choose draft number:")
    try:
        choice = input("> ").strip()
    except EOFError:
        print()
        return

    if not choice.isdigit():
        print("Invalid draft selection.")
        return

    selected_index = int(choice) - 1
    if selected_index < 0 or selected_index >= len(draft_paths):
        print("Invalid draft selection.")
        return

    selected_draft = draft_paths[selected_index]
    try:
        draft_text = selected_draft.read_text(encoding="utf-8")
    except OSError as exc:
        print(f"Could not read draft: {exc}")
        return

    chapter_entries = split_manuscript_into_chapters(draft_text)
    if not chapter_entries:
        print("Draft could not be restored: no chapter markers found.")
        return

    try:
        current_manuscript_snapshot = build_manuscript_text(load_sorted_chapter_paths())
        create_operation_backup("chapter_restore", content=current_manuscript_snapshot)
        create_operation_backup("draft_restore", source_path=selected_draft)
        CHAPTERS_DIR.mkdir(parents=True, exist_ok=True)
        draft_chapter_numbers = {chapter_number for chapter_number, _ in chapter_entries}
        for existing_path in load_sorted_chapter_paths():
            existing_number = extract_chapter_number(existing_path)
            if existing_number is not None and existing_number not in draft_chapter_numbers:
                existing_path.unlink(missing_ok=True)
        for chapter_number, chapter_text in chapter_entries:
            chapter_path = CHAPTERS_DIR / f"chapter_{chapter_number}.txt"
            atomic_write(chapter_path, chapter_text + ("\n" if chapter_text else ""))
    except OSError as exc:
        print(f"Draft restore failed: {exc}")
        return

    print("Draft restored successfully.")


def handle_save_draft() -> None:
    """Legacy alias for /draft-save."""
    handle_draft_save()


def handle_list_drafts() -> list[Path]:
    """Legacy alias for /draft-list."""
    return handle_draft_list()


def handle_restore_draft() -> None:
    """Legacy alias for /draft-load."""
    handle_draft_load()


def handle_ideas(command_text: str = "") -> None:
    """Capture a freeform idea or list saved ideas without AI involvement."""
    if "--list" in command_text:
        if not IDEAS_PATH.exists():
            print("No ideas file found.")
            return
        ideas_text = IDEAS_PATH.read_text(encoding="utf-8").rstrip()
        if not ideas_text:
            print("No ideas saved yet.")
            return
        print("IDEAS LIST:\n")
        print(ideas_text)
        return

    print("Paste idea. Type END on a new line when finished.")
    idea_text = collect_multiline_input(end_marker="END")
    idea_text = clean_terminal_text(idea_text)

    if not idea_text:
        print("No idea entered.")
        return

    try:
        append_idea(idea_text)
    except OSError:
        print("Idea could not be saved.")


def load_structured_inspiration_data() -> dict[str, str]:
    """Load inspiration files by category across all inspiration-book folders."""
    category_filenames = OrderedDict(
        [
            ("prose", "prose.txt"),
            ("dialogue", "dialogue.txt"),
            ("description", "description.txt"),
            ("pacing", "pacing.txt"),
            ("tension", "tension.txt"),
            ("devices", "devices.txt"),
        ]
    )
    combined = {category: "" for category in category_filenames}

    if not INSPIRATIONS_DIR.exists() or not INSPIRATIONS_DIR.is_dir():
        return combined

    for book_dir in sorted(path for path in INSPIRATIONS_DIR.iterdir() if path.is_dir()):
        for category, filename in category_filenames.items():
            file_path = book_dir / filename
            if not file_path.exists() or not file_path.is_file():
                continue
            try:
                content = clean_terminal_text(file_path.read_text(encoding="utf-8")).strip()
            except OSError:
                continue
            if not content:
                continue
            if combined[category]:
                combined[category] += "\n\n"
            combined[category] += content

    return combined


def detect_inspiration_categories(user_text: str) -> dict[str, tuple[bool, str]]:
    """Heuristically detect which inspiration categories are present in user text."""
    lowered = user_text.lower()
    lines = [line.strip() for line in user_text.splitlines() if line.strip()]
    words = re.findall(r"[A-Za-z']+", lowered)
    word_set = set(words)

    has_dialogue = any(marker in user_text for marker in INSPIRATION_DIALOGUE_MARKERS)

    has_devices = False
    if any(marker in lowered for marker in INSPIRATION_DEVICES_MARKERS):
        has_devices = True
    if any(re.match(r"^[A-Z][A-Z ]{2,}:\s+", line) for line in lines):
        has_devices = True

    has_description = any(term in word_set for term in INSPIRATION_SENSORY_TERMS)

    has_tension = any(term in lowered for term in INSPIRATION_TENSION_TERMS)
    if "?" in user_text:
        has_tension = True

    return {
        "prose": (True, "always present"),
        "dialogue": (has_dialogue, "quotation marks detected" if has_dialogue else "no dialogue present"),
        "description": (
            has_description,
            "environment or sensory detail detected" if has_description else "little to no sensory/environmental detail detected",
        ),
        "pacing": (True, "always present (sentence structure)"),
        "tension": (
            has_tension,
            "stakes/uncertainty/pressure detected" if has_tension else "no clear stakes, uncertainty, or pressure detected",
        ),
        "devices": (has_devices, "structural devices detected" if has_devices else "no structural devices used"),
    }


def handle_inspiration(client: Any) -> None:
    """Compare user scene techniques against structured inspirations grouped by category."""
    print("Paste your scene. Type END on a new line when finished.")
    user_text = collect_multiline_input(end_marker="END")
    if not user_text:
        print("No scene provided.")
        return

    combined = load_structured_inspiration_data()
    if all(not section.strip() for section in combined.values()):
        print("No inspiration data found.")
        return
    category_presence = detect_inspiration_categories(user_text)
    category_presence_block = "\n".join(
        f"- {name.upper()}: {'present' if present else 'not present'} ({reason})"
        for name, (present, reason) in category_presence.items()
    )

    user_payload = (
        "Analyze the user's writing against the structured inspiration sections.\n\n"
        "Use the provided category presence detection before scoring.\n"
        "Only score categories marked present.\n\n"
        "DETECTED CATEGORY PRESENCE:\n"
        f"{category_presence_block}\n\n"
        f"USER TEXT:\n{user_text}\n\n"
        f"PROSE TECHNIQUES:\n{combined['prose'] or '(none)'}\n\n"
        f"DIALOGUE TECHNIQUES:\n{combined['dialogue'] or '(none)'}\n\n"
        f"DESCRIPTION TECHNIQUES:\n{combined['description'] or '(none)'}\n\n"
        f"PACING TECHNIQUES:\n{combined['pacing'] or '(none)'}\n\n"
        f"TENSION TECHNIQUES:\n{combined['tension'] or '(none)'}\n\n"
        f"NARRATIVE DEVICES:\n{combined['devices'] or '(none)'}"
    )

    try:
        analysis = request_chat_completion(
            client=client,
            messages=[
                {"role": "system", "content": INSPIRATION_ANALYSIS_SYSTEM_PROMPT},
                {"role": "user", "content": user_payload},
            ],
            temperature=SCENE_TEMPERATURE,
        )
    except Exception as exc:  # Keep terminal app stable for the user.
        print(f"Inspiration analysis failed: {exc}")
        return

    print()
    print(analysis.strip())


def handle_novel_stats() -> None:
    """Print chapter and wordcount telemetry for the current novel project."""
    ensure_project_files()
    chapter_paths = load_sorted_chapter_paths()
    if not chapter_paths:
        print("No chapter files found.")
        return

    chapter_wordcounts: list[tuple[str, int]] = []
    total_words = 0
    for chapter_path in chapter_paths:
        chapter_text = clean_terminal_text(chapter_path.read_text(encoding="utf-8"))
        wordcount = len(chapter_text.split())
        chapter_wordcounts.append((chapter_path.name, wordcount))
        total_words += wordcount

    total_chapters = len(chapter_wordcounts)
    average_words = total_words / total_chapters if total_chapters else 0.0
    longest_chapter = max(chapter_wordcounts, key=lambda item: item[1])
    shortest_chapter = min(chapter_wordcounts, key=lambda item: item[1])

    print(f"Total chapters: {total_chapters}")
    print(f"Total wordcount: {total_words}")
    print(f"Average words per chapter: {average_words:.2f}")
    print(f"Longest chapter: {longest_chapter[0]} ({longest_chapter[1]} words)")
    print(f"Shortest chapter: {shortest_chapter[0]} ({shortest_chapter[1]} words)")


def command_research_world(client: Any) -> None:
    """Analyse world.txt for plausibility and save a structured report."""
    ensure_project_files()
    world_file = NOVEL_PROJECT_DIR / "world.txt"

    if not world_file.exists() or not world_file.is_file():
        print("No world.txt file found.")
        return

    try:
        with open(world_file, "r", encoding="utf-8") as file_handle:
            world_text = clean_terminal_text(file_handle.read()).strip()
    except OSError as exc:
        print(f"Could not read world.txt: {exc}")
        return

    if not world_text:
        print("world.txt is empty.")
        return

    messages = [
        {
            "role": "system",
            "content": (
                "You are a strict worldbuilding plausibility auditor. "
                "Analyse the provided world file for physical plausibility "
                "(technology, geography, biology, environment), social plausibility "
                "(economics, politics, culture), internal consistency, logical contradictions, "
                "and realism versus intentional stylisation. Provide concrete suggestions "
                "to strengthen believability.\n\n"
                "Return output in this exact format:\n\n"
                "=== WORLD PLAUSIBILITY REPORT ===\n\n"
                "Overall Plausibility Rating: (High / Medium / Low)\n\n"
                "Strengths:\n"
                "- bullet points\n\n"
                "Concerns:\n"
                "- bullet points\n\n"
                "Contradictions Detected:\n"
                "- bullet points\n\n"
                "Suggestions:\n"
                "- bullet points"
            ),
        },
        {"role": "user", "content": f"world.txt contents:\n\n{world_text}"},
    ]

    try:
        report = request_chat_completion(
            client=client,
            messages=messages,
            temperature=RESEARCH_TEMPERATURE,
        ).strip()
    except Exception as exc:  # Keep terminal app stable for the user.
        print(f"World plausibility analysis failed: {exc}")
        return

    timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    report_path = WORLD_PLAUSIBILITY_REPORTS_DIR / f"world_plausibility_{timestamp}.txt"

    try:
        WORLD_PLAUSIBILITY_REPORTS_DIR.mkdir(parents=True, exist_ok=True)
        atomic_write(report_path, report + "\n")
    except OSError as exc:
        print(f"Could not save world plausibility report: {exc}")
        return

    print("World plausibility analysis complete.")
    print("Full report saved to analysis folder.")


def handle_research(client: Any, command_text: str = "") -> None:
    """Route /research options to concrete research handlers."""
    parts = command_text.strip().split()
    command = parts[0] if parts else ""
    args = parts[1:] if len(parts) > 1 else []
    if command == "/research" and "--world" in args:
        command_research_world(client)
        return
    print("Unsupported research option. Use /research --world")


def handle_system_health() -> None:
    """Run a deep non-AI filesystem diagnostic of novel project health."""
    now_utc = datetime.utcnow()

    chapters_dir = Path.home() / "writing" / "novel_project" / "chapters"
    canon_memory_path = Path.home() / "writing" / "novel_project" / "memory" / "canon_memory.txt"
    drafts_dir = Path.home() / "writing" / "novel_project" / "drafts"
    research_dir = Path.home() / "writing" / "novel_project" / "research"
    analysis_dir = Path.home() / "writing" / "novel_project" / "analysis"
    continuity_logs_dir = analysis_dir / "continuity_reports"
    timeline_logs_dir = analysis_dir / "timeline_logs"
    rebuild_logs_dir = analysis_dir / "rebuild_logs"

    def safe_read_text(path: Path) -> str:
        try:
            return path.read_text(encoding="utf-8")
        except OSError:
            return ""

    def word_count(text: str) -> int:
        return len(re.findall(r"\b\w+\b", text))

    def format_bytes(byte_count: int) -> str:
        mb = byte_count / (1024 * 1024)
        return f"{mb:.2f} MB"

    def list_files(path: Path) -> list[Path]:
        if not path.exists() or not path.is_dir():
            return []
        return [item for item in path.iterdir() if item.is_file()]

    severity_rank = {"GOOD": 0, "STABLE": 1, "WARNING": 2, "CRITICAL": 3}
    overall_status = "GOOD"
    warnings: list[str] = []

    # --------------------------------------------------------
    # Check area 1 — Manuscript health
    # --------------------------------------------------------
    chapter_entries: list[tuple[int, Path]] = []
    if chapters_dir.exists() and chapters_dir.is_dir():
        for path in chapters_dir.iterdir():
            if not path.is_file():
                continue
            match = CHAPTER_FILENAME_PATTERN.fullmatch(path.name)
            if match is None:
                continue
            chapter_entries.append((int(match.group(1)), path))
    chapter_entries.sort(key=lambda item: item[0])

    chapter_numbers = [number for number, _ in chapter_entries]
    missing_numbers: list[int] = []
    if chapter_numbers:
        expected_numbers = set(range(chapter_numbers[0], chapter_numbers[-1] + 1))
        missing_numbers = sorted(expected_numbers - set(chapter_numbers))

    words_per_chapter: list[tuple[int, int]] = []
    total_words = 0
    flagged_chapters: list[str] = []
    for chapter_number, chapter_path in chapter_entries:
        chapter_text = safe_read_text(chapter_path)
        chapter_words = word_count(chapter_text)
        words_per_chapter.append((chapter_number, chapter_words))
        total_words += chapter_words
        if chapter_words == 0:
            flagged_chapters.append(f"chapter_{chapter_number}: EMPTY")
        elif chapter_words < 1200:
            flagged_chapters.append(f"chapter_{chapter_number}: SHORT")
        elif chapter_words > 7000:
            flagged_chapters.append(f"chapter_{chapter_number}: VERY LONG")

    if total_words > 150000:
        warnings.append("Total novel size exceeds 150k words (HIGH scale warning).")

    # --------------------------------------------------------
    # Check area 2 — Canon memory health
    # --------------------------------------------------------
    total_facts = 0
    duplicate_facts_count = 0
    unresolved_states = 0
    chapter_headers_without_facts = 0
    canon_memory_characters = 0

    if not canon_memory_path.exists():
        warnings.append("Canon memory file missing.")
    else:
        canon_content = safe_read_text(canon_memory_path)
        canon_memory_characters = len(canon_content)
        parsed_chapters = parse_canon_memory(canon_content)

        normalized_facts: list[str] = []
        for chapter in parsed_chapters:
            chapter_fact_total = sum(
                len(facts)
                for facts in chapter.get("categories", {}).values()
            )
            if chapter_fact_total == 0:
                chapter_headers_without_facts += 1
            total_facts += chapter_fact_total

            for category_facts in chapter.get("categories", {}).values():
                for fact in category_facts:
                    normalized = get_fact_text(fact).strip().lower()
                    if normalized:
                        normalized_facts.append(normalized)

            unresolved_states += sum(
                1
                for story_state in chapter.get("story_states", [])
                if str(story_state.get("state", "ACTIVE")).upper() == "ACTIVE"
            )

        fact_counts = Counter(normalized_facts)
        duplicate_facts_count = sum(count - 1 for count in fact_counts.values() if count > 1)

    # --------------------------------------------------------
    # Check area 3 — Draft system health
    # --------------------------------------------------------
    draft_files = list_files(drafts_dir)
    total_drafts = len(draft_files)
    latest_draft_time: datetime | None = None
    timestamp_keys: list[str] = []
    for draft_file in draft_files:
        try:
            stat = draft_file.stat()
        except OSError:
            continue
        modified = datetime.utcfromtimestamp(stat.st_mtime)
        if latest_draft_time is None or modified > latest_draft_time:
            latest_draft_time = modified

        timestamp_match = re.fullmatch(r"draft_(\d{8}_\d{4})(?:_.+)?\.txt", draft_file.name)
        if timestamp_match is not None:
            timestamp_keys.append(timestamp_match.group(1))

    if latest_draft_time is None:
        warnings.append("No drafts found.")
        last_draft_display = "NONE"
    else:
        last_draft_display = latest_draft_time.strftime("%Y-%m-%d %H:%M UTC")
        if latest_draft_time < now_utc - timedelta(days=7):
            warnings.append("No draft saved in the last 7 days.")

    duplicate_timestamp_collisions = sum(
        1
        for count in Counter(timestamp_keys).values()
        if count > 1
    )
    if duplicate_timestamp_collisions > 0:
        warnings.append("Duplicate draft timestamp naming collision risk detected.")

    # --------------------------------------------------------
    # Check area 4 — Research system health
    # --------------------------------------------------------
    research_files = list_files(research_dir)
    total_research_topics = 0
    research_characters = 0
    for research_file in research_files:
        if research_file.parent == research_dir and research_file.suffix.lower() == ".txt":
            total_research_topics += 1
        research_characters += len(safe_read_text(research_file))

    integrity_reports_present = (research_dir / "integrity_reports").exists()
    if research_characters > 200000:
        warnings.append("Research corpus exceeds 200k characters (realism audit scaling risk).")

    # --------------------------------------------------------
    # Check area 5 — Analysis log health
    # --------------------------------------------------------
    def collect_log_stats(path: Path) -> tuple[int, int]:
        files = list_files(path)
        total_size = 0
        for file_path in files:
            try:
                total_size += file_path.stat().st_size
            except OSError:
                continue
        return len(files), total_size

    continuity_count, continuity_size = collect_log_stats(continuity_logs_dir)
    timeline_count, timeline_size = collect_log_stats(timeline_logs_dir)
    rebuild_count, rebuild_size = collect_log_stats(rebuild_logs_dir)

    for log_name, log_size in (
        ("continuity_reports", continuity_size),
        ("timeline_logs", timeline_size),
        ("rebuild_logs", rebuild_size),
    ):
        if log_size > 20 * 1024 * 1024:
            warnings.append(f"{log_name} exceeds 20MB (log cleanup recommended).")

    # --------------------------------------------------------
    # Check area 6 — Chunking / performance risk
    # --------------------------------------------------------
    average_chapter_words = (
        total_words / len(words_per_chapter) if words_per_chapter else 0.0
    )
    chunking_risk = "LOW"
    memory_scale_risk = "LOW"
    rebuild_scale_risk = "LOW"

    if len(words_per_chapter) > 25 or average_chapter_words > 5000:
        chunking_risk = "HIGH"
        warnings.append("Chunking/performance risk elevated for full-novel operations.")

    if canon_memory_characters > 80000:
        memory_scale_risk = "HIGH"
        warnings.append("Canon memory exceeds 80k characters (memory drift risk).")

    if rebuild_count > 100:
        rebuild_scale_risk = "HIGH"
        warnings.append("Rebuild logs exceed 100 files (maintenance required).")

    if any("HIGH" in warning for warning in warnings):
        overall_status = "CRITICAL"
    elif warnings:
        overall_status = "WARNING"
    elif len(words_per_chapter) == 0 and total_drafts == 0 and total_research_topics == 0:
        overall_status = "STABLE"

    if severity_rank[overall_status] < severity_rank["WARNING"] and chapter_headers_without_facts > 0:
        overall_status = "STABLE"

    if overall_status == "CRITICAL":
        recommended_action = "Run /draft-save immediately and clean high-risk areas."
    elif overall_status == "WARNING":
        recommended_action = "Run /draft-save soon and resolve warnings."
    elif overall_status == "STABLE":
        recommended_action = "Proceed with writing and monitor weekly."
    else:
        recommended_action = "Continue normal workflow."

    missing_numbers_display = ", ".join(str(number) for number in missing_numbers) if missing_numbers else "None"
    flagged_chapter_display = "; ".join(flagged_chapters) if flagged_chapters else "None"

    print("SYSTEM HEALTH REPORT")
    print()
    print("Manuscript:")
    print(f"- total chapters: {len(words_per_chapter)}")
    print(f"- missing chapter numbers: {missing_numbers_display}")
    print(f"- total words: {total_words}")
    print(f"- flagged chapters: {flagged_chapter_display}")
    print()
    print("Canon Memory:")
    print(f"- total facts: {total_facts}")
    print(f"- duplicates: {duplicate_facts_count}")
    print(f"- unresolved states: {unresolved_states}")
    print()
    print("Draft System:")
    print(f"- total drafts: {total_drafts}")
    print(f"- last draft: {last_draft_display}")
    print()
    print("Research System:")
    print(f"- total research topics: {total_research_topics}")
    print(f"- integrity reports present: {'YES' if integrity_reports_present else 'NO'}")
    print()
    print("Analysis Logs:")
    print(f"- continuity logs: {continuity_count} files ({format_bytes(continuity_size)})")
    print(f"- timeline logs: {timeline_count} files ({format_bytes(timeline_size)})")
    print(f"- rebuild logs: {rebuild_count} files ({format_bytes(rebuild_size)})")
    print()
    print("Performance Risk:")
    print(f"- chunking risk: {chunking_risk}")
    print(f"- memory scale risk: {memory_scale_risk}")
    print(f"- rebuild scale risk: {rebuild_scale_risk}")
    print()
    print(f"OVERALL STATUS: {overall_status}")
    print()
    print("Recommended Next Action:")
    print(f"- {recommended_action}")

    if warnings:
        print()
        print("Warnings:")
        for warning in warnings:
            print(f"- {warning}")


TIMELINE_STOPWORDS = {
    "a", "an", "and", "approaches", "as", "at", "be", "been", "being", "by",
    "for", "from", "has", "have", "in", "into", "is", "it", "its", "later",
    "of", "on", "or", "the", "their", "them", "to", "toward", "with",
    "active", "resolved", "increase", "increases", "increased", "rises", "rose",
    "rising", "shows", "showing", "signs", "becomes", "becoming", "create",
    "creates", "created", "indicates", "indicating", "reveals", "revealing",
    "marks", "marked", "set", "sets", "setting", "stage",
}


ENVIRONMENT_PLOT_KEYWORDS = {
    "alarm", "anomaly", "breach", "collapse", "countdown", "danger", "debris",
    "environment", "evacuation", "hazard", "impact", "lockdown", "radiation",
    "route", "safe", "scan", "sensor", "shelter", "signal", "storm", "survival",
    "threat", "window",
}


DISCOVERY_KEYWORDS = {
    "discover", "discovery", "find", "found", "identify", "identified", "learn",
    "learns", "realise", "realises", "realize", "realizes", "reveal", "reveals",
    "revealed", "truth", "clue", "evidence", "decode", "decoded", "message",
    "signal", "secret", "unknown", "uncover", "understand", "insight",
}


CONFLICT_KEYWORDS = {
    "argument", "attack", "betrayal", "clash", "collision", "conflict", "crisis",
    "danger", "fear", "fight", "fracture", "hostile", "hostility", "opposes",
    "alarm", "radiation",
    "pressure", "risk", "rupture", "strain", "standoff", "struggle", "tension",
    "threat", "urgent", "urgency", "warning",
}


MISSION_KEYWORDS = {
    "approach", "arrival", "countdown", "crew", "deadline", "deploy", "deployment",
    "engine", "escape", "evacuate", "journey", "launch", "mission", "objective",
    "plan", "preparation", "prepare", "prepared", "progress", "ready", "readiness",
    "reactor", "repair", "response", "route", "schedule", "shield", "ship",
    "survival", "system", "window",
}


CHARACTER_KEYWORDS = {
    "afraid", "doubt", "hesitation", "injury", "pain", "panic", "secret", "shame",
    "sick", "weakness", "wound",
}


RELATIONSHIP_KEYWORDS = {
    "alliance", "blame", "bond", "crew", "distrust", "loyalty", "promise", "responsibility",
    "rift", "team", "trust",
}


def tokenize_timeline_text(text: str) -> set[str]:
    """Return a compact set of meaningful tokens for timeline thread matching."""
    return {
        token
        for token in re.findall(r"[a-z0-9']+", text.lower())
        if len(token) > 3 and token not in TIMELINE_STOPWORDS
    }



def extract_temporal_references(text: str) -> set[str]:
    """Extract explicit temporal markers used for safer timeline merging."""
    return {
        match.group(0).lower()
        for match in re.finditer(
            r"\b(?:day|week|month|year|hour|minute)\s+\d+\b|\b(?:before|after|during|later|earlier|tonight|tomorrow|yesterday)\b",
            text.lower(),
        )
    }


def infer_event_type(tokens: set[str]) -> str:
    """Infer coarse event type for timeline merge matching."""
    if tokens & INJURY_KEYWORDS:
        return "injury"
    if tokens & RELATIONSHIP_KEYWORDS:
        return "relationship"
    if tokens & CONFLICT_KEYWORDS:
        return "conflict"
    if tokens & DISCOVERY_KEYWORDS:
        return "discovery"
    if tokens & MISSION_KEYWORDS:
        return "mission"
    return "generic"


def extract_unique_nouns(tokens: set[str]) -> set[str]:
    """Extract less-generic continuity anchors from tokens."""
    generic_tokens = TIMELINE_STOPWORDS | {"mission", "crew", "team", "system", "base"}
    return {token for token in tokens if len(token) >= 5 and token not in generic_tokens}


def extract_character_tokens(tokens: set[str]) -> set[str]:
    """Extract likely character-name tokens for merge logic."""
    return {token for token in tokens if token in CHARACTER_KEYWORDS or token.startswith(("dr", "capt", "cmdr"))}


def is_plot_relevant_event(category: str, text: str) -> bool:
    """Filter out low-impact atmosphere notes from the timeline view."""
    if category not in {"World", "Location"}:
        return True
    tokens = tokenize_timeline_text(text)
    return (
        bool(tokens & ENVIRONMENT_PLOT_KEYWORDS)
        or bool(tokens & DISCOVERY_KEYWORDS)
        or bool(tokens & CONFLICT_KEYWORDS)
    )



def infer_timeline_category(category: str, text: str, state: str) -> str:
    """Map canon-memory categories into the narrative timeline categories."""
    tokens = tokenize_timeline_text(text)

    if category in {
        "Relationship",
        "Relationship State — Active",
        "Relationship State — Resolved",
    }:
        return "Relationship"
    if category in {
        "Character",
        "Injury",
        "Psychological State — Active",
        "Psychological State — Resolved",
    }:
        if tokens & RELATIONSHIP_KEYWORDS:
            return "Relationship"
        return "Character"
    if category in {
        "Mission State — Active",
        "Mission State — Resolved",
        "Technology State — Active",
        "Technology State — Resolved",
    }:
        if tokens & CONFLICT_KEYWORDS and state == "ACTIVE":
            return "Conflict"
        return "Mission"
    if category in {"World", "Location"}:
        if tokens & DISCOVERY_KEYWORDS:
            return "Discovery"
        if tokens & CONFLICT_KEYWORDS:
            return "Conflict"
        return "Environment"
    if category == "Object":
        if tokens & DISCOVERY_KEYWORDS:
            return "Discovery"
        if tokens & MISSION_KEYWORDS:
            return "Mission"
        return "Environment"

    if tokens & DISCOVERY_KEYWORDS:
        return "Discovery"
    if tokens & CONFLICT_KEYWORDS:
        return "Conflict"
    if tokens & CHARACTER_KEYWORDS:
        return "Character"
    if tokens & RELATIONSHIP_KEYWORDS:
        return "Relationship"
    if tokens & MISSION_KEYWORDS or category == "Timeline":
        return "Mission"
    return "Environment"



def build_timeline_event(
    chapter_number: int,
    source_category: str,
    fact: dict[str, Any],
) -> dict[str, Any] | None:
    """Convert one canon-memory fact into a narrative timeline event."""
    text = get_fact_text(fact)
    state = str(fact.get("state", "ACTIVE")).upper()
    if not text or not is_plot_relevant_event(source_category, text):
        return None

    tokens = tokenize_timeline_text(text)
    return {
        "chapter": chapter_number,
        "source_category": source_category,
        "text": text,
        "state": "RESOLVED" if state == "RESOLVED" else "ACTIVE",
        "narrative_category": infer_timeline_category(source_category, text, state),
        "tokens": tokens,
        "event_type": infer_event_type(tokens),
        "character_tokens": extract_character_tokens(tokens),
        "unique_nouns": extract_unique_nouns(tokens),
        "temporal_references": extract_temporal_references(text),
    }



def events_share_thread(event_a: dict[str, Any], event_b: dict[str, Any]) -> bool:
    """Return True when two timeline events likely describe the same narrative thread."""
    if facts_are_similar(event_a["text"], event_b["text"]):
        return True

    if event_a["unique_nouns"] & event_b["unique_nouns"]:
        return True

    if (
        event_a["event_type"] == event_b["event_type"]
        and event_a["character_tokens"] & event_b["character_tokens"]
    ):
        return True

    if event_a["temporal_references"] & event_b["temporal_references"]:
        return True

    return False



def merge_timeline_events(chapters: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """Merge duplicate timeline threads into a cleaner chronological event list."""
    merged_events: list[dict[str, Any]] = []
    seen_threads: list[dict[str, Any]] = []

    prioritized_categories = [
        "Timeline",
        "Mission State — Active",
        "Mission State — Resolved",
        "Technology State — Active",
        "Technology State — Resolved",
        "Psychological State — Active",
        "Psychological State — Resolved",
        "Relationship State — Active",
        "Relationship State — Resolved",
        "Character",
        "Relationship",
        "Injury",
        "Foreshadowing Setup",
        "Foreshadowing Payoff",
        "Object",
        "World",
        "Location",
    ]

    for chapter in sorted(chapters, key=lambda item: item["number"]):
        chapter_categories = chapter["categories"]
        ordered_category_names = [
            name for name in prioritized_categories if name in chapter_categories
        ]
        ordered_category_names.extend(
            name for name in chapter_categories if name not in ordered_category_names
        )

        for category_name in ordered_category_names:
            for fact in chapter_categories.get(category_name, []):
                event = build_timeline_event(chapter["number"], category_name, fact)
                if event is None:
                    continue

                existing_thread = next(
                    (thread for thread in seen_threads if events_share_thread(thread, event)),
                    None,
                )

                if existing_thread is None:
                    thread_event = dict(event)
                    seen_threads.append(thread_event)
                    merged_events.append(event)
                    continue

                if event["state"] == existing_thread["state"]:
                    existing_thread["tokens"] = existing_thread["tokens"] | event["tokens"]
                    if event["narrative_category"] in {"Conflict", "Discovery"}:
                        existing_thread["narrative_category"] = event["narrative_category"]
                    if len(event["text"]) > len(existing_thread["text"]):
                        existing_thread["text"] = event["text"]
                    continue

                existing_thread["state"] = event["state"]
                existing_thread["text"] = event["text"]
                existing_thread["chapter"] = event["chapter"]
                existing_thread["tokens"] = event["tokens"]
                existing_thread["narrative_category"] = event["narrative_category"]
                merged_events.append(event)

    return merged_events



def render_timeline_overview(chapters: list[dict[str, Any]]) -> str:
    """Render a readable, chronological story map from canon memory."""
    timeline_events = merge_timeline_events(chapters)
    if not timeline_events:
        return "No timeline events recorded yet."

    lines: list[str] = []
    current_chapter: int | None = None
    for event in timeline_events:
        if event["chapter"] != current_chapter:
            if lines:
                lines.append("")
            current_chapter = event["chapter"]
            lines.append(f"Chapter {current_chapter}")

        bullet = "✓" if event["state"] == "RESOLVED" else "•"
        lines.append(
            f"{bullet} [{event['narrative_category']}] {event['text']} ({event['state']})"
        )

    return "\n".join(lines)


def handle_timeline_view() -> None:
    ensure_project_files()

    if not CANON_MEMORY_PATH.exists():
        print("No canon memory found.")
        return

    content = CANON_MEMORY_PATH.read_text(encoding="utf-8")
    chapters = parse_canon_memory(content)

    print("\nTIMELINE OVERVIEW\n")
    print(render_timeline_overview(chapters))


def handle_story_state() -> None:
    ensure_project_files()

    if not CANON_MEMORY_PATH.exists():
        print("No canon memory found.")
        return

    content = CANON_MEMORY_PATH.read_text(encoding="utf-8")
    chapters = parse_canon_memory(content)

    print("\nSTORY STATE\n")

    active_found = False

    for chapter in sorted(chapters, key=lambda c: c["number"]):
        for story_state in chapter.get("story_states", []):
            if str(story_state.get("state", "ACTIVE")).upper() != "ACTIVE":
                continue
            description = str(story_state.get("description", "")).strip()
            if not description:
                continue
            active_found = True
            print(f"- {description}")
            print(f"  FIRST_SEEN: Chapter {int(story_state.get('first_seen', chapter['number']))}")
            print()

    if not active_found:
        print("No active narrative states recorded.")


def handle_research_topic(client: Any) -> None:
    """Run an isolated factual science research workflow."""
    depth_choice = prompt_for_research_choice(
        "Choose research depth:",
        RESEARCH_DEPTH_OPTIONS,
    )
    if depth_choice is None:
        return

    style_choice = prompt_for_research_choice(
        "Choose output style:",
        RESEARCH_STYLE_OPTIONS,
    )
    if style_choice is None:
        return

    print("Paste research question. Type END on new line when finished.")
    question = collect_multiline_input(end_marker="END")
    if not question:
        print("No research question entered.")
        return

    try:
        research_notes = request_chat_completion(
            client=client,
            messages=build_research_messages(
                question=question,
                depth_choice=depth_choice,
                style_choice=style_choice,
            ),
            temperature=RESEARCH_TEMPERATURE,
        )
    except Exception as exc:  # Keep terminal app stable for the user.
        print(f"Research request failed: {exc}")
        return

    print()
    print(research_notes)

    if not prompt_for_confirmation("Save this research topic? (y/n)"):
        return

    print("Topic filename (example: neutron_star_collision):")
    try:
        filename_input = input("> ").strip()
    except EOFError:
        print()
        return

    if not filename_input:
        print("No topic filename entered.")
        return

    safe_topic_name = re.sub(r"[^a-zA-Z0-9_-]+", "_", filename_input).strip("_")
    if not safe_topic_name:
        print("Invalid topic filename.")
        return

    RESEARCH_DIR.mkdir(parents=True, exist_ok=True)
    output_path = RESEARCH_DIR / f"{safe_topic_name}.txt"
    if output_path.exists():
        timestamp_suffix = datetime.utcnow().strftime("%Y%m%d_%H%M")
        output_path = RESEARCH_DIR / f"{safe_topic_name}_{timestamp_suffix}.txt"

    file_body = (
        f"TOPIC: {safe_topic_name}\n\n"
        "QUESTION:\n"
        f"{question}\n\n"
        "RESEARCH NOTES:\n"
        f"{research_notes}\n"
    )

    try:
        atomic_write(output_path, file_body)
    except OSError as exc:
        print(f"Could not save research topic: {exc}")
        return

    print("Research topic saved.")


def handle_research_scene(client: Any) -> None:
    """Run chunk-safe hard-science realism analysis on pasted scene text."""
    print("Paste scene text. Type END on new line when finished.")
    scene_text = collect_multiline_input(end_marker="END")
    if not scene_text:
        print("No scene provided.")
        return

    scene_chunks = chunk_text_blocks([scene_text], max_chars=10000)
    if not scene_chunks:
        print("No scene provided.")
        return

    canon_memory_block = (
        read_text_file(CANON_MEMORY_PATH)
        if CANON_MEMORY_PATH.exists()
        else "(not provided)"
    )
    world_rules_block = (
        read_text_file(WORLD_RULES_PATH)
        if WORLD_RULES_PATH.exists()
        else "(not provided)"
    )

    try:
        chunk_reports: list[str] = []
        prior_findings = ""
        for chunk_index, scene_chunk in enumerate(scene_chunks, start=1):
            if len(scene_chunks) > 1:
                print(f"Analyzing chunk {chunk_index}/{len(scene_chunks)}...")
            chunk_report = request_chat_completion(
                client=client,
                messages=build_research_scene_messages(
                    scene_chunk=scene_chunk,
                    canon_memory_block=canon_memory_block,
                    world_rules_block=world_rules_block,
                    prior_findings=prior_findings,
                ),
                temperature=RESEARCH_SCENE_TEMPERATURE,
            )
            chunk_reports.append(chunk_report)
            prior_findings = "\n\n".join(chunk_reports[-2:])

        if len(chunk_reports) == 1:
            final_report = chunk_reports[0]
        else:
            final_report = request_chat_completion(
                client=client,
                messages=[
                    {"role": "system", "content": RESEARCH_SCENE_SYSTEM_PROMPT},
                    {
                        "role": "user",
                        "content": (
                            "Merge these chunk-level scientific realism analyses into one final report.\n"
                            "Keep the exact output structure below. Remove duplicates. Preserve concrete findings.\n"
                            "Do not add prose rewrites, story suggestions, or narrative critique.\n\n"
                            "Output format (strict):\n"
                            "SCIENCE REALISM REPORT\n\n"
                            "Relevant Real Science\n"
                            "- bullet points\n\n"
                            "Realistic Constraints\n"
                            "- bullet points\n\n"
                            "Already Realistic Elements\n"
                            "- bullet points\n\n"
                            "Unrealistic or Risky Elements\n"
                            "- bullet points\n\n"
                            "Scientific Detail Opportunities\n"
                            "- micro realism ideas\n"
                            "- instrumentation behaviour\n"
                            "- environmental reactions\n\n"
                            + "\n\n".join(
                                f"Chunk {index} report:\n{chunk_report}"
                                for index, chunk_report in enumerate(chunk_reports, start=1)
                            )
                        ),
                    },
                ],
                temperature=RESEARCH_SCENE_TEMPERATURE,
            )
    except Exception as exc:  # Keep terminal app stable for the user.
        print(f"Research scene analysis failed: {exc}")
        return

    print()
    print(final_report)


def handle_research_apply(client: Any) -> None:
    """Apply saved research notes to a scene or chapter for realism conflict auditing."""
    if not RESEARCH_DIR.exists() or not RESEARCH_DIR.is_dir():
        print("No research topics found.")
        return

    topic_files = sorted(path for path in RESEARCH_DIR.iterdir() if path.is_file())
    if not topic_files:
        print("No research topics found.")
        return

    print("Available research topics:")
    for index, topic_path in enumerate(topic_files, start=1):
        print(f"{index}. {topic_path.name}")

    print("Select topic number.")
    try:
        topic_selection = input("> ").strip()
    except EOFError:
        print()
        return

    if not topic_selection.isdigit():
        print("Invalid selection.")
        return

    topic_index = int(topic_selection)
    if topic_index < 1 or topic_index > len(topic_files):
        print("Invalid selection.")
        return

    selected_topic_path = topic_files[topic_index - 1]
    research_notes = read_text_file(selected_topic_path)

    print("Choose input type:")
    print("1 = Paste scene")
    print("2 = Choose chapter file")
    try:
        input_type = input("> ").strip()
    except EOFError:
        print()
        return

    scene_text = ""
    if input_type == "1":
        print("Paste scene text. Type END on new line when finished.")
        scene_text = collect_multiline_input(end_marker="END")
        if not scene_text:
            print("No scene provided.")
            return
    elif input_type == "2":
        chapter_number = prompt_for_chapter_number("Ask chapter number.")
        if chapter_number is None:
            return
        chapter_path = CHAPTERS_DIR / f"chapter_{chapter_number}.txt"
        if not chapter_path.exists() or not chapter_path.is_file():
            print("Chapter file not found.")
            return
        scene_text = clean_terminal_text(chapter_path.read_text(encoding="utf-8"))
        if not scene_text:
            print("Chapter file is empty.")
            return
    else:
        print("Invalid selection.")
        return

    try:
        report = request_chat_completion(
            client=client,
            messages=build_research_apply_messages(
                research_notes=research_notes,
                scene_text=scene_text,
            ),
            temperature=RESEARCH_SCENE_TEMPERATURE,
        )
    except Exception as exc:  # Keep terminal app stable for the user.
        print(f"Research realism audit failed: {exc}")
        return

    print()
    print(report)


def handle_research_integrity(client: Any) -> None:
    """Audit all saved research topics for scientific consistency issues only."""
    RESEARCH_DIR.mkdir(parents=True, exist_ok=True)
    RESEARCH_INTEGRITY_REPORTS_DIR.mkdir(parents=True, exist_ok=True)

    topic_files = sorted(
        path
        for path in RESEARCH_DIR.iterdir()
        if path.is_file() and path.suffix.lower() == ".txt"
    )
    if not topic_files:
        print("No research topics found.")
        return

    corpus_parts: list[str] = []
    for topic_path in topic_files:
        topic_text = clean_terminal_text(topic_path.read_text(encoding="utf-8"))
        if not topic_text:
            continue
        corpus_parts.append(f"FILE: {topic_path.name}\n{topic_text}")

    if not corpus_parts:
        print("No research topics found.")
        return

    research_corpus = ("\n\n" + ("-" * 60) + "\n\n").join(corpus_parts)

    try:
        report = request_chat_completion(
            client=client,
            messages=build_research_integrity_messages(research_corpus),
            temperature=RESEARCH_TEMPERATURE,
        )
    except Exception as exc:  # Keep terminal app stable for the user.
        print(f"Research integrity audit failed: {exc}")
        return

    print()
    print(report)

    if not prompt_for_confirmation("Save integrity report? (y/n)"):
        return

    timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M")
    output_path = RESEARCH_INTEGRITY_REPORTS_DIR / f"integrity_{timestamp}.txt"
    suffix = 1
    while output_path.exists():
        output_path = RESEARCH_INTEGRITY_REPORTS_DIR / f"integrity_{timestamp}_{suffix}.txt"
        suffix += 1

    try:
        atomic_write(output_path, report)
    except OSError as exc:
        print(f"Could not save integrity report: {exc}")
        return

    print("Research integrity report saved.")


def handle_world_add() -> None:
    """Capture and save one structured world rule entry."""
    print("Enter world rule category:")
    try:
        category = input("> ").strip()
    except EOFError:
        print()
        return

    if not category:
        print("No category entered.")
        return

    print("Enter rule text:")
    try:
        rule_text = input("> ").strip()
    except EOFError:
        print()
        return

    if not rule_text:
        print("No rule text entered.")
        return

    try:
        append_world_rule(category, rule_text)
    except (OSError, ValueError) as exc:
        print(f"World rule could not be saved: {exc}")
        return

    print(f"World rule saved to {WORLD_RULES_PATH}.")



def handle_export_chapter() -> None:
    """Prepare a chapter text file for manual WordGrinder export."""
    ensure_project_files()
    chapter_number = prompt_for_chapter_number()
    if chapter_number is None:
        return

    export_path = CHAPTERS_DIR / f"chapter_{chapter_number}.txt"
    if export_path.exists() and not prompt_for_confirmation(
        "Overwrite existing export? (y/n)"
    ):
        print("Export cancelled.")
        return

    try:
        if export_path.exists():
            create_operation_backup("chapter_restore", source_path=export_path)
        atomic_write(export_path, "")
    except OSError as exc:
        print(f"Could not prepare export file: {exc}")
        return

    print()
    print("Now open your WordGrinder chapter and press:")
    print("CTRL+SHIFT+E (export)")
    print("Choose:")
    print("Plain Text")
    print("Save to:")
    print(export_path)
    print()
    print("Press ENTER after exporting the chapter.")

    try:
        input()
    except EOFError:
        print()

    if export_path.exists():
        print("Chapter export complete.")
        return

    print("Chapter export file was not found.")


def handle_export_book_docx() -> None:
    """Export all numbered chapter files to a DOCX manuscript."""
    ensure_project_files()

    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.shared import Inches, Pt
    except ImportError:
        print("python-docx not installed. Run: pip install python-docx")
        return

    chapter_paths = load_sorted_chapter_paths()
    if not chapter_paths:
        print("No chapter files found in ~/writing/novel_project/chapters/")
        return
    warn_for_missing_chapter_files(chapter_paths)

    title = "Untitled Novel"
    if CANON_MEMORY_PATH.exists() and CANON_MEMORY_PATH.is_file():
        canon_lines = CANON_MEMORY_PATH.read_text(encoding="utf-8").splitlines()
        if canon_lines:
            raw_first_line = clean_terminal_text(canon_lines[0])
            if raw_first_line:
                if raw_first_line.lower().startswith("title:"):
                    parsed_title = raw_first_line.partition(":")[2].strip()
                    if parsed_title:
                        title = parsed_title
                else:
                    title = raw_first_line

    document = Document()

    normal_style = document.styles["Normal"]
    normal_font = normal_style.font
    normal_font.name = "Times New Roman"
    normal_font.size = Pt(12)
    normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    normal_style.paragraph_format.first_line_indent = Inches(0.5)

    for section in document.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.add_run(title)

    subtitle_paragraph = document.add_paragraph()
    subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_paragraph.add_run("A Novel")

    by_paragraph = document.add_paragraph()
    by_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    by_paragraph.add_run("by")

    blank_author_paragraph = document.add_paragraph()
    blank_author_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    blank_author_paragraph.add_run("")

    document.add_page_break()

    scene_break_pattern = re.compile(r"^(?:-{3,}|\*{3,}|(?:\*\s*){3,}|(?:—\s*){3,})$")

    for path in chapter_paths:
        chapter_number = extract_chapter_number(path)
        if chapter_number is None:
            continue

        chapter_text = clean_terminal_text(path.read_text(encoding="utf-8"))

        document.add_page_break()
        heading = document.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.paragraph_format.first_line_indent = Inches(0)
        heading_run = heading.add_run(f"CHAPTER {chapter_number}")
        heading_run.font.name = "Times New Roman"
        heading_run.font.size = Pt(16)
        heading_run.bold = True

        for raw_line in chapter_text.splitlines():
            line = raw_line.strip()
            if not line:
                document.add_paragraph("")
                continue

            if scene_break_pattern.fullmatch(line):
                scene_break = document.add_paragraph()
                scene_break.alignment = WD_ALIGN_PARAGRAPH.CENTER
                scene_break.paragraph_format.first_line_indent = Inches(0)
                scene_break.add_run("---")
                continue

            body_paragraph = document.add_paragraph(line)
            body_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    output_path = MANUSCRIPT_DIR / "novel.docx"
    if output_path.exists():
        timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        output_path = MANUSCRIPT_DIR / f"novel_{timestamp}.docx"

    try:
        document.save(str(output_path))
    except Exception as exc:  # Keep terminal app stable for the user.
        print(f"DOCX export failed: {exc}")
        return

    output_display = str(output_path).replace(str(Path.home()), "~", 1)
    print("DOCX manuscript export complete:")
    print(output_display)


# ============================================================
# Main application loop
# ============================================================


def print_welcome() -> None:
    """Show a simple startup message."""
    print("Novel AI Assistant")
    print(f"Model: {MODEL_NAME}")
    print(f"Script path: {NOVEL_AI_SCRIPT_PATH}")
    print(f"Project path: {NOVEL_PROJECT_DIR}")
    print("Type /help for commands or type exit to quit.")


ALL_COMMANDS: list[str] = []


COMMAND_HELP_DETAILS = {
    "/scene-summary": """Purpose: Analyse a pasted scene and extract continuity-critical memory suggestions and story-state changes.
Files read: Canon memory, recent scene summaries, screenplay source (if present), and user-pasted scene text.
Files written: Canon memory, story state memory, scene summaries, and chapter files when user confirms updates.
AI usage: Yes.
Canon memory impact: Can append, resolve, or update canon facts and state-tracking entries.
Manuscript impact: May save chapter text updates when user chooses to persist generated/edited chapter content.
Safety level: Modifies data.
When to use: After drafting a substantial scene to keep memory and continuity synchronized.""",
    "/proofread": """Purpose: Rewrite pasted text into clean, publication-ready novel formatting with UK English corrections.
Files read: User-pasted text only.
Files written: None.
AI usage: Yes.
Canon memory impact: None.
Manuscript impact: None unless user manually pastes rewritten output into project files.
Safety level: Safe.
When to use: Immediately after writing a scene/chapter draft before continuity processing. Supports /proofread nocopy to skip clipboard copy.""",
    "/proofread nocopy": """Purpose: Rewrite pasted text into clean, publication-ready novel formatting without clipboard auto-copy.
Files read: User-pasted text only.
Files written: None.
AI usage: Yes.
Canon memory impact: None.
Manuscript impact: None unless user manually pastes rewritten output into project files.
Safety level: Safe.
When to use: Same as /proofread when clipboard copy should be skipped.""",
    "/rebuild-memory": """Purpose: Reconstruct canon memory from existing chapter files after major edits.
Files read: Chapter files, existing canon memory, and related project memory artifacts.
Files written: Canon memory, story state memory, timeline threads, scene summaries, and rebuild logs.
AI usage: Yes.
Canon memory impact: Replaces and rebuilds canonical continuity records.
Manuscript impact: Does not edit manuscript prose directly.
Safety level: Modifies data.
When to use: After large rewrites, chapter reordering, or memory drift concerns.""",
    "/rebuild-memory full": """Purpose: Force a full-novel canon memory rebuild from existing chapter files.
Files read: All chapter files, existing canon memory, and related project memory artifacts.
Files written: Canon memory, story state memory, timeline threads, scene summaries, and rebuild logs.
AI usage: Yes.
Canon memory impact: Replaces and rebuilds canonical continuity records across the full novel.
Manuscript impact: Does not edit manuscript prose directly.
Safety level: Modifies data.
When to use: After large rewrites spanning multiple chapters or major structural reordering.""",
    "/rebuild-memory single": """Purpose: Rebuild canon memory with single-chapter targeting.
Files read: Selected chapter file, existing canon memory, and related project memory artifacts.
Files written: Canon memory, story state memory, timeline threads, scene summaries, and rebuild logs.
AI usage: Yes.
Canon memory impact: Updates continuity records with targeted single-chapter rebuild intent.
Manuscript impact: Does not edit manuscript prose directly.
Safety level: Modifies data.
When to use: After heavy edits to one chapter where full novel rebuild is unnecessary.""",
    "/continuity-check": """Purpose: Compare selected chapter content against memory/context to detect factual continuity issues.
Files read: Canon memory, selected chapter text, and nearby chapter context.
Files written: Continuity report file.
AI usage: Yes.
Canon memory impact: None.
Manuscript impact: None.
Safety level: Safe.
When to use: Before locking a chapter or progressing to the next chapter.""",
    "/story-state": """Purpose: Display current persistent unresolved and resolved narrative pressure states.
Files read: Story state memory file.
Files written: None.
AI usage: No.
Canon memory impact: None.
Manuscript impact: None.
Safety level: Safe.
When to use: When you need fast orientation on active narrative pressure.""",
    "/timeline-view": """Purpose: Display timeline-oriented continuity entries in chronological-style view.
Files read: Canon memory and timeline-thread memory files.
Files written: None.
AI usage: No.
Canon memory impact: None.
Manuscript impact: None.
Safety level: Safe.
When to use: When sequence, timing, or event order is unclear.""",
    "/chapter-summary": """Purpose: Summarize a selected chapter’s key movement and outcomes.
Files read: Chapter file selected by the user.
Files written: None.
AI usage: Yes.
Canon memory impact: None.
Manuscript impact: None.
Safety level: Safe.
When to use: During revision planning and chapter-level structural review.""",
    "/ideas": """Purpose: Save a raw idea note into persistent idea storage.
Files read: Existing ideas file (if present).
Files written: Ideas memory file.
AI usage: No.
Canon memory impact: None.
Manuscript impact: None.
Safety level: Modifies data.
When to use: Any time you want to capture inspiration for later.""",
    "/ideas --list": """Purpose: View all saved ideas exactly as currently stored.
Files read: Ideas memory file.
Files written: None.
AI usage: No.
Canon memory impact: None.
Manuscript impact: None.
Safety level: Safe.
When to use: Any time you want to review previously saved ideas.""",
    "/idea-resurface": """Purpose: Retrieve and rank previously saved ideas relevant to current context.
Files read: Ideas memory file and optional user prompt text.
Files written: None.
AI usage: Yes.
Canon memory impact: None.
Manuscript impact: None.
Safety level: Safe.
When to use: When story progress stalls and prior idea inventory may help.""",
    "/inspiration": """Purpose: Analyse writing techniques in a pasted scene against structured inspiration categories.
Files read: Inspiration category files under ~/writing/novel_project/inspirations/<book>/ and user-pasted scene text.
Files written: None.
AI usage: Yes.
Canon memory impact: None.
Manuscript impact: None.
Safety level: Safe.
When to use: When you want craft-level technique feedback (prose/dialogue/description/pacing/tension/devices) without plot or rewrite suggestions.""",
    "/world-add": """Purpose: Append a confirmed world rule/fact to world memory storage.
Files read: Existing world memory file.
Files written: World memory file.
AI usage: No.
Canon memory impact: Adds persistent world canon notes.
Manuscript impact: None.
Safety level: Modifies data.
When to use: After confirming stable world constraints you want consistently enforced.""",
    "/draft-save": """Purpose: Save a timestamped backup snapshot of the manuscript.
Files read: Current manuscript file.
Files written: Draft/backup files.
AI usage: No.
Canon memory impact: None.
Manuscript impact: Creates backup only; does not alter manuscript text.
Safety level: Modifies data.
When to use: Before risky edits or major rewrite passes.""",
    "/draft-list": """Purpose: List available saved manuscript draft snapshots.
Files read: Draft/backup directory contents.
Files written: None.
AI usage: No.
Canon memory impact: None.
Manuscript impact: None.
Safety level: Safe.
When to use: Before restoring or auditing backup history.""",
    "/draft-load": """Purpose: Restore manuscript content from a selected draft snapshot.
Files read: Draft/backup files.
Files written: Manuscript file (restored content).
AI usage: No.
Canon memory impact: None.
Manuscript impact: Overwrites current manuscript content with selected snapshot.
Safety level: Destructive.
When to use: When you need rollback to a known stable manuscript state.""",
    "/draft-pass": """Purpose: Run focused AI revision diagnostics on manuscript text using a selected pass mode.
Files read: Manuscript text and user-selected mode options.
Files written: None.
AI usage: Yes.
Canon memory impact: None.
Manuscript impact: None unless user manually applies recommendations.
Safety level: Safe.
When to use: During revision phases focused on structure, tension, or clarity.""",
    "/book-integrity": """Purpose: Audit full manuscript for structural, continuity, arc, and tension integrity.
Files read: Full manuscript and relevant continuity context.
Files written: Book integrity report file.
AI usage: Yes.
Canon memory impact: None.
Manuscript impact: None.
Safety level: Safe.
When to use: Periodically at major milestones or pre-distribution review.""",
    "/build-book": """Purpose: Compile ordered chapter files into a single manuscript text artifact.
Files read: Chapter files.
Files written: Manuscript file.
AI usage: No.
Canon memory impact: None.
Manuscript impact: Regenerates manuscript from chapters.
Safety level: Modifies data.
When to use: After chapter updates to produce a current full-book draft.""",
    "/export-book --docx": """Purpose: Export manuscript/chapter compilation into DOCX format for sharing or submission.
Files read: Chapter files or manuscript source used by exporter.
Files written: DOCX export file in manuscript directory.
AI usage: No.
Canon memory impact: None.
Manuscript impact: No source text changes; generates export artifact.
Safety level: Modifies data.
When to use: When preparing editor/agent/beta-reader distribution output.""",
    "/world-consistency": """Purpose: Audit whole-book world-logic consistency in chunks and synthesize issues.
Files read: Full manuscript/book text.
Files written: World consistency report file.
AI usage: Yes.
Canon memory impact: None.
Manuscript impact: None.
Safety level: Safe.
When to use: During continuity QA for setting, technology, and rules.""",
    "/character-consistency": """Purpose: Audit character behaviour/motivation continuity across manuscript chunks.
Files read: Full manuscript/book text.
Files written: Character consistency report file.
AI usage: Yes.
Canon memory impact: None.
Manuscript impact: None.
Safety level: Safe.
When to use: During revision when character logic drift is suspected.""",
    "/recap": """Purpose: Generate a present-state orientation recap from current narrative context.
Files read: Canon memory and story state context.
Files written: None.
AI usage: Yes.
Canon memory impact: None.
Manuscript impact: None.
Safety level: Safe.
When to use: When re-entering project context or resetting narrative focus.""",
    "/research-topic": """Purpose: Run structured scientific research capture on a user-defined topic.
Files read: User research prompt and existing research topic files (for conflict checks as applicable).
Files written: Research topic files.
AI usage: Yes.
Canon memory impact: None.
Manuscript impact: None.
Safety level: Modifies data.
When to use: Before writing scenes that require technical realism grounding.""",
    "/research-scene": """Purpose: Evaluate pasted scene realism against scientific/engineering constraints.
Files read: User-pasted scene and relevant research context.
Files written: None.
AI usage: Yes.
Canon memory impact: None.
Manuscript impact: None.
Safety level: Safe.
When to use: After drafting technically sensitive scenes.""",
    "/research-apply": """Purpose: Apply saved research topics to evaluate a scene for realism conflicts.
Files read: Stored research topics and user-pasted scene.
Files written: None.
AI usage: Yes.
Canon memory impact: None.
Manuscript impact: None.
Safety level: Safe.
When to use: During revision when strict alignment with saved research is required.""",
    "/research-integrity": """Purpose: Cross-check stored research topics for internal contradiction.
Files read: Research topic files.
Files written: Research integrity report files.
AI usage: Yes.
Canon memory impact: None.
Manuscript impact: None.
Safety level: Modifies data.
When to use: Before consolidating world rules that depend on multiple research topics.""",
    "/research --world": """Purpose: Analyse world.txt for plausibility, consistency, and believability.
Files read: Project root world.txt.
Files written: World plausibility report files.
AI usage: Yes.
Canon memory impact: None.
Manuscript impact: None.
Safety level: Modifies data.
When to use: During worldbuilding QA before drafting or revision.""",
    "/help": """Purpose: Show available command list.
Files read: None.
Files written: None.
AI usage: No.
Canon memory impact: None.
Manuscript impact: None.
Safety level: Safe.
When to use: Any time command discovery is needed.""",
    "/help --workflow": """Purpose: Show static end-to-end workflow guidance for operational command cadence.
Files read: None.
Files written: None.
AI usage: No.
Canon memory impact: None.
Manuscript impact: None.
Safety level: Safe.
When to use: When choosing workflow phase and tool order.""",
    "/help --when": """Purpose: Show static per-command usage timing guidance.
Files read: None.
Files written: None.
AI usage: No.
Canon memory impact: None.
Manuscript impact: None.
Safety level: Safe.
When to use: When deciding which command fits current writing situation.""",
    "/system --health": """Purpose: Display local project health diagnostics and risk signals.
Files read: Core project files/directories and memory/manuscript metadata.
Files written: None.
AI usage: No.
Canon memory impact: None.
Manuscript impact: None.
Safety level: Safe.
When to use: Before long sessions or when system/performance drift is suspected.""",
    "/system": """Purpose: Base system namespace command for system diagnostics/options.
Files read: Depends on selected option.
Files written: None.
AI usage: No.
Canon memory impact: None.
Manuscript impact: None.
Safety level: Safe.
When to use: Use with /system --tree, /system --map, or /system --health.""",
    "/help --describe": """Purpose: Interactive static command manual lookup by numbered command selection.
Files read: Built-in ALL_COMMANDS and COMMAND_HELP constants only.
Files written: None.
AI usage: No.
Canon memory impact: None.
Manuscript impact: None.
Safety level: Safe.
When to use: When you need precise operational behavior for a specific command.""",
    "exit": """Purpose: Terminate the interactive assistant session.
Files read: None.
Files written: None.
AI usage: No.
Canon memory impact: None.
Manuscript impact: None.
Safety level: Safe.
When to use: When ending the current terminal session.""",
}


def _extract_when_to_use(help_text: str) -> str:
    """Extract the 'When to use' guidance line from a help text block."""
    for line in help_text.splitlines():
        stripped = line.strip()
        if stripped.lower().startswith("when to use:"):
            return stripped.split(":", 1)[1].strip()
    return ""


COMMAND_DESCRIPTIONS: dict[str, str] = dict(COMMAND_HELP_DETAILS)
COMMAND_WHEN: dict[str, str] = {
    command_name: _extract_when_to_use(help_text)
    for command_name, help_text in COMMAND_HELP_DETAILS.items()
}


def handle_help_describe() -> None:
    """Show descriptions for all discovered commands."""
    print("COMMAND DESCRIPTIONS")
    print()
    for command_name in ALL_COMMANDS:
        print(command_name)
        print(COMMAND_DESCRIPTIONS.get(command_name, "No description available"))
        print()


def _extract_help_generation_signals(command_func: Callable[..., Any]) -> dict[str, Any]:
    """Collect behavior signals from command function source for description generation."""
    source = inspect.getsource(command_func)
    source_lower = source.lower()

    printed_headers = re.findall(r'print\((?:"|\')([^"\']+)(?:"|\')\)', source)
    header_markers = [header for header in printed_headers if "===" in header or "DIRECTORY" in header.upper()]

    directory_constant_refs = sorted(set(re.findall(r"\b[A-Z_]+_DIR\b", source)))
    literal_directory_refs = sorted(
        set(
            match
            for match in re.findall(r'(?:"|\')([a-zA-Z0-9_/\-]+/)(?:"|\')', source)
            if "/" in match
        )
    )

    architecture_terms = sorted(
        {
            term
            for term in ("architecture", "map", "reads", "writes", "debug", "onboarding", "data flow")
            if term in source_lower
        }
    )

    return {
        "function_name": command_func.__name__,
        "headers": header_markers,
        "directory_constants": directory_constant_refs,
        "literal_directories": literal_directory_refs,
        "architecture_terms": architecture_terms,
    }


def _generate_system_introspection_help(command_func: Callable[..., Any]) -> str:
    """Generate static help text for system introspection commands from function behavior."""
    signals = _extract_help_generation_signals(command_func)
    function_name = signals["function_name"]

    emits_tree_output = "tree" in function_name or "tree_lines" in inspect.getsource(command_func)
    emits_architecture_map = "map" in function_name or "architecture" in " ".join(signals["architecture_terms"])

    discovered_directories = signals["directory_constants"] + signals["literal_directories"]
    directory_scope = "multiple project directories" if discovered_directories else "project paths when available"

    header_phrase = ""
    if signals["headers"]:
        cleaned_headers = ", ".join(header.strip("= ").lower() for header in signals["headers"][:2])
        header_phrase = f" based on printed sections such as {cleaned_headers}"

    if emits_tree_output:
        purpose = (
            "Display the project directory tree and explain the role of each directory"
            f"{header_phrase}."
        )
        when_to_use = "For architecture transparency, onboarding, and filesystem debugging."
    elif emits_architecture_map:
        purpose = (
            "Display the logical architecture map of command interactions with project directories"
            f"{header_phrase}."
        )
        when_to_use = "For pipeline debugging and system behavior visualization."
    else:
        purpose = f"Display system introspection output derived from {function_name}{header_phrase}."
        when_to_use = "For developer introspection and operational debugging."

    files_read = (
        f"Function source-defined metadata for {directory_scope}."
        if discovered_directories
        else "Function source-defined metadata only."
    )

    return "\n".join(
        [
            f"Purpose: {purpose}",
            "Command category: System Introspection.",
            f"Files read: {files_read}",
            "Files written: None.",
            "AI usage: No.",
            "Canon memory impact: None.",
            "Manuscript impact: None.",
            "Safety level: Safe.",
            f"When to use: {when_to_use}",
        ]
    )


AVAILABLE_COMMAND_FUNCTIONS: list[Callable[..., Any]] = []
HELP_COMMAND_DISCOVERY: OrderedDict[str, Callable[..., Any]] = OrderedDict()


def rebuild_help_descriptions() -> None:
    """Rebuild generated static help descriptions from discovered command functions."""
    for command_name, command_func in HELP_COMMAND_DISCOVERY.items():
        generated_help = _generate_system_introspection_help(command_func)
        COMMAND_DESCRIPTIONS[command_name] = generated_help
        COMMAND_WHEN[command_name] = _extract_when_to_use(generated_help)


HELP_SECTION_ORDER: tuple[str, ...] = (
    "System",
    "Exports",
    "Drafting",
    "Writing",
    "Continuity & Integrity",
    "World",
    "Help & Workflow",
)


HELP_SECTION_TAXONOMY: OrderedDict[str, list[str]] = OrderedDict(
    [
        (
            "System",
            [
                "exit",
                "/system",
                "/system --tree",
                "/system --map",
                "/system --health",
                "/novel-stats",
                "/story-state",
            ],
        ),
        (
            "Exports",
            [
                "/build-book",
                "/export-book",
                "/export-book --docx",
                "/export-chapter",
            ],
        ),
        (
            "Drafting",
            [
                "/draft-list",
                "/draft-load",
                "/draft-pass",
                "/draft-save",
                "/drafts",
                "/restore-draft",
                "/save-draft",
            ],
        ),
        (
            "Writing",
            [
                "/chapter-summary",
                "/scene-summary",
                "/recap",
                "/proofread",
                "/proofread nocopy",
                "/ideas",
                "/ideas --list",
                "/idea-resurface",
                "/inspiration",
            ],
        ),
        (
            "Continuity & Integrity",
            [
                "/continuity-check",
                "/book-integrity",
                "/character-consistency",
                "/world-consistency",
                "/timeline-view",
                "/rebuild-summaries",
                "/rebuild-memory",
                "/rebuild-memory full",
                "/rebuild-memory single",
            ],
        ),
        (
            "World",
            [
                "/research-topic",
                "/research-scene",
                "/research-apply",
                "/research-integrity",
                "/research",
                "/research --world",
                "/world-add",
            ],
        ),
        (
            "Help & Workflow",
            [
                "/help",
                "/help --describe",
                "/help --when",
                "/help --workflow",
            ],
        ),
    ]
)


def _print_columns(items: list[str], width: int) -> None:
    """Print items in terminal-width-aware columns."""
    if not items:
        return

    col_width = 28
    num_cols = max(1, width // col_width)
    rows = math.ceil(len(items) / num_cols)

    for row_index in range(rows):
        line = ""
        for col_index in range(num_cols):
            item_index = col_index * rows + row_index
            if item_index < len(items):
                line += items[item_index].ljust(col_width)
        print(line.rstrip())


def _categorize_command(command_name: str) -> str:
    """Categorize commands using rule-based keyword matching."""
    lowered = command_name.lower()
    category_rules: list[tuple[str, tuple[str, ...]]] = [
        ("System", ("system", "stats", "state")),
        ("Exports", ("export",)),
        ("Drafting", ("draft",)),
        ("Writing", ("proofread", "inspiration", "ideas")),
        ("Continuity & Integrity", ("continuity", "consistency", "integrity", "rebuild")),
        ("World", ("world", "research")),
        ("Help & Workflow", ("help",)),
    ]
    for category, keywords in category_rules:
        if any(keyword in lowered for keyword in keywords):
            return category
    return "Other"


def _build_help_sections() -> OrderedDict[str, OrderedDict[str, list[str]]]:
    """Build dynamic help sections from discovered command handlers."""
    sections: OrderedDict[str, OrderedDict[str, list[str]]] = OrderedDict(
        (section, OrderedDict()) for section in HELP_SECTION_ORDER
    )
    sections["Other"] = OrderedDict()

    command_set = sorted(set(ALL_COMMANDS), key=lambda item: (item.split()[0], item))
    for command_name in command_set:
        pieces = command_name.split(maxsplit=1)
        base_command = pieces[0]
        variant = pieces[1] if len(pieces) > 1 else ""

        category = _categorize_command(command_name)
        category_map = sections.setdefault(category, OrderedDict())
        category_map.setdefault(base_command, [])

        if variant and base_command in command_set:
            if variant not in category_map[base_command]:
                category_map[base_command].append(variant)
        elif command_name != base_command:
            category_map.setdefault(command_name, [])

    return sections


def _print_dynamic_help_catalog(mode: str = "list") -> None:
    """Print dynamic command catalog with optional description/when details."""
    sections = _build_help_sections()
    print("=== NOVEL AI COMMANDS ===")
    print()

    for section_title, grouped_commands in sections.items():
        if not grouped_commands:
            continue
        print(f"[ {section_title} ]")
        for base_command, variants in grouped_commands.items():
            print(f"- {base_command}")
            if mode == "describe":
                print(f"    {COMMAND_DESCRIPTIONS.get(base_command, 'No description available')}")
            elif mode == "when":
                print(f"    {COMMAND_WHEN.get(base_command, '') or 'No usage guidance available'}")
            for variant in sorted(variants):
                full_variant_command = f"{base_command} {variant}"
                print(f"  ├── {variant}")
                if mode == "describe":
                    print(f"      {COMMAND_DESCRIPTIONS.get(full_variant_command, 'No description available')}")
                elif mode == "when":
                    print(f"      {COMMAND_WHEN.get(full_variant_command, '') or 'No usage guidance available'}")
        print()


def print_help() -> None:
    """Show dynamically discovered command list."""
    _print_dynamic_help_catalog(mode="list")


def update_help_commands_from_handlers(command_handlers: dict[str, Callable[[str], None]]) -> None:
    """Synchronize help-menu command list from active command handlers."""
    global ALL_COMMANDS
    ALL_COMMANDS = sorted(set(command_handlers.keys()), key=lambda item: (item.split()[0], item))


def validate_help_taxonomy(command_handlers: dict[str, Callable[[str], None]]) -> None:
    """Retained for compatibility; dynamic help no longer depends on static taxonomy."""
    del command_handlers


def command_matches_input(command_name: str, user_input: str) -> bool:
    """Flexible command routing with prefix support for option-bearing commands."""
    normalized_user_input = user_input.strip()
    if normalized_user_input == command_name:
        return True
    if normalized_user_input.startswith(f"{command_name} "):
        return True
    if " " in command_name and normalized_user_input.startswith(command_name):
        return True
    return False


def handle_help_workflow() -> None:
    """Show the static master writing workflow guidance."""
    print(
        """==================================================
NOVEL AI MASTER WRITING WORKFLOW
==================================================

PHASE 1 — RAW WRITING

Write freely inside WordGrinder.
Do NOT interrupt flow with AI.

Export chapter text as:

chapter_<number>.txt

--------------------------------------------------

PHASE 2 — SESSION CLEANUP

/proofread
→ full rewrite for grammar, punctuation, UK spelling
→ novel paragraph/dialogue formatting
→ auto-copy clean text to clipboard (use /proofread nocopy to skip)

/scene-summary
→ extract canon facts
→ update story state
→ maintain continuity memory

/chapter-summary
→ understand structural movement
→ tension change
→ resolved threads

/continuity-check
→ detect timeline, injury, location, world contradictions

--------------------------------------------------

PHASE 3 — STORY CONTROL (STRATEGIC USE)

Use only when needed.

/recap
→ narrative re-orientation
→ remember current story position

/story-state
→ view active unresolved arcs

/timeline-view
→ understand event order across novel

/idea-resurface
→ reintroduce previously saved ideas

--------------------------------------------------

PHASE 4 — DRAFT SAFETY + REVISION

/draft-save
→ create recoverable manuscript snapshot

/draft-list
→ view stored drafts

/draft-load
→ restore earlier manuscript state

/draft-pass
→ focused rewrite diagnostics
→ structure / tension / clarity / character logic

--------------------------------------------------

PHASE 5 — LARGE SCALE STORY AUDITS

Use occasionally, not daily.

/book-integrity
→ full-novel structural health check

/world-consistency
→ science-fiction world logic audit

/character-consistency
→ psychological behaviour audit

/rebuild-memory
or /rebuild-memory full
or /rebuild-memory single
→ resynchronise canon memory after major rewrites

/rebuild-summaries
→ regenerate chapter summary artifacts after major chapter edits

--------------------------------------------------

PHASE 6 — SCIENTIFIC REALISM ENGINE

/research-topic
→ pure real-world science research

/research-scene
→ scientific realism analysis for pasted scene

/research-apply
→ check scene against saved research

/research-integrity
→ detect contradictions across research topics

/research --world
→ world.txt plausibility and consistency audit

--------------------------------------------------

PHASE 7 — MANUSCRIPT CREATION

/build-book
→ compile readable full manuscript

/export-book --docx
→ create professional submission manuscript

--------------------------------------------------

GOLDEN WRITING RHYTHM

WRITE  
→ PROOFREAD  
→ SCENE SUMMARY  
→ CONTINUITY CHECK  
→ CONTINUE WRITING  

Strategic commands only when necessary.

=================================================="""
    )
    print()
    print("LIVE COMMAND INDEX")
    _print_dynamic_help_catalog(mode="list")


def handle_help_when() -> None:
    """Show per-command usage guidance for all discovered commands."""
    print("=== NOVEL AI — WHEN TO USE COMMANDS ===")
    print()
    _print_dynamic_help_catalog(mode="when")


def handle_help_generate(client: Any, command_text: str) -> None:
    """Generate and persist command description and usage guidance for a command."""
    requested_command = command_text.replace("/help --generate", "", 1).strip()
    if not requested_command:
        print("Usage: /help --generate <command>")
        return
    if requested_command not in ALL_COMMANDS:
        print(f"Unknown command: {requested_command}")
        return

    messages = [
        {
            "role": "system",
            "content": (
                "Generate concise command documentation. Return exactly two lines:\n"
                "DESCRIPTION: <text>\nWHEN: <text>"
            ),
        },
        {"role": "user", "content": f"Command: {requested_command}"},
    ]
    try:
        response = request_chat_completion(
            client=client,
            messages=messages,
            temperature=0.2,
            loading_message=f"Generating help for {requested_command}...",
        )
    except Exception as exc:
        print(f"Help generation failed: {exc}")
        return

    description_match = re.search(r"^DESCRIPTION:\s*(.+)$", response, re.IGNORECASE | re.MULTILINE)
    when_match = re.search(r"^WHEN:\s*(.+)$", response, re.IGNORECASE | re.MULTILINE)
    if not description_match or not when_match:
        print("Generation failed: invalid model output.")
        return

    COMMAND_DESCRIPTIONS[requested_command] = description_match.group(1).strip()
    COMMAND_WHEN[requested_command] = when_match.group(1).strip()
    print(f"Generated help saved for {requested_command}.")


def handle_export_book(command_text: str = "") -> None:
    """Route /export-book options to concrete export handlers."""
    if "--docx" in command_text:
        handle_export_book_docx()
        return
    print("Unsupported export option. Use /export-book --docx")


def command_system_tree() -> None:
    """Display the canonical project directory tree and folder purposes."""
    print("=== NOVEL AI SYSTEM STRUCTURE ===")
    print()

    tree_lines = [
        f"{WRITING_DIR.name}/",
        f"└── {NOVEL_PROJECT_DIR.name}/",
        f"    ├── {CHAPTERS_DIR.name}/",
        f"    ├── {MANUSCRIPT_DIR.name}/",
        f"    ├── {DRAFTS_DIR.name}/",
        "    ├── autosaves/",
        f"    ├── {PROJECT_MEMORY_DIR.name}/",
        f"    │   └── {CANON_MEMORY_BACKUPS_DIR.name}/",
        f"    ├── {PROJECT_ANALYSIS_DIR.name}/",
        f"    │   ├── {CONTINUITY_REPORTS_DIR.name}/",
        f"    │   ├── {TIMELINE_LOGS_DIR.name}/",
        f"    │   ├── {BOOK_INTEGRITY_REPORTS_DIR.name}/",
        f"    │   └── {REBUILD_LOG_DIR.name}/",
        f"    ├── {RESEARCH_DIR.name}/",
        f"    │   └── {RESEARCH_INTEGRITY_REPORTS_DIR.name}/",
        f"    ├── {FULL_NOVEL_PROCESSOR_LOG_DIR.name}/",
        f"    ├── {PROJECT_BACKUPS_DIR.name}/",
        "    └── sources/",
    ]
    for line in tree_lines:
        print(line)

    print()
    print("=== DIRECTORY PURPOSE ===")
    print()

    directory_descriptions = OrderedDict(
        [
            ("chapters/", "Stores finalized individual chapter files."),
            ("manuscript/", "Stores compiled full novel manuscript outputs."),
            ("drafts/", "Creative workspace for experimental or unfinished writing."),
            (
                "autosaves/",
                "Automatic safety saves created during writing or processor activity.",
            ),
            ("memory/", "Persistent AI canon memory (characters, world, plot state)."),
            ("memory/backups/", "Safety snapshots of AI memory before modification."),
            (
                "analysis/",
                "Outputs from timeline engines, continuity scanners and rebuild systems.",
            ),
            ("analysis/continuity_reports/", "Canon contradiction detection reports."),
            ("analysis/timeline_logs/", "Chronological tracking of story events."),
            (
                "analysis/book_integrity_reports/",
                "Structural manuscript health reports.",
            ),
            (
                "analysis/rebuild_logs/",
                "Logs from rebuild engines altering novel structure.",
            ),
            ("research/", "Worldbuilding and factual research storage."),
            (
                "research/integrity_reports/",
                "Reports checking research consistency with canon.",
            ),
            (
                "logs/",
                "System execution logs for processor runs and major operations.",
            ),
            (
                "backups/",
                "Full project restore snapshots for disaster recovery.",
            ),
            ("sources/", "Imported external text or screenplay material."),
        ]
    )
    for directory_name, description in directory_descriptions.items():
        print(f"{directory_name}\n  {description}")


def command_system_map() -> None:
    """Display a logical map of command read/write relationships by project directory."""
    print("=== NOVEL AI SYSTEM ARCHITECTURE MAP ===")
    print()

    architecture_sections = [
        (
            "CHAPTERS DIRECTORY",
            [
                ("Writes", ["chapter generation commands", "draft promotion system", "rebuild chapter engine"]),
                ("Reads", ["manuscript compiler", "timeline engine", "continuity scanner"]),
            ],
        ),
        (
            "MANUSCRIPT DIRECTORY",
            [
                ("Writes", ["manuscript compile command", "rebuild engine", "processor full run"]),
                ("Reads", ["integrity scanner", "backup system", "export systems"]),
            ],
        ),
        (
            "DRAFTS DIRECTORY",
            [
                ("Writes", ["AI writing generation", "experimental scene builder", "processor intermediate outputs"]),
                ("Reads", ["draft promotion system", "draft comparison tools", "tone analysis engine"]),
            ],
        ),
        (
            "AUTOSAVES DIRECTORY",
            [
                ("Writes", ["automatic save triggers during writing", "processor safety checkpoints"]),
                ("Reads", ["recovery tools", "restore assistant"]),
            ],
        ),
        (
            "MEMORY DIRECTORY",
            [
                ("Writes", ["memory processor", "continuity repair engine", "canon update commands"]),
                ("Reads", ["writing generation prompts", "timeline builder", "research alignment tools"]),
            ],
        ),
        (
            "MEMORY BACKUPS DIRECTORY",
            [
                ("Writes", ["memory update safeguard", "rebuild memory engine"]),
                ("Reads", ["memory recovery system", "continuity rollback logic"]),
            ],
        ),
        (
            "ANALYSIS DIRECTORY",
            [
                ("Writes", ["timeline builder", "continuity scanner", "integrity engine", "rebuild diagnostics"]),
                ("Reads", ["system health command", "developer inspection tools"]),
            ],
        ),
        (
            "RESEARCH DIRECTORY",
            [
                ("Writes", ["research ingestion commands", "worldbuilding tools"]),
                ("Reads", ["writing prompt construction", "lore validation systems"]),
            ],
        ),
        (
            "LOGS DIRECTORY",
            [
                ("Writes", ["processor runs", "rebuild events", "integrity scans", "major command execution"]),
                ("Reads", ["system health reporting", "debug tools"]),
            ],
        ),
        (
            "BACKUPS DIRECTORY",
            [
                ("Writes", ["backup command", "rebuild safeguard", "processor pre-operation snapshot"]),
                ("Reads", ["restore system", "corruption recovery engine"]),
            ],
        ),
        (
            "SOURCES DIRECTORY",
            [
                ("Writes", ["screenplay import", "external text ingestion"]),
                ("Reads", ["conversion engines", "writing inspiration tools"]),
            ],
        ),
    ]

    for index, (directory_header, relationships) in enumerate(architecture_sections):
        print(directory_header)
        for relationship_label, command_groups in relationships:
            print(f"{relationship_label}:")
            for command_group in command_groups:
                print(f"- {command_group}")
            print()
        if index != len(architecture_sections) - 1:
            print("-" * 50)
            print()


AVAILABLE_COMMAND_FUNCTIONS.append(command_system_tree)
AVAILABLE_COMMAND_FUNCTIONS.append(command_system_map)
HELP_COMMAND_DISCOVERY["/system --tree"] = command_system_tree
HELP_COMMAND_DISCOVERY["/system --map"] = command_system_map


def handle_system(command_text: str = "") -> None:
    """Route /system options to concrete system handlers."""
    if "--health" in command_text:
        handle_system_health()
        return
    if "--map" in command_text:
        command_system_map()
        return
    if "--tree" in command_text:
        command_system_tree()
        return
    print("Unsupported system option. Use /system --tree, /system --map, or /system --health")


LONG_RUNNING_COMMAND_PREFIXES: tuple[str, ...] = (
    "/scene-summary",
    "/recap",
    "/chapter-summary",
    "/rebuild-summaries",
    "/rebuild-memory",
    "/continuity-check",
    "/book-integrity",
    "/world-consistency",
    "/character-consistency",
    "/proofread",
    "/research-topic",
    "/research-scene",
    "/research-apply",
    "/research-integrity",
    "/research",
    "/idea-resurface",
    "/draft-pass",
    "/build-book",
    "/inspiration",
)


def command_requires_loading(user_input: str) -> bool:
    """Return True when command should run under the global loading wrapper."""
    normalized = user_input.strip().lower()
    return any(normalized.startswith(prefix) for prefix in LONG_RUNNING_COMMAND_PREFIXES)


def execute_command_with_optional_loading(
    command_name: str,
    command_handler: Callable[[str], None],
    user_input: str,
) -> None:
    """Centralized command execution layer with reusable loading support."""
    if command_requires_loading(user_input):
        run_with_loading(f"Running {command_name}...", lambda: command_handler(user_input))
        return
    command_handler(user_input)


def main() -> None:
    """Run the terminal assistant."""
    ensure_project_files()

    try:
        client = create_client()
    except RuntimeError as exc:
        print(exc)
        return

    conversation_history: list[dict[str, str]] = []

    command_handlers: dict[str, Callable[[str], None]] = {
        "/scene-summary": lambda command_text="": handle_scene_summary(client),
        "/recap": lambda command_text="": handle_recap(client),
        "/chapter-summary": lambda command_text="": handle_chapter_summary(client),
        "/rebuild-summaries": lambda command_text="": handle_rebuild_summaries(client),
        "/rebuild-memory": lambda command_text="": handle_rebuild_memory(client, command_text),
        "/rebuild-memory full": lambda command_text="": handle_rebuild_memory(client, command_text),
        "/rebuild-memory single": lambda command_text="": handle_rebuild_memory(client, command_text),
        "/continuity-check": lambda command_text="": handle_continuity_check(client),
        "/book-integrity": lambda command_text="": handle_book_integrity(client),
        "/world-consistency": lambda command_text="": handle_world_consistency(client),
        "/character-consistency": lambda command_text="": handle_character_consistency(client),
        "/proofread": lambda command_text="": handle_proofread(client, command_text),
        "/proofread nocopy": lambda command_text="": handle_proofread(client, command_text),
        "/research-topic": lambda command_text="": handle_research_topic(client),
        "/research-scene": lambda command_text="": handle_research_scene(client),
        "/research-apply": lambda command_text="": handle_research_apply(client),
        "/research-integrity": lambda command_text="": handle_research_integrity(client),
        "/research": lambda command_text="": handle_research(client, command_text),
        "/research --world": lambda command_text="": handle_research(client, command_text),
        "/idea-resurface": lambda command_text="": handle_idea_resurface(client),
        "/inspiration": lambda command_text="": handle_inspiration(client),
        "/draft-pass": lambda command_text="": handle_draft_pass(client, command_text),
        "/build-book": lambda command_text="": handle_build_book(),
        "/draft-save": lambda command_text="": handle_draft_save(),
        "/draft-list": lambda command_text="": handle_draft_list(),
        "/draft-load": lambda command_text="": handle_draft_load(),
        "/save-draft": lambda command_text="": handle_draft_save(),
        "/drafts": lambda command_text="": handle_draft_list(),
        "/restore-draft": lambda command_text="": handle_draft_load(),
        "/ideas --list": handle_ideas,
        "/ideas": handle_ideas,
        "/world-add": lambda command_text="": handle_world_add(),
        "/timeline-view": lambda command_text="": handle_timeline_view(),
        "/story-state": lambda command_text="": handle_story_state(),
        "/export-chapter": lambda command_text="": handle_export_chapter(),
        "/system --map": handle_system,
        "/system --tree": handle_system,
        "/system": handle_system,
        "/export-book": handle_export_book,
        "/export-book --docx": handle_export_book,
        "/novel-stats": lambda command_text="": handle_novel_stats(),
        "/system --health": handle_system,
        "/help --describe": lambda command_text="": handle_help_describe(),
        "/help --generate": lambda command_text="": handle_help_generate(client, command_text),
        "/help --workflow": lambda command_text="": handle_help_workflow(),
        "/help --when": lambda command_text="": handle_help_when(),
        "/help": lambda command_text="": print_help(),
    }
    update_help_commands_from_handlers(command_handlers)
    validate_help_taxonomy(command_handlers)
    rebuild_help_descriptions()

    print_welcome()

    while True:
        user_input = ""
        try:
            user_input = input("\nYou: ").strip()
        except EOFError:
            print("\nGoodbye.")
            break
        except KeyboardInterrupt:
            print("\nGoodbye.")
            break

        if not user_input:
            continue

        if user_input.lower() == "exit":
            print("Goodbye.")
            break


        routed = False
        for command_name in sorted(command_handlers.keys(), key=len, reverse=True):
            if command_matches_input(command_name, user_input):
                execute_command_with_optional_loading(
                    command_name,
                    command_handlers[command_name],
                    user_input,
                )
                routed = True
                break

        if routed:
            continue

        memory_block = load_memory_block()
        conversation_history = conversation_history[-MAX_CONVERSATION_TURNS * 2:]
        messages = build_main_messages(memory_block, conversation_history, user_input)

        try:
            reply = request_chat_completion(
                client=client,
                messages=messages,
                temperature=MAIN_TEMPERATURE,
            )
        except Exception as exc:  # Keep terminal app stable for the user.
            print(f"Assistant request failed: {exc}")
            continue

        print(f"\nAssistant: {reply}")

        conversation_history.append({"role": "user", "content": user_input})
        conversation_history.append({"role": "assistant", "content": reply})


if __name__ == "__main__":
    main()
